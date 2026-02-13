"""Alfred API Server - Main FastAPI application with auth, integrations, and tool calling."""

import logging
import os
import re
import sys
import uuid
from contextlib import asynccontextmanager
from datetime import datetime
from pathlib import Path

from fastapi import FastAPI, WebSocket, WebSocketDisconnect, UploadFile, File, Depends, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, StreamingResponse, Response, RedirectResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from slowapi import Limiter
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from slowapi.middleware import SlowAPIMiddleware

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from core.brain.router import ask, classify_query, ModelTier, get_system_prompt
from core.memory.store import store_memory, recall, store_conversation
from core.memory.conversations import (
    init_db as init_conversations_db,
    create_conversation, list_conversations as list_convos,
    get_conversation, add_message, archive_conversation, update_title,
    list_archived_conversations, restore_conversation, delete_conversation_permanently,
    search_conversations, create_project, list_projects, get_project,
    update_project, delete_project, add_reference, list_references,
    get_reference, update_reference, delete_reference, search_references,
    move_conversation_to_project, list_conversations_by_project, get_project_context,
    UPLOADS_DIR,
)

# TTS audio cache directory for Twilio voice calls
TTS_AUDIO_DIR = Path(__file__).parent.parent.parent / "data" / "audio" / "tts"
TTS_AUDIO_DIR.mkdir(parents=True, exist_ok=True)
from core.security.auth import (
    create_user, verify_user, create_access_token, get_current_user,
    require_auth, setup_initial_user,
    # 2FA and Passkey functions
    setup_totp, verify_totp, enable_totp, disable_totp, is_totp_enabled,
    get_passkey_registration_options, verify_passkey_registration,
    get_passkey_login_options, verify_passkey_login,
    list_passkeys, delete_passkey, get_user_auth_methods,
)
from core.security.google_oauth import get_auth_url, handle_callback, is_connected
from config.settings import settings

# Register all tools on import
from core.tools.definitions import register_all
register_all()

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(name)s] %(levelname)s: %(message)s")
logger = logging.getLogger("alfred")


@asynccontextmanager
async def lifespan(app: FastAPI):
    logger.info("Alfred is starting up...")
    logger.info(f"Local LLM: {settings.ollama_model} at {settings.ollama_host}")
    logger.info(f"Cloud LLM: {settings.anthropic_model}")
    # Initialize conversation history database
    init_conversations_db()
    logger.info("Conversation history database initialized")
    # Create initial admin user if needed
    pwd = setup_initial_user()
    if pwd:
        logger.info("Initial admin user created — check logs for password")
    # Pre-warm voice models in background thread so startup isn't blocked
    import asyncio
    loop = asyncio.get_event_loop()
    def _warmup_voice():
        try:
            from interfaces.voice.stt import warmup as stt_warmup
            stt_warmup(model_size=settings.whisper_model)
        except Exception as e:
            logger.warning(f"STT warmup failed: {e}")
        try:
            from interfaces.voice.tts import warmup as tts_warmup
            tts_warmup()
        except Exception as e:
            logger.warning(f"TTS warmup failed: {e}")
        # Pre-generate static Twilio voice phrases
        try:
            pregenerate_static_phrases()
        except Exception as e:
            logger.warning(f"Static TTS phrase pre-generation failed: {e}")
    loop.run_in_executor(None, _warmup_voice)
    # Initialize agent pool for multi-agent orchestration
    try:
        from core.orchestration.agents import initialize_agent_pool
        await initialize_agent_pool()
        logger.info("Agent pool initialized with 3 workers")
    except Exception as e:
        logger.warning(f"Agent pool initialization failed: {e}")
    yield
    # Shutdown agent pool
    try:
        from core.orchestration.agents import get_agent_pool
        pool = get_agent_pool()
        await pool.stop()
        logger.info("Agent pool stopped")
    except Exception:
        pass
    logger.info("Alfred is shutting down...")


limiter = Limiter(key_func=get_remote_address)

app = FastAPI(title="Alfred", version="0.3.0", lifespan=lifespan)
app.state.limiter = limiter
app.add_middleware(SlowAPIMiddleware)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["https://aialfred.groundrushcloud.com", "http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["Authorization", "Content-Type"],
)

# Static files for PWA icons
_static_dir = Path(__file__).parent.parent.parent / "static"
if _static_dir.exists():
    app.mount("/static", StaticFiles(directory=str(_static_dir)), name="static")

# Serve React frontend build if available
_frontend_dist = Path(__file__).parent.parent.parent / "frontend" / "dist"


@app.get("/manifest.json")
@app.get("/manifest.webmanifest")
async def pwa_manifest():
    # Serve Vite-generated manifest if available
    for name in ("manifest.webmanifest", "manifest.json"):
        vite_manifest = _frontend_dist / name
        if vite_manifest.exists():
            return Response(content=vite_manifest.read_text(), media_type="application/manifest+json",
                           headers={"Cache-Control": "public, max-age=86400"})
    return JSONResponse({
        "name": "Alfred",
        "short_name": "Alfred",
        "description": "Personal AI Assistant",
        "start_url": "/",
        "display": "standalone",
        "background_color": "#0a0a0a",
        "theme_color": "#0a0a0a",
        "orientation": "any",
        "icons": [
            {"src": "/static/icon-192.png", "sizes": "192x192", "type": "image/png", "purpose": "any maskable"},
            {"src": "/static/icon-512.png", "sizes": "512x512", "type": "image/png", "purpose": "any maskable"},
        ],
    }, headers={"Cache-Control": "public, max-age=86400"})


@app.get("/sw.js")
async def service_worker():
    # Serve Vite-generated service worker if available
    vite_sw = _frontend_dist / "sw.js"
    if vite_sw.exists():
        return Response(content=vite_sw.read_text(), media_type="application/javascript",
                        headers={"Cache-Control": "no-cache", "Service-Worker-Allowed": "/"})
    return Response(content=SERVICE_WORKER_JS, media_type="application/javascript",
                    headers={"Cache-Control": "no-cache", "Service-Worker-Allowed": "/"})


@app.get("/registerSW.js")
async def register_service_worker():
    """Serve Vite PWA registration script."""
    vite_reg = _frontend_dist / "registerSW.js"
    if vite_reg.exists():
        return Response(content=vite_reg.read_text(), media_type="application/javascript",
                        headers={"Cache-Control": "no-cache"})
    return Response(content="", media_type="application/javascript")


@app.get("/{filename}.png")
async def serve_root_png(filename: str):
    """Serve PNG files from frontend build root (PWA icons)."""
    for base in (_frontend_dist, _static_dir):
        f = base / f"{filename}.png"
        if f.exists():
            return Response(content=f.read_bytes(), media_type="image/png",
                           headers={"Cache-Control": "public, max-age=86400"})
    raise HTTPException(status_code=404)


@app.get("/{filename}.jpg")
async def serve_root_jpg(filename: str):
    """Serve JPG files from frontend build root."""
    for base in (_frontend_dist, _static_dir):
        f = base / f"{filename}.jpg"
        if f.exists():
            return Response(content=f.read_bytes(), media_type="image/jpeg",
                           headers={"Cache-Control": "public, max-age=86400"})
    raise HTTPException(status_code=404)


@app.get("/workbox-{filename}")
async def serve_workbox(filename: str):
    """Serve Workbox runtime files from frontend build."""
    workbox_file = _frontend_dist / f"workbox-{filename}"
    if workbox_file.exists():
        return Response(content=workbox_file.read_bytes(), media_type="application/javascript",
                        headers={"Cache-Control": "public, max-age=31536000"})
    raise HTTPException(status_code=404)


# ==================== Models ====================

class ChatRequest(BaseModel):
    message: str
    tier: str | None = None
    session_id: str = "default"
    # Optional image data for vision queries
    image_path: str | None = None
    image_base64: str | None = None
    image_media_type: str | None = None
    # Optional document path for analysis
    document_path: str | None = None
    # Optional TTS settings (for Shortcuts/API clients)
    voice: str | None = None  # Voice ID (e.g., "Gwen_Stacy", "design:Natalie")
    tts_backend: str | None = None  # TTS backend ("kokoro", "qwen3")

class ChatResponse(BaseModel):
    response: str
    tier: str
    timestamp: str
    images: list[dict] | None = None  # [{base64, filename, download_url}]
    ui_action: dict | None = None  # {action, value} for frontend to execute

class MemoryRequest(BaseModel):
    text: str
    category: str = "general"

class RecallRequest(BaseModel):
    query: str
    category: str = "general"
    n_results: int = 5

class LoginRequest(BaseModel):
    username: str
    password: str

class SetupRequest(BaseModel):
    username: str
    password: str

class AddServerRequest(BaseModel):
    name: str
    host: str
    username: str = "root"
    port: int = 22
    key_path: str = ""
    description: str = ""

class ConversationTitleRequest(BaseModel):
    title: str


class ProjectRequest(BaseModel):
    name: str
    description: str = ""
    color: str = "#3b82f6"


class ProjectUpdateRequest(BaseModel):
    name: str | None = None
    description: str | None = None
    color: str | None = None


class ReferenceNoteRequest(BaseModel):
    title: str
    content: str


class ReferenceUpdateRequest(BaseModel):
    title: str | None = None
    content: str | None = None


# 2FA and Passkey models
class TOTPVerifyRequest(BaseModel):
    code: str


class TOTPLoginRequest(BaseModel):
    username: str
    code: str


class PasskeyCredentialRequest(BaseModel):
    credential: dict
    name: str | None = None


class PasskeyLoginRequest(BaseModel):
    credential: dict
    username: str | None = None


class ConversationProjectRequest(BaseModel):
    project_id: str | None = None


# Conversation history per session
_sessions: dict[str, list[dict]] = {}


async def get_session_messages(session_id: str) -> list[dict]:
    if session_id not in _sessions:
        msgs = [{"role": "system", "content": await get_system_prompt()}]
        # Try to restore from SQLite
        conv = get_conversation(session_id)
        if conv:
            for m in conv["messages"]:
                msgs.append({"role": m["role"], "content": m["content"]})
        _sessions[session_id] = msgs
    return _sessions[session_id]


# ==================== Auth Endpoints ====================

@app.post("/auth/login")
@limiter.limit("5/minute")
async def login(request: Request, req: LoginRequest):
    user = verify_user(req.username, req.password)
    if not user:
        raise HTTPException(status_code=401, detail="Invalid credentials")

    # Check if 2FA is required
    if user.get("totp_enabled"):
        # Return partial auth - client needs to provide TOTP code
        return JSONResponse({
            "requires_2fa": True,
            "username": user["username"],
            "message": "Please enter your 2FA code"
        })

    # No 2FA - complete login
    token = create_access_token({"sub": user["username"], "role": user["role"]})
    response = JSONResponse({"token": token, "username": user["username"], "role": user["role"]})
    response.set_cookie("alfred_token", token, httponly=True, secure=True, samesite="lax", max_age=86400)
    return response


@app.post("/auth/setup")
async def setup(req: SetupRequest):
    """Create the first admin user (only works when no users exist)."""
    from core.security.auth import _load_users
    users = _load_users()
    if users:
        raise HTTPException(status_code=400, detail="Setup already complete")
    create_user(req.username, req.password, "admin")
    token = create_access_token({"sub": req.username, "role": "admin"})
    response = JSONResponse({"token": token, "username": req.username, "message": "Setup complete"})
    response.set_cookie("alfred_token", token, httponly=True, secure=True, samesite="lax", max_age=86400)
    return response


@app.get("/auth/me")
async def me(user: dict = Depends(get_current_user)):
    if user is None:
        return {"authenticated": False}
    return {"authenticated": True, "username": user.get("sub", user.get("username")), "role": user.get("role")}


@app.get("/auth/auto")
async def auto_login(request: Request):
    """Auto-login for trusted local network clients.

    Checks client IP against private subnets. If trusted, issues a JWT cookie
    for the admin user. If not, returns {auto_login: false}.
    """
    import ipaddress
    client_ip = request.client.host if request.client else ""
    # Check X-Forwarded-For for proxied requests
    forwarded = request.headers.get("x-forwarded-for", "")
    if forwarded:
        client_ip = forwarded.split(",")[0].strip()

    trusted = False
    try:
        addr = ipaddress.ip_address(client_ip)
        private_nets = [
            ipaddress.ip_network("10.0.0.0/8"),
            ipaddress.ip_network("172.16.0.0/12"),
            ipaddress.ip_network("192.168.0.0/16"),
            ipaddress.ip_network("127.0.0.0/8"),
            ipaddress.ip_network("::1/128"),
        ]
        trusted = any(addr in net for net in private_nets)
    except ValueError:
        trusted = False

    if not trusted:
        return JSONResponse({"auto_login": False})

    # Find admin user
    from core.security.auth import _load_users
    users = _load_users()
    admin_user = None
    for username, user_data in users.items():
        if user_data.get("role") == "admin":
            admin_user = username
            break

    if not admin_user:
        return JSONResponse({"auto_login": False})

    token = create_access_token({"sub": admin_user, "role": "admin"})
    response = JSONResponse({"auto_login": True, "username": admin_user})
    response.set_cookie("alfred_token", token, httponly=True, secure=False, samesite="lax", max_age=86400)
    return response


@app.post("/auth/logout")
async def logout():
    response = JSONResponse({"message": "Logged out"})
    response.delete_cookie("alfred_token")
    return response


@app.post("/auth/change-password")
async def change_password(req: LoginRequest, user: dict = Depends(require_auth)):
    """Change password. The LoginRequest.password field is the new password."""
    from core.security.auth import _load_users, _save_users, pwd_context
    users = _load_users()
    username = user.get("sub", user.get("username"))
    if username not in users:
        raise HTTPException(status_code=404, detail="User not found")
    users[username]["password_hash"] = pwd_context.hash(req.password)
    _save_users(users)
    return {"message": "Password changed"}


# ==================== 2FA (TOTP) ====================

@app.post("/auth/2fa/setup")
async def totp_setup(user: dict = Depends(require_auth)):
    """Generate TOTP secret and QR code for setup."""
    username = user.get("sub", user.get("username"))
    result = setup_totp(username)
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])
    return result


@app.post("/auth/2fa/verify")
async def totp_verify(req: TOTPVerifyRequest, user: dict = Depends(require_auth)):
    """Verify a TOTP code (for testing during setup)."""
    username = user.get("sub", user.get("username"))
    if verify_totp(username, req.code):
        return {"valid": True}
    raise HTTPException(status_code=400, detail="Invalid code")


@app.post("/auth/2fa/enable")
async def totp_enable(req: TOTPVerifyRequest, user: dict = Depends(require_auth)):
    """Enable 2FA after verifying the TOTP code."""
    username = user.get("sub", user.get("username"))
    result = enable_totp(username, req.code)
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])
    return result


@app.post("/auth/2fa/disable")
async def totp_disable(req: TOTPVerifyRequest, user: dict = Depends(require_auth)):
    """Disable 2FA after verifying the TOTP code."""
    username = user.get("sub", user.get("username"))
    result = disable_totp(username, req.code)
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])
    return result


@app.get("/auth/2fa/status")
async def totp_status(user: dict = Depends(require_auth)):
    """Check if 2FA is enabled for current user."""
    username = user.get("sub", user.get("username"))
    return {"enabled": is_totp_enabled(username)}


@app.post("/auth/2fa/login")
@limiter.limit("5/minute")
async def totp_login(request: Request, req: TOTPLoginRequest):
    """Complete login with TOTP verification (step 2 of 2FA login)."""
    if not verify_totp(req.username, req.code):
        raise HTTPException(status_code=401, detail="Invalid 2FA code")
    # Get user role for token
    from core.security.auth import _load_users
    users = _load_users()
    user = users.get(req.username)
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    # Create full access token
    token = create_access_token({"sub": req.username, "role": user["role"]})
    response = JSONResponse({"token": token, "username": req.username, "role": user["role"]})
    response.set_cookie("alfred_token", token, httponly=True, secure=True, samesite="lax", max_age=86400)
    return response


# ==================== Passkeys (WebAuthn) ====================

@app.post("/auth/passkey/register/begin")
async def passkey_register_begin(user: dict = Depends(require_auth)):
    """Start passkey registration."""
    username = user.get("sub", user.get("username"))
    result = get_passkey_registration_options(username)
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])
    return result


@app.post("/auth/passkey/register/complete")
async def passkey_register_complete(req: PasskeyCredentialRequest, user: dict = Depends(require_auth)):
    """Complete passkey registration."""
    username = user.get("sub", user.get("username"))
    if req.name:
        req.credential["name"] = req.name
    result = verify_passkey_registration(username, req.credential)
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])
    return result


@app.post("/auth/passkey/login/begin")
@limiter.limit("10/minute")
async def passkey_login_begin(request: Request, username: str = None):
    """Get authentication options for passkey login."""
    result = get_passkey_login_options(username)
    return result


@app.post("/auth/passkey/login/complete")
@limiter.limit("5/minute")
async def passkey_login_complete(request: Request, req: PasskeyLoginRequest):
    """Complete passkey login."""
    result = verify_passkey_login(req.credential, req.username)
    if "error" in result:
        raise HTTPException(status_code=401, detail=result["error"])
    # Create access token
    token = create_access_token({"sub": result["username"], "role": result["role"]})
    response = JSONResponse({"token": token, "username": result["username"], "role": result["role"]})
    response.set_cookie("alfred_token", token, httponly=True, secure=True, samesite="lax", max_age=86400)
    return response


@app.get("/auth/passkey/list")
async def passkey_list(user: dict = Depends(require_auth)):
    """List all passkeys for current user."""
    username = user.get("sub", user.get("username"))
    return {"passkeys": list_passkeys(username)}


@app.delete("/auth/passkey/{credential_id}")
async def passkey_delete(credential_id: str, user: dict = Depends(require_auth)):
    """Delete a passkey."""
    username = user.get("sub", user.get("username"))
    result = delete_passkey(username, credential_id)
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])
    return result


@app.get("/auth/methods")
async def auth_methods(username: str = None):
    """Get available auth methods for a user (for login page)."""
    if not username:
        return {"exists": False, "totp_enabled": False, "has_passkeys": False}
    return get_user_auth_methods(username)


# ==================== Google OAuth ====================

@app.get("/auth/google")
async def google_auth(user: dict = Depends(require_auth)):
    """Start Google OAuth flow."""
    if not settings.google_client_id:
        raise HTTPException(status_code=400, detail="Google OAuth not configured")
    auth_url = get_auth_url()
    return RedirectResponse(auth_url)


@app.get("/auth/google/callback")
async def google_callback(code: str = None, error: str = None):
    """Handle Google OAuth callback."""
    if error:
        return HTMLResponse(f"<h2>Google Authorization Failed</h2><p>{error}</p><a href='/'>Back to Alfred</a>")
    if not code:
        return HTMLResponse("<h2>No authorization code received</h2><a href='/'>Back to Alfred</a>")

    success = handle_callback(code)
    if success:
        return HTMLResponse(
            "<h2>Google Connected Successfully</h2>"
            "<p>Gmail and Calendar are now linked to Alfred.</p>"
            "<script>setTimeout(()=>window.location='/',2000)</script>"
        )
    return HTMLResponse("<h2>Connection Failed</h2><p>Check logs for details.</p><a href='/'>Back to Alfred</a>")


# ==================== Integration Status ====================

@app.get("/integrations/status")
async def integration_status(user: dict = Depends(require_auth)):
    """Check which integrations are connected."""
    google = is_connected()

    # Check server config
    from integrations.servers.manager import list_servers
    servers = []
    try:
        servers = list_servers()
    except Exception:
        pass

    # Check CRM connection
    crm_connected = False
    try:
        from integrations.base_crm.client import is_connected as crm_is_connected
        crm_connected = crm_is_connected()
    except Exception:
        pass

    # Check n8n connection
    n8n_connected = False
    n8n_workflow_count = 0
    try:
        from integrations.n8n.client import is_connected as n8n_is_connected, list_workflows as n8n_list
        n8n_connected = n8n_is_connected()
        if n8n_connected:
            n8n_workflow_count = len(n8n_list(limit=100))
    except Exception:
        pass

    # Check Nextcloud connection
    nextcloud_connected = False
    try:
        from integrations.nextcloud.client import is_connected as nc_is_connected
        nextcloud_connected = nc_is_connected()
    except Exception:
        pass

    # Check Stripe connection
    stripe_connected = False
    try:
        from integrations.stripe.client import is_connected as stripe_is_connected
        stripe_connected = stripe_is_connected()
    except Exception:
        pass

    # Check LightRAG connection
    lightrag_connected = False
    lightrag_doc_count = 0
    try:
        from integrations.lightrag.client import is_connected as lightrag_is_connected, get_document_status
        lightrag_connected = await lightrag_is_connected()
        if lightrag_connected:
            status = await get_document_status()
            if status.get("success"):
                lightrag_doc_count = status.get("status", {}).get("status_counts", {}).get("processed", 0)
    except Exception:
        pass

    # Check Twilio connection
    twilio_connected = False
    twilio_phone = ""
    try:
        from integrations.twilio.client import health_check as twilio_health
        twilio_status = twilio_health()
        twilio_connected = twilio_status.get("success", False)
        twilio_phone = twilio_status.get("phone_number", "")
    except Exception:
        pass

    return {
        "google": {
            "connected": google,
            "services": ["Gmail", "Calendar"] if google else [],
        },
        "servers": {
            "count": len(servers),
            "names": [s["name"] for s in servers],
        },
        "base_crm": {
            "configured": bool(settings.base_crm_api_key),
            "connected": crm_connected,
            "label": "Twenty CRM",
        },
        "n8n": {
            "configured": bool(settings.n8n_api_key),
            "connected": n8n_connected,
            "workflow_count": n8n_workflow_count,
            "label": "n8n Automations",
        },
        "nextcloud": {
            "configured": bool(settings.nextcloud_url),
            "connected": nextcloud_connected,
            "label": "Nextcloud",
        },
        "stripe": {
            "configured": bool(settings.stripe_api_key),
            "connected": stripe_connected,
            "label": "Stripe Payments",
        },
        "lightrag": {
            "configured": True,  # Always configured via env
            "connected": lightrag_connected,
            "document_count": lightrag_doc_count,
            "label": "LightRAG Knowledge Graph",
        },
        "twilio": {
            "configured": bool(os.getenv("TWILIO_ACCOUNT_SID")),
            "connected": twilio_connected,
            "phone_number": twilio_phone,
            "label": "Twilio SMS/Voice",
        },
        "agents": {
            "active": True,
            "max_concurrent": 3,
            "label": "Agent Orchestration",
        },
        "learning": {
            "active": True,
            "label": "Adaptive Learning",
        },
        "ollama": {
            "model": settings.ollama_model,
            "host": settings.ollama_host,
        },
        "anthropic": {
            "configured": bool(settings.anthropic_api_key and settings.anthropic_api_key != "sk-ant-CHANGEME"),
            "model": settings.anthropic_model,
        },
        "claude_code": {
            "configured": _check_claude_code_cli(),
            "label": "Claude Code (Max)",
        },
        "tts": {
            "backend": settings.tts_model,
            "qwen3_available": _check_qwen3_available(),
        },
    }


def _check_qwen3_available() -> bool:
    """Check if Qwen3-TTS server is available."""
    try:
        import requests
        resp = requests.get("http://localhost:7860/docs", timeout=2)
        return resp.status_code == 200
    except Exception:
        return False


def _check_claude_code_cli() -> bool:
    """Check if Claude Code CLI is available (for Max subscription users)."""
    import shutil
    import os
    from pathlib import Path
    # Check standard which first
    if shutil.which("claude"):
        return True
    # Also check common nvm/npm paths since uvicorn may not inherit full PATH
    common_paths = [
        Path.home() / ".nvm/versions/node/v20.12.2/bin/claude",
        Path.home() / ".local/bin/claude",
        Path("/usr/local/bin/claude"),
    ]
    for path in common_paths:
        if path.exists() and os.access(path, os.X_OK):
            return True
    return False


# ==================== Twilio Webhooks ====================

@app.post("/webhooks/twilio/sms")
async def twilio_sms_webhook(request: Request):
    """Handle incoming SMS from Twilio.

    This endpoint is called by Twilio when someone sends an SMS to your Twilio number.
    No auth required - we validate the Twilio signature instead.
    """
    from integrations.twilio.client import validate_twilio_signature, send_sms, TWILIO_PHONE_NUMBER
    from twilio.twiml.messaging_response import MessagingResponse

    # Get the request data
    form_data = await request.form()
    params = dict(form_data)

    # Validate Twilio signature
    signature = request.headers.get("X-Twilio-Signature", "")
    url = str(request.url)

    # For local testing, skip validation if no signature
    if signature and not validate_twilio_signature(url, params, signature):
        logger.warning("Invalid Twilio signature on SMS webhook")
        return Response(content="Invalid signature", status_code=403)

    # Extract message details
    from_number = params.get("From", "")
    to_number = params.get("To", "")
    body = params.get("Body", "").strip()

    logger.info(f"SMS received from {from_number}: {body[:50]}...")

    # Process the message through Alfred
    try:
        response = await ask(
            query=body,
            messages=None,
            smart_routing=True,
        )
        reply_text = response.get("response", "Sorry, I couldn't process that request.")
    except Exception as e:
        logger.error(f"Error processing SMS: {e}")
        reply_text = "Sorry, I encountered an error. Please try again."

    # Send TwiML response
    twiml = MessagingResponse()
    # Truncate if too long (SMS limit is 1600 chars for concatenated)
    if len(reply_text) > 1500:
        reply_text = reply_text[:1497] + "..."
    twiml.message(reply_text)

    return Response(content=str(twiml), media_type="application/xml")


@app.post("/webhooks/twilio/voice")
async def twilio_voice_webhook(request: Request):
    """Handle incoming voice calls from Twilio.

    This endpoint is called when someone calls your Twilio number.
    Alfred will answer and have a conversation using Kokoro TTS (bm_daniel voice).
    """
    from twilio.twiml.voice_response import VoiceResponse, Gather

    # Get the request data
    form_data = await request.form()
    params = dict(form_data)

    from_number = params.get("From", "")
    speech_result = params.get("SpeechResult", "")
    call_status = params.get("CallStatus", "")

    logger.info(f"Voice webhook: from={from_number}, status={call_status}, speech={speech_result[:50] if speech_result else 'none'}")

    response = VoiceResponse()

    # Build absolute URL for action and audio
    base_url = "https://aialfred.groundrushcloud.com"

    if not speech_result:
        # Initial greeting - use time-based greeting (pre-generated, instant)
        greeting_audio_id = get_greeting_audio_id()
        goodbye_audio_id = get_static_audio_id("goodbye")

        gather = Gather(
            input="speech",
            action=f"{base_url}/webhooks/twilio/voice",
            method="POST",
            speech_timeout="3",
            timeout=5,
            language="en-US",
        )
        gather.play(f"{base_url}/audio/tts/{greeting_audio_id}.wav")
        response.append(gather)
        # If no input, say goodbye
        response.play(f"{base_url}/audio/tts/{goodbye_audio_id}.wav")
    else:
        # Play "One moment" immediately, then redirect to process
        # This gives instant feedback while Alfred thinks
        from urllib.parse import quote
        thinking_audio_id = get_random_thinking_audio_id()
        response.play(f"{base_url}/audio/tts/{thinking_audio_id}.wav")
        encoded_speech = quote(speech_result, safe='')
        response.redirect(f"{base_url}/webhooks/twilio/voice/process?speech={encoded_speech}", method="POST")

    return Response(content=str(response), media_type="application/xml")


@app.post("/webhooks/twilio/voice/process")
async def twilio_voice_process(request: Request):
    """Process speech and return Alfred's response.

    Called after playing "One moment" to reduce perceived latency.
    """
    from twilio.twiml.voice_response import VoiceResponse, Gather
    from urllib.parse import unquote

    # Get speech from query params
    speech_result = request.query_params.get("speech", "")
    if speech_result:
        speech_result = unquote(speech_result)

    base_url = "https://aialfred.groundrushcloud.com"
    response = VoiceResponse()

    if not speech_result:
        # No speech to process
        response.play(f"{base_url}/audio/tts/{get_static_audio_id('error')}.wav")
        return Response(content=str(response), media_type="application/xml")

    # Process speech through Alfred
    # Add phone-specific instructions to keep responses brief (Twilio has 15s timeout)
    phone_query = f"[PHONE CALL - Keep response under 3 sentences. Be concise and conversational. No lists or detailed breakdowns.] {speech_result}"
    logger.info(f"Processing speech: {speech_result}")
    try:
        import asyncio
        # Timeout after 12 seconds to stay under Twilio's 15s limit
        alfred_response = await asyncio.wait_for(
            ask(
                query=phone_query,
                messages=None,
                smart_routing=True,
            ),
            timeout=12.0
        )
        reply_text = alfred_response.get("response", "I'm sorry, I couldn't process that.")
    except asyncio.TimeoutError:
        logger.warning("Phone call processing timed out after 12s")
        reply_text = "I'm still working on that. Can you ask me again in a simpler way?"
    except Exception as e:
        logger.error(f"Error processing voice: {e}")
        reply_text = None

    # Generate audio for response (dynamic) or use pre-generated error
    if reply_text:
        reply_audio_id = generate_tts_audio(reply_text)
    else:
        reply_audio_id = get_static_audio_id("error")

    # Follow-up is always the same - use pre-generated (instant)
    followup_audio_id = get_static_audio_id("followup")

    # Speak the response and listen for more
    gather = Gather(
        input="speech",
        action=f"{base_url}/webhooks/twilio/voice",
        method="POST",
        speech_timeout="3",
        timeout=5,
        language="en-US",
    )
    gather.play(f"{base_url}/audio/tts/{reply_audio_id}.wav")
    response.append(gather)
    response.play(f"{base_url}/audio/tts/{followup_audio_id}.wav")

    return Response(content=str(response), media_type="application/xml")


@app.post("/webhooks/twilio/voice/status")
async def twilio_voice_status_webhook(request: Request):
    """Handle voice call status updates from Twilio."""
    form_data = await request.form()
    params = dict(form_data)

    call_sid = params.get("CallSid", "")
    call_status = params.get("CallStatus", "")

    logger.info(f"Call {call_sid} status: {call_status}")

    return Response(content="OK", status_code=200)


@app.post("/webhooks/twilio/voice/outbound")
@app.get("/webhooks/twilio/voice/outbound")
async def twilio_outbound_twiml(request: Request):
    """Generate TwiML for outbound calls using Kokoro TTS.

    Called by Twilio when making outbound calls via make_call().
    Accepts 'message' in query params or form data.
    """
    from twilio.twiml.voice_response import VoiceResponse
    from urllib.parse import unquote

    # Try to get message from query params first, then form data
    message = request.query_params.get("message", "")
    if not message:
        try:
            form_data = await request.form()
            message = form_data.get("message", "")
        except Exception:
            pass

    # Build absolute URL for audio
    base_url = "https://aialfred.groundrushcloud.com"

    if not message:
        # Use pre-generated default (instant)
        audio_id = get_static_audio_id("default_outbound")
        message = STATIC_PHRASES["default_outbound"]
    else:
        # URL decode and generate audio (uses cache for repeated phrases)
        message = unquote(message)
        audio_id = generate_tts_audio(message)

    response = VoiceResponse()
    response.play(f"{base_url}/audio/tts/{audio_id}.wav")

    logger.info(f"Generated outbound TwiML with Kokoro TTS for message: {message[:50]}...")

    return Response(content=str(response), media_type="application/xml")


# TTS audio cache: hash -> audio_id (for repeated phrases)
_tts_cache: dict[str, str] = {}

# Pre-generated static phrases (populated on startup)
STATIC_PHRASES = {
    "greeting_morning": "Good morning, sir.  How can I be of service?",
    "greeting_afternoon": "Good afternoon, sir.  How can I be of service?",
    "greeting_evening": "Good evening, sir.  How can I be of service?",
    "goodbye": "I didn't hear anything. Goodbye.",
    "followup": "Is there anything else I can help with?",
    "error": "I encountered an error. Please try again.",
    "default_outbound": "Hello, this is Alfred calling.",
}

# 50 natural acknowledgment variations for thinking time
THINKING_PHRASES = [
    "Let me check on that for you.",
    "One moment please.",
    "Let me look into that.",
    "Give me just a second.",
    "Let me see what I can find.",
    "Hold on, let me check.",
    "Just a moment.",
    "Let me pull that up.",
    "One sec.",
    "Let me take a look.",
    "Alright, checking now.",
    "Let me find that for you.",
    "Sure, one moment.",
    "Let me get that information.",
    "Okay, looking into it.",
    "Let me see here.",
    "Just a second.",
    "Let me dig into that.",
    "Hang on a moment.",
    "Let me check my sources.",
    "Working on that now.",
    "Let me look that up.",
    "Sure thing, one moment.",
    "Let me find out.",
    "Checking on that.",
    "Let me see what I've got.",
    "One moment, please.",
    "Let me get back to you on that.",
    "Alright, let me check.",
    "Give me a moment.",
    "Let me look into this.",
    "Sure, let me check.",
    "Just one second.",
    "Let me pull up that information.",
    "Okay, one moment.",
    "Let me search for that.",
    "Hold on just a second.",
    "Let me take a quick look.",
    "Alright, one sec.",
    "Let me review that.",
    "Sure, checking now.",
    "Let me gather that info.",
    "One second please.",
    "Let me verify that.",
    "Okay, let me see.",
    "Let me check real quick.",
    "Alright, looking now.",
    "Let me fetch that.",
    "Sure, give me a moment.",
    "Let me process that.",
]

# Pre-generated thinking phrase audio IDs (populated on startup)
_thinking_audio_ids: list[str] = []
_static_audio_ids: dict[str, str] = {}


def _hash_text(text: str) -> str:
    """Create a short hash of text for caching."""
    import hashlib
    return hashlib.md5(text.encode()).hexdigest()[:16]


def generate_tts_audio(text: str, use_cache: bool = True) -> str:
    """Generate TTS audio using Kokoro and return the audio_id.

    Uses caching to avoid regenerating the same phrases.
    Returns the audio_id (without extension) that can be used to build the URL.
    """
    from interfaces.voice.tts import speak

    # Check cache first
    if use_cache:
        text_hash = _hash_text(text)
        if text_hash in _tts_cache:
            cached_id = _tts_cache[text_hash]
            # Verify file still exists
            if (TTS_AUDIO_DIR / f"{cached_id}.wav").exists():
                logger.debug(f"TTS cache hit for: {text[:30]}...")
                return cached_id

    # Generate unique ID for this audio
    audio_id = str(uuid.uuid4())
    audio_path = TTS_AUDIO_DIR / f"{audio_id}.wav"

    # Generate audio using Kokoro TTS
    audio_data = speak(text)

    # Save to file
    audio_path.write_bytes(audio_data)
    logger.info(f"Generated TTS audio: {audio_id} ({len(audio_data)} bytes)")

    # Cache it
    if use_cache:
        _tts_cache[text_hash] = audio_id

    return audio_id


def get_static_audio_id(phrase_key: str) -> str:
    """Get pre-generated audio ID for static phrases (fast path)."""
    if phrase_key in _static_audio_ids:
        return _static_audio_ids[phrase_key]
    # Fallback to generating if not pre-cached
    return generate_tts_audio(STATIC_PHRASES.get(phrase_key, phrase_key))


def get_greeting_audio_id() -> str:
    """Get the appropriate greeting based on time of day."""
    from datetime import datetime
    hour = datetime.now().hour

    if 5 <= hour < 12:
        return get_static_audio_id("greeting_morning")
    elif 12 <= hour < 17:
        return get_static_audio_id("greeting_afternoon")
    else:
        return get_static_audio_id("greeting_evening")


def pregenerate_static_phrases():
    """Pre-generate audio for static phrases on startup."""
    import random

    logger.info("Pre-generating static TTS phrases...")
    for key, text in STATIC_PHRASES.items():
        audio_id = generate_tts_audio(text, use_cache=True)
        _static_audio_ids[key] = audio_id
        logger.info(f"Pre-generated '{key}': {audio_id}")

    # Pre-generate all 50 thinking variations
    logger.info(f"Pre-generating {len(THINKING_PHRASES)} thinking phrase variations...")
    for i, text in enumerate(THINKING_PHRASES):
        audio_id = generate_tts_audio(text, use_cache=True)
        _thinking_audio_ids.append(audio_id)
        if (i + 1) % 10 == 0:
            logger.info(f"Pre-generated {i + 1}/{len(THINKING_PHRASES)} thinking phrases...")

    logger.info(f"Pre-generated {len(STATIC_PHRASES)} static + {len(THINKING_PHRASES)} thinking phrases")


def get_random_thinking_audio_id() -> str:
    """Get a random pre-generated thinking phrase audio ID."""
    import random
    if _thinking_audio_ids:
        return random.choice(_thinking_audio_ids)
    # Fallback if not pre-generated
    return generate_tts_audio(random.choice(THINKING_PHRASES))


@app.get("/audio/tts/{audio_id}.wav")
async def serve_tts_audio(audio_id: str):
    """Serve generated TTS audio for Twilio voice calls.

    This endpoint serves WAV audio files generated by Kokoro TTS.
    Audio files are temporary and can be cleaned up periodically.
    """
    # Validate audio_id format (UUID)
    try:
        uuid.UUID(audio_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid audio ID")

    audio_path = TTS_AUDIO_DIR / f"{audio_id}.wav"

    if not audio_path.exists():
        raise HTTPException(status_code=404, detail="Audio not found")

    audio_data = audio_path.read_bytes()

    return Response(
        content=audio_data,
        media_type="audio/wav",
        headers={
            "Content-Disposition": f"inline; filename={audio_id}.wav",
            "Cache-Control": "no-cache",
        }
    )


@app.get("/settings/tts")
async def get_tts_settings(user: dict = Depends(require_auth)):
    """Get current TTS settings."""
    return {
        "backend": settings.tts_model,
        "qwen3_available": _check_qwen3_available(),
    }


def _safe_env_update(key: str, value: str) -> None:
    """Safely update a key in .env file with atomic write and backup."""
    env_path = Path("/home/aialfred/alfred/config/.env")
    backup_path = Path("/home/aialfred/alfred/config/.env.backup")

    if env_path.exists():
        original = env_path.read_text()
        original_lines = len(original.strip().split('\n'))

        # Update or append the key
        pattern = rf"^{re.escape(key)}=.*$"
        if re.search(pattern, original, re.MULTILINE):
            new_content = re.sub(pattern, f"{key}={value}", original, flags=re.MULTILINE)
        else:
            new_content = original.rstrip() + f"\n{key}={value}\n"

        # Safety check: new content shouldn't be drastically shorter
        new_lines = len(new_content.strip().split('\n'))
        if new_lines < original_lines - 1:
            logger.error(f"ENV update aborted: would reduce from {original_lines} to {new_lines} lines")
            raise HTTPException(status_code=500, detail="ENV update failed safety check")

        # Atomic write: write to temp, then rename
        import tempfile
        with tempfile.NamedTemporaryFile(mode='w', dir=env_path.parent, delete=False) as tmp:
            tmp.write(new_content)
            tmp_path = Path(tmp.name)

        # Update backup
        backup_path.write_text(original)

        # Atomic rename
        tmp_path.rename(env_path)
    else:
        env_path.write_text(f"{key}={value}\n")


@app.put("/settings/tts")
async def set_tts_settings(backend: str, user: dict = Depends(require_auth)):
    """Set TTS backend (kokoro, qwen3, piper)."""
    if backend not in ["kokoro", "qwen3", "piper"]:
        raise HTTPException(status_code=400, detail="Invalid TTS backend")

    _safe_env_update("TTS_MODEL", backend)

    # Update runtime settings
    settings.tts_model = backend

    return {"backend": backend, "message": f"TTS backend set to {backend}"}


# ==================== Knowledge Management (LightRAG) ====================

class KnowledgeTextRequest(BaseModel):
    text: str
    description: str = ""


class KnowledgeQueryRequest(BaseModel):
    query: str
    mode: str = "hybrid"  # naive, local, global, hybrid
    top_k: int = 10


@app.get("/knowledge/status")
async def knowledge_status(user: dict = Depends(require_auth)):
    """Get LightRAG knowledge base status."""
    from integrations.lightrag.client import health_check, get_document_status, is_connected

    health = await health_check()
    if not health.get("healthy"):
        return {"connected": False, "error": health.get("error")}

    status = await get_document_status()
    return {
        "connected": True,
        "document_counts": status.get("status", {}).get("status_counts", {}),
        "details": health.get("details", {}).get("configuration", {}),
    }


@app.get("/knowledge/documents")
async def list_knowledge_documents(
    limit: int = 20,
    offset: int = 0,
    user: dict = Depends(require_auth)
):
    """List documents in the knowledge base."""
    from integrations.lightrag.client import list_documents

    result = await list_documents(limit=limit, offset=offset)
    if result.get("success"):
        return result.get("documents", {})
    raise HTTPException(status_code=500, detail=result.get("error", "Failed to list documents"))


@app.post("/knowledge/upload/text")
async def upload_knowledge_text(req: KnowledgeTextRequest, user: dict = Depends(require_auth)):
    """Upload text content to the knowledge base."""
    from integrations.lightrag.client import upload_text

    if not req.text or len(req.text) < 10:
        raise HTTPException(status_code=400, detail="Text must be at least 10 characters")

    result = await upload_text(req.text, req.description)
    if result.get("success"):
        return {"message": "Text uploaded successfully", "result": result.get("result")}
    raise HTTPException(status_code=500, detail=result.get("error", "Upload failed"))


@app.post("/knowledge/upload/file")
async def upload_knowledge_file(
    file: UploadFile = File(...),
    user: dict = Depends(require_auth)
):
    """Upload a file to the knowledge base (PDF, TXT, MD, DOCX)."""
    from integrations.lightrag.client import upload_file
    import tempfile
    import os

    # Validate file type
    allowed_types = [".pdf", ".txt", ".md", ".docx", ".doc"]
    ext = Path(file.filename).suffix.lower()
    if ext not in allowed_types:
        raise HTTPException(status_code=400, detail=f"File type {ext} not supported. Allowed: {allowed_types}")

    # Save to temp file and upload
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name

    try:
        result = await upload_file(tmp_path)
        if result.get("success"):
            return {"message": f"File '{file.filename}' uploaded successfully", "result": result.get("result")}
        raise HTTPException(status_code=500, detail=result.get("error", "Upload failed"))
    finally:
        os.unlink(tmp_path)


@app.post("/knowledge/query")
async def query_knowledge(req: KnowledgeQueryRequest, user: dict = Depends(require_auth)):
    """Query the knowledge base."""
    from integrations.lightrag.client import query

    result = await query(req.query, mode=req.mode, top_k=req.top_k)
    if result.get("success"):
        return result.get("result")
    raise HTTPException(status_code=500, detail=result.get("error", "Query failed"))


@app.get("/knowledge/context")
async def get_knowledge_context_api(query: str, top_k: int = 5, user: dict = Depends(require_auth)):
    """Get relevant context for a query (without LLM response)."""
    from integrations.lightrag.client import query_context

    result = await query_context(query, top_k=top_k)
    if result.get("success"):
        return result.get("result")
    raise HTTPException(status_code=500, detail=result.get("error", "Context retrieval failed"))


@app.get("/knowledge/entities")
async def get_knowledge_entities(limit: int = 20, user: dict = Depends(require_auth)):
    """Get popular entities from the knowledge graph."""
    from integrations.lightrag.client import get_popular_entities

    result = await get_popular_entities(limit=limit)
    if result.get("success"):
        return result.get("entities")
    raise HTTPException(status_code=500, detail=result.get("error", "Failed to get entities"))


@app.get("/knowledge/search")
async def search_knowledge_graph(label: str, user: dict = Depends(require_auth)):
    """Search the knowledge graph for entities matching a label."""
    from integrations.lightrag.client import search_graph

    result = await search_graph(label)
    if result.get("success"):
        return result.get("entities")
    raise HTTPException(status_code=500, detail=result.get("error", "Search failed"))


@app.delete("/knowledge/document/{doc_id}")
async def delete_knowledge_document(doc_id: str, user: dict = Depends(require_auth)):
    """Delete a document from the knowledge base."""
    from integrations.lightrag.client import delete_document

    result = await delete_document(doc_id)
    if result.get("success"):
        return {"message": "Document deleted", "result": result.get("result")}
    raise HTTPException(status_code=500, detail=result.get("error", "Delete failed"))


# ==================== Agent Orchestration ====================

class SpawnAgentRequest(BaseModel):
    goal: str
    agent_type: str = "general"  # coder, researcher, analyst, writer, planner, executor, general
    context: str = ""
    model_override: str | None = None


@app.post("/agents/spawn")
async def spawn_agent(request: SpawnAgentRequest, user: dict = Depends(require_auth)):
    """Spawn a specialized agent to work on a task."""
    from core.orchestration.agents import get_agent_pool, AgentType

    pool = get_agent_pool()

    # Map string to AgentType enum
    agent_type_map = {
        "coder": AgentType.CODER,
        "researcher": AgentType.RESEARCHER,
        "analyst": AgentType.ANALYST,
        "writer": AgentType.WRITER,
        "planner": AgentType.PLANNER,
        "executor": AgentType.EXECUTOR,
        "general": AgentType.GENERAL,
    }
    agent_type = agent_type_map.get(request.agent_type.lower(), AgentType.GENERAL)

    task_id = await pool.spawn_agent(
        goal=request.goal,
        agent_type=agent_type,
        context=request.context,
        model_override=request.model_override,
    )

    return {"task_id": task_id, "agent_type": agent_type.value, "status": "pending"}


@app.get("/agents/tasks")
async def list_agent_tasks(status: str | None = None, user: dict = Depends(require_auth)):
    """List all agent tasks, optionally filtered by status."""
    from core.orchestration.agents import get_agent_pool, AgentStatus

    pool = get_agent_pool()

    status_filter = None
    if status:
        status_map = {
            "pending": AgentStatus.PENDING,
            "running": AgentStatus.RUNNING,
            "completed": AgentStatus.COMPLETED,
            "failed": AgentStatus.FAILED,
            "cancelled": AgentStatus.CANCELLED,
        }
        status_filter = status_map.get(status.lower())

    tasks = pool.list_tasks(status=status_filter)
    return {"tasks": tasks, "count": len(tasks)}


@app.get("/agents/tasks/{task_id}")
async def get_agent_task(task_id: str, wait: bool = False, user: dict = Depends(require_auth)):
    """Get the status and result of an agent task."""
    from core.orchestration.agents import get_agent_pool

    pool = get_agent_pool()
    result = await pool.get_task_result(task_id, wait=wait, timeout=30)

    if not result:
        raise HTTPException(status_code=404, detail=f"Task {task_id} not found")

    return result


@app.post("/agents/tasks/{task_id}/cancel")
async def cancel_agent_task(task_id: str, user: dict = Depends(require_auth)):
    """Cancel a pending or running agent task."""
    from core.orchestration.agents import get_agent_pool

    pool = get_agent_pool()
    success = await pool.cancel_task(task_id)

    if success:
        return {"message": f"Task {task_id} cancelled"}
    raise HTTPException(status_code=400, detail=f"Could not cancel task {task_id}")


@app.get("/agents/pool")
async def get_agent_pool_status(user: dict = Depends(require_auth)):
    """Get the agent pool status and active agents."""
    from core.orchestration.agents import get_agent_pool

    pool = get_agent_pool()
    agents = pool.list_agents()
    tasks = pool.list_tasks()

    pending = len([t for t in tasks if t["status"] == "pending"])
    running = len([t for t in tasks if t["status"] == "running"])
    completed = len([t for t in tasks if t["status"] == "completed"])
    failed = len([t for t in tasks if t["status"] == "failed"])

    return {
        "max_concurrent": pool.max_concurrent,
        "active_agents": len(agents),
        "agents": agents,
        "task_counts": {
            "pending": pending,
            "running": running,
            "completed": completed,
            "failed": failed,
            "total": len(tasks),
        },
    }


# ==================== Daily Briefing ====================

class BriefingRequest(BaseModel):
    include_calendar: bool = True
    include_email: bool = True
    include_crm: bool = True
    include_servers: bool = True
    include_revenue: bool = True
    include_weather: bool = True
    weather_location: str = "Atlanta, GA"


@app.get("/briefing")
async def get_daily_briefing(
    include_weather: bool = True,
    weather_location: str = "Atlanta, GA",
    user: dict = Depends(require_auth),
):
    """Get a personalized daily briefing."""
    from core.briefing.daily import generate_briefing

    briefing = await generate_briefing(
        include_weather=include_weather,
        weather_location=weather_location,
    )
    return briefing.to_dict()


@app.post("/briefing")
async def generate_custom_briefing(request: BriefingRequest, user: dict = Depends(require_auth)):
    """Generate a custom briefing with selected sections."""
    from core.briefing.daily import generate_briefing

    briefing = await generate_briefing(
        include_calendar=request.include_calendar,
        include_email=request.include_email,
        include_crm=request.include_crm,
        include_servers=request.include_servers,
        include_revenue=request.include_revenue,
        include_weather=request.include_weather,
        weather_location=request.weather_location,
    )
    return briefing.to_dict()


@app.get("/briefing/quick")
async def get_quick_briefing(user: dict = Depends(require_auth)):
    """Get a quick text briefing suitable for voice output."""
    from core.briefing.daily import generate_quick_briefing

    text = await generate_quick_briefing()
    return {"text": text}


# ==================== Learning System ====================

class FeedbackRequest(BaseModel):
    feedback_type: str  # 'positive', 'negative', 'correction'
    original_response: str = ""
    correction: str = ""
    context: str = ""
    category: str = ""
    conversation_id: str = ""
    message_id: str = ""


class PreferenceRequest(BaseModel):
    category: str
    key: str
    value: str


@app.post("/learning/feedback")
async def submit_feedback(request: FeedbackRequest, user: dict = Depends(require_auth)):
    """Submit feedback on Alfred's response."""
    from core.learning.feedback import record_feedback

    feedback_id = record_feedback(
        feedback_type=request.feedback_type,
        original_response=request.original_response,
        correction=request.correction,
        context=request.context,
        category=request.category,
        conversation_id=request.conversation_id,
        message_id=request.message_id,
    )
    return {"feedback_id": feedback_id, "message": "Feedback recorded. Thank you!"}


@app.get("/learning/feedback/stats")
async def get_feedback_statistics(user: dict = Depends(require_auth)):
    """Get feedback statistics."""
    from core.learning.feedback import get_feedback_stats
    return get_feedback_stats()


@app.get("/learning/preferences")
async def get_all_preferences(category: str = None, user: dict = Depends(require_auth)):
    """Get user preferences."""
    from core.learning.preferences import get_preferences
    return get_preferences(category)


@app.put("/learning/preferences")
async def set_preference(request: PreferenceRequest, user: dict = Depends(require_auth)):
    """Set a user preference."""
    from core.learning.preferences import update_preference

    success = update_preference(
        category=request.category,
        key=request.key,
        value=request.value,
        source="explicit",
    )
    return {"success": success, "message": f"Preference {request.category}.{request.key} updated"}


@app.get("/learning/patterns")
async def get_detected_patterns(user: dict = Depends(require_auth)):
    """Get detected usage patterns."""
    from core.learning.patterns import detect_patterns, get_pattern_stats

    return {
        "patterns": detect_patterns(),
        "stats": get_pattern_stats(),
    }


@app.get("/learning/suggestions")
async def get_workflow_suggestions_endpoint(user: dict = Depends(require_auth)):
    """Get pending workflow automation suggestions."""
    from core.learning.patterns import get_workflow_suggestions
    return {"suggestions": get_workflow_suggestions()}


@app.post("/learning/suggestions/{suggestion_id}")
async def respond_to_workflow_suggestion(
    suggestion_id: int,
    accept: bool,
    user: dict = Depends(require_auth),
):
    """Accept or dismiss a workflow suggestion."""
    from core.learning.patterns import respond_to_suggestion

    success = respond_to_suggestion(suggestion_id, accept)
    action = "accepted" if accept else "dismissed"
    return {"success": success, "message": f"Suggestion {action}"}


@app.get("/learning/stats")
async def get_learning_stats(user: dict = Depends(require_auth)):
    """Get overall learning system statistics."""
    from core.learning.feedback import get_feedback_stats
    from core.learning.patterns import get_pattern_stats

    return {
        "feedback": get_feedback_stats(),
        "patterns": get_pattern_stats(),
    }


# Kokoro voice options
KOKORO_VOICES = [
    {"id": "bm_daniel", "name": "Daniel", "desc": "British Male"},
    {"id": "bm_george", "name": "George", "desc": "British Male"},
    {"id": "bm_lewis", "name": "Lewis", "desc": "British Male"},
    {"id": "bf_emma", "name": "Emma", "desc": "British Female"},
    {"id": "bf_isabella", "name": "Isabella", "desc": "British Female"},
    {"id": "am_adam", "name": "Adam", "desc": "American Male"},
    {"id": "am_michael", "name": "Michael", "desc": "American Male"},
    {"id": "af_heart", "name": "Heart", "desc": "American Female"},
    {"id": "af_bella", "name": "Bella", "desc": "American Female"},
    {"id": "af_nicole", "name": "Nicole", "desc": "American Female"},
    {"id": "af_sarah", "name": "Sarah", "desc": "American Female"},
    {"id": "af_sky", "name": "Sky", "desc": "American Female"},
]


@app.get("/settings/voices")
async def get_voices(user: dict = Depends(require_auth)):
    """Get available voices for the current TTS backend."""
    import requests as req
    # Read directly from env file to avoid caching issues
    env_path = Path("/home/aialfred/alfred/config/.env")
    backend = settings.tts_model
    if env_path.exists():
        for line in env_path.read_text().splitlines():
            if line.startswith("TTS_MODEL="):
                backend = line.split("=", 1)[1].strip()
                break

    if backend == "qwen3":
        voices = []

        # List cloned voices from Qwen3 resources directory
        resources_dir = Path("/home/aialfred/qwen3-tts/resources")
        if resources_dir.exists():
            for f in resources_dir.iterdir():
                if f.suffix in [".mp3", ".wav"]:
                    voice_id = f.stem
                    voices.append({"id": voice_id, "name": voice_id, "desc": "Cloned voice", "type": "clone"})

        # List designed voices from Qwen3 API
        try:
            resp = req.get("http://localhost:7860/voice_design/voices", timeout=5)
            if resp.status_code == 200:
                for v in resp.json().get("voices", []):
                    voices.append({
                        "id": f"design:{v['name']}",
                        "name": f"{v['name']} (designed)",
                        "desc": v.get("description", "Designed voice")[:50],
                        "type": "design"
                    })
        except Exception as e:
            logger.warning(f"Could not fetch designed voices: {e}")

        if not voices:
            voices = [{"id": "demo_speaker0", "name": "Demo Speaker", "desc": "Default voice", "type": "clone"}]
        # Read current voice from env file
        current_voice = settings.tts_voice
        for line in env_path.read_text().splitlines():
            if line.startswith("TTS_VOICE="):
                current_voice = line.split("=", 1)[1].strip()
                break
        return {"backend": backend, "voices": voices, "current": current_voice}
    else:
        # Kokoro voices - read current voice from env file
        current_voice = settings.tts_voice
        if env_path.exists():
            for line in env_path.read_text().splitlines():
                if line.startswith("TTS_VOICE="):
                    current_voice = line.split("=", 1)[1].strip()
                    break
        return {"backend": backend, "voices": KOKORO_VOICES, "current": current_voice}


@app.put("/settings/voice")
async def set_voice(voice: str, user: dict = Depends(require_auth)):
    """Set the TTS voice."""
    _safe_env_update("TTS_VOICE", voice)

    # Update runtime settings
    settings.tts_voice = voice

    return {"voice": voice, "message": f"Voice set to {voice}"}


# ==================== Server Management ====================

@app.post("/servers/add")
async def add_server(req: AddServerRequest, user: dict = Depends(require_auth)):
    """Register a server for management."""
    from integrations.servers.manager import add_server as _add
    _add(req.name, req.host, req.username, req.port, req.key_path, req.description)
    return {"message": f"Server '{req.name}' registered", "name": req.name}


# ==================== File Upload/Download ====================

@app.post("/upload/document")
async def upload_document(file: UploadFile = File(...), user: dict = Depends(require_auth)):
    """Upload a document for analysis (PDF, DOCX, XLSX, CSV, TXT, MD, JSON)."""
    from core.tools.files import save_upload, parse_document
    allowed_ext = {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".csv", ".txt", ".md", ".json"}
    ext = Path(file.filename).suffix.lower() if file.filename else ""
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")
    content = await file.read()
    if len(content) > 10 * 1024 * 1024:  # 10MB limit
        raise HTTPException(status_code=400, detail="File too large (max 10MB)")
    path = save_upload(content, file.filename)
    # Parse the document
    result = parse_document(path)

    # Index to LightRAG for long-term memory
    lightrag_status = "not_indexed"
    if result["text"] and not result["error"]:
        try:
            from integrations.lightrag.client import upload_text
            lr_result = await upload_text(
                result["text"],
                description=f"Document: {file.filename}"
            )
            lightrag_status = "indexed" if lr_result["success"] else f"failed: {lr_result.get('error', 'unknown')}"
            logger.info(f"LightRAG indexing: {lightrag_status}")
        except Exception as e:
            lightrag_status = f"failed: {e}"
            logger.warning(f"LightRAG indexing failed: {e}")

    return {
        "path": path,
        "filename": file.filename,
        "text_preview": result["text"][:2000] + ("..." if len(result["text"]) > 2000 else ""),
        "metadata": result["metadata"],
        "error": result["error"],
        "lightrag_status": lightrag_status,
    }


@app.post("/upload/image")
async def upload_image(file: UploadFile = File(...), user: dict = Depends(require_auth)):
    """Upload an image for vision analysis."""
    from core.tools.files import save_upload
    import base64
    allowed_ext = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
    ext = Path(file.filename).suffix.lower() if file.filename else ""
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"Unsupported image type: {ext}")
    content = await file.read()
    if len(content) > 20 * 1024 * 1024:  # 20MB limit
        raise HTTPException(status_code=400, detail="Image too large (max 20MB)")
    path = save_upload(content, file.filename)
    # Return base64 for inline display
    b64 = base64.b64encode(content).decode("utf-8")
    media_type = f"image/{ext.lstrip('.')}"
    if ext == ".jpg":
        media_type = "image/jpeg"
    return {
        "path": path,
        "filename": file.filename,
        "base64": b64,
        "media_type": media_type,
        "size_bytes": len(content),
    }


@app.get("/download/{filename}")
async def download_file(filename: str, user: dict = Depends(require_auth)):
    """Download a generated document."""
    from core.tools.files import GENERATED_DIR
    # Security: only allow alphanumeric, dots, underscores, hyphens
    safe_chars = set("abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789._-")
    if not all(c in safe_chars for c in filename):
        raise HTTPException(status_code=400, detail="Invalid filename")
    path = GENERATED_DIR / filename
    if not path.exists() or not path.is_file():
        raise HTTPException(status_code=404, detail="File not found")
    # Determine content type
    ext = path.suffix.lower()
    content_types = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".csv": "text/csv",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".json": "application/json",
    }
    ct = content_types.get(ext, "application/octet-stream")
    return Response(
        content=path.read_bytes(),
        media_type=ct,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# ==================== Conversation History Endpoints ====================

@app.get("/conversations")
async def conversations_list(limit: int = 50, offset: int = 0, project_id: str | None = None, user: dict = Depends(require_auth)):
    """List conversations ordered by most recent, optionally filtered by project."""
    return list_convos(limit=limit, offset=offset, project_id=project_id)


@app.post("/conversations")
async def conversations_create(user: dict = Depends(require_auth)):
    """Create a new conversation."""
    return create_conversation()


# Static routes MUST come before dynamic {conv_id} routes
@app.get("/conversations/search")
async def conversations_search(q: str, limit: int = 20, user: dict = Depends(require_auth)):
    """Search conversations using full-text search."""
    if not q or len(q) < 2:
        raise HTTPException(status_code=400, detail="Query must be at least 2 characters")
    return search_conversations(q, limit)


@app.get("/conversations/archived")
async def conversations_archived_list(limit: int = 50, offset: int = 0, user: dict = Depends(require_auth)):
    """List archived conversations."""
    return list_archived_conversations(limit=limit, offset=offset)


# Dynamic routes with {conv_id} parameter
@app.get("/conversations/{conv_id}")
async def conversations_get(conv_id: str, user: dict = Depends(require_auth)):
    """Get full conversation with messages."""
    conv = get_conversation(conv_id)
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    return conv


@app.delete("/conversations/{conv_id}")
async def conversations_delete(conv_id: str, user: dict = Depends(require_auth)):
    """Archive (soft-delete) a conversation."""
    if not archive_conversation(conv_id):
        raise HTTPException(status_code=404, detail="Conversation not found")
    return {"archived": True}


@app.put("/conversations/{conv_id}/title")
async def conversations_rename(conv_id: str, req: ConversationTitleRequest, user: dict = Depends(require_auth)):
    """Rename a conversation."""
    if not update_title(conv_id, req.title):
        raise HTTPException(status_code=404, detail="Conversation not found")
    return {"updated": True}


@app.put("/conversations/{conv_id}/project")
async def conversations_move_to_project(conv_id: str, req: ConversationProjectRequest, user: dict = Depends(require_auth)):
    """Move a conversation to a project or remove from project."""
    if not move_conversation_to_project(conv_id, req.project_id):
        raise HTTPException(status_code=404, detail="Conversation not found")
    return {"updated": True, "project_id": req.project_id}


@app.post("/conversations/{conv_id}/restore")
async def conversations_restore(conv_id: str, user: dict = Depends(require_auth)):
    """Restore an archived conversation."""
    if not restore_conversation(conv_id):
        raise HTTPException(status_code=404, detail="Archived conversation not found")
    return {"restored": True}


@app.delete("/conversations/{conv_id}/permanent")
async def conversations_delete_permanent(conv_id: str, user: dict = Depends(require_auth)):
    """Permanently delete a conversation."""
    if not delete_conversation_permanently(conv_id):
        raise HTTPException(status_code=404, detail="Conversation not found")
    return {"deleted": True}


# ==================== Projects ====================

@app.get("/projects")
async def projects_list(user: dict = Depends(require_auth)):
    """List all projects."""
    return list_projects()


@app.post("/projects")
async def projects_create(req: ProjectRequest, user: dict = Depends(require_auth)):
    """Create a new project."""
    return create_project(req.name, req.description, req.color)


@app.get("/projects/{project_id}")
async def projects_get(project_id: str, user: dict = Depends(require_auth)):
    """Get project details."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project


@app.put("/projects/{project_id}")
async def projects_update(project_id: str, req: ProjectUpdateRequest, user: dict = Depends(require_auth)):
    """Update a project."""
    if not update_project(project_id, req.name, req.description, req.color):
        raise HTTPException(status_code=404, detail="Project not found")
    return {"updated": True}


@app.delete("/projects/{project_id}")
async def projects_delete(project_id: str, user: dict = Depends(require_auth)):
    """Delete a project and all its references."""
    if not delete_project(project_id):
        raise HTTPException(status_code=404, detail="Project not found")
    return {"deleted": True}


@app.get("/projects/{project_id}/references")
async def projects_references_list(project_id: str, user: dict = Depends(require_auth)):
    """List all references for a project."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return list_references(project_id)


@app.post("/projects/{project_id}/references")
async def projects_references_add_note(project_id: str, req: ReferenceNoteRequest, user: dict = Depends(require_auth)):
    """Add a note reference to a project."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return add_reference(project_id, "note", req.title, req.content)


@app.post("/projects/{project_id}/references/upload")
async def projects_references_upload(project_id: str, file: UploadFile = File(...), user: dict = Depends(require_auth)):
    """Upload a file reference to a project."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Validate file type
    allowed_ext = {".pdf", ".txt", ".md", ".docx", ".png", ".jpg", ".jpeg", ".gif", ".webp"}
    ext = Path(file.filename).suffix.lower() if file.filename else ""
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")

    content = await file.read()
    if len(content) > 10 * 1024 * 1024:  # 10MB limit
        raise HTTPException(status_code=400, detail="File too large (max 10MB)")

    # Save file
    import uuid as uuid_mod
    file_id = uuid_mod.uuid4().hex[:12]
    safe_name = "".join(c for c in file.filename if c.isalnum() or c in "._-")[:50]
    relative_path = f"{project_id}/{file_id}_{safe_name}"
    file_path = UPLOADS_DIR / relative_path
    file_path.parent.mkdir(parents=True, exist_ok=True)
    file_path.write_bytes(content)

    # Extract text for searchability
    extracted_text = ""
    mime_type = file.content_type or "application/octet-stream"

    if ext == ".pdf":
        try:
            import PyPDF2
            import io
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            extracted_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            logger.warning(f"PDF text extraction failed: {e}")
    elif ext == ".docx":
        try:
            import docx
            import io
            doc = docx.Document(io.BytesIO(content))
            extracted_text = "\n".join(para.text for para in doc.paragraphs)
        except Exception as e:
            logger.warning(f"DOCX text extraction failed: {e}")
    elif ext in {".txt", ".md"}:
        try:
            extracted_text = content.decode("utf-8", errors="replace")
        except Exception:
            pass

    ref = add_reference(
        project_id,
        "file",
        file.filename,
        extracted_text or None,
        relative_path,
        mime_type,
        len(content),
    )
    return ref


@app.get("/projects/{project_id}/references/search")
async def projects_references_search(project_id: str, q: str, limit: int = 20, user: dict = Depends(require_auth)):
    """Search references within a project."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    if not q or len(q) < 2:
        raise HTTPException(status_code=400, detail="Query must be at least 2 characters")
    return search_references(project_id, q, limit)


@app.get("/projects/{project_id}/conversations")
async def projects_conversations_list(project_id: str, limit: int = 50, offset: int = 0, user: dict = Depends(require_auth)):
    """List conversations in a project."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return list_conversations_by_project(project_id, limit, offset)


# ==================== References ====================

@app.get("/references/{ref_id}")
async def references_get(ref_id: int, user: dict = Depends(require_auth)):
    """Get reference details."""
    ref = get_reference(ref_id)
    if not ref:
        raise HTTPException(status_code=404, detail="Reference not found")
    return ref


@app.put("/references/{ref_id}")
async def references_update(ref_id: int, req: ReferenceUpdateRequest, user: dict = Depends(require_auth)):
    """Update a reference."""
    if not update_reference(ref_id, req.title, req.content):
        raise HTTPException(status_code=404, detail="Reference not found")
    return {"updated": True}


@app.delete("/references/{ref_id}")
async def references_delete(ref_id: int, user: dict = Depends(require_auth)):
    """Delete a reference."""
    if not delete_reference(ref_id):
        raise HTTPException(status_code=404, detail="Reference not found")
    return {"deleted": True}


@app.get("/references/{ref_id}/download")
async def references_download(ref_id: int, user: dict = Depends(require_auth)):
    """Download a file reference."""
    ref = get_reference(ref_id)
    if not ref or not ref.get("file_path"):
        raise HTTPException(status_code=404, detail="File reference not found")
    file_path = UPLOADS_DIR / ref["file_path"]
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")
    return Response(
        content=file_path.read_bytes(),
        media_type=ref.get("file_type", "application/octet-stream"),
        headers={"Content-Disposition": f'attachment; filename="{ref["title"]}"'}
    )


# ==================== Core Endpoints ====================

@app.get("/")
async def root():
    # Serve React frontend if build exists
    index_html = _frontend_dist / "index.html"
    if index_html.exists():
        return HTMLResponse(content=index_html.read_text(), headers={"Cache-Control": "no-cache, no-store, must-revalidate"})
    # Fallback to embedded UI
    return HTMLResponse(content=CHAT_HTML, headers={"Cache-Control": "no-cache, no-store, must-revalidate"})


@app.get("/health")
async def health():
    return {
        "status": "online",
        "name": "Alfred",
        "version": "0.3.0",
        "timestamp": datetime.now().isoformat(),
        "llm_local": settings.ollama_model,
        "llm_cloud": settings.anthropic_model,
    }


@app.post("/chat", response_model=ChatResponse)
@limiter.limit("30/minute")
async def chat(request: Request, req: ChatRequest, user: dict = Depends(require_auth)):
    """Send a message to Alfred."""
    messages = await get_session_messages(req.session_id)
    messages.append({"role": "user", "content": req.message})

    # Handle image - check if user wants to upload vs analyze
    if req.image_base64 and req.image_media_type:
        msg_lower = (req.message or "").lower()
        upload_keywords = ["upload", "save", "put", "store", "send to", "nextcloud", "cloud", "drive"]
        wants_upload = any(kw in msg_lower for kw in upload_keywords)

        if wants_upload and req.image_path:
            # User wants to upload the file - pass to tool flow with file context
            # Extract filename from path for suggested destination
            from pathlib import Path
            filename = Path(req.image_path).name
            augmented_msg = f"[Attached file for upload - local_file_path: {req.image_path}, filename: {filename}]\nUser request: {req.message}"
            messages[-1]["content"] = augmented_msg
            # Continue to normal tool flow below (don't return here)
        else:
            # User wants image analysis - use OpenAI vision (gpt-4o-mini supports images)
            from core.brain.router import query_openai_vision
            response = await query_openai_vision(
                req.message or "Describe this image in detail.",
                req.image_base64,
                req.image_media_type
            )
            messages.append({"role": "assistant", "content": response})
            store_conversation(req.message, response, req.session_id)
            try:
                add_message(req.session_id, "user", req.message, "openai")
                add_message(req.session_id, "assistant", response, "openai")
            except Exception as e:
                logger.warning(f"Failed to persist vision message: {e}")
            return ChatResponse(response=response, tier="openai", timestamp=datetime.now().isoformat())

    # Handle document analysis request
    if req.document_path:
        from core.tools.files import parse_document
        doc_result = parse_document(req.document_path)
        if doc_result["error"]:
            response = f"Error reading document: {doc_result['error']}"
        else:
            # Include document content in context
            doc_text = doc_result["text"][:8000]  # Limit for context window
            if len(doc_result["text"]) > 8000:
                doc_text += "\n... (document truncated)"
            augmented_msg = f"[Document content from {doc_result['metadata']['filename']}]:\n{doc_text}\n\n{req.message}"
            messages[-1]["content"] = augmented_msg

    # Check memory for relevant context
    memories = recall(req.message)
    if memories:
        context = "\n".join(f"- {m['text']}" for m in memories[:3])
        messages[-1]["content"] = (
            f"[Relevant context from memory:\n{context}]\n\n" + messages[-1]["content"]
        )

    # Inject project context if conversation belongs to a project
    conv = get_conversation(req.session_id)
    if conv and conv.get("project_id"):
        project_context = get_project_context(conv["project_id"])
        if project_context:
            # Prepend project context to the system message
            if messages and messages[0]["role"] == "system":
                messages[0]["content"] = f"{messages[0]['content']}\n\n[Project Context]\n{project_context}"
            else:
                messages.insert(0, {"role": "system", "content": f"[Project Context]\n{project_context}"})

    tier = ModelTier(req.tier) if req.tier else classify_query(req.message, session_id=req.session_id)

    # Start long processing watcher — fires notifications if ask() takes too long
    from core.notifications.watcher import LongProcessingWatcher
    watcher = LongProcessingWatcher(query=req.message)
    watcher.start()
    try:
        result = await ask(req.message, messages=messages, tier=tier)
    finally:
        watcher.cancel()

    # Handle both old string return and new dict return
    if isinstance(result, dict):
        response = result["response"]
        images = result.get("images")
        ui_action = result.get("ui_action")
    else:
        response = result
        images = None
        ui_action = None

    messages.append({"role": "assistant", "content": response})
    store_conversation(req.message, response, req.session_id)

    # Persist to SQLite conversation history
    try:
        add_message(req.session_id, "user", req.message, tier.value)
        add_message(req.session_id, "assistant", response, tier.value)
    except Exception as e:
        logger.warning(f"Failed to persist conversation message: {e}")

    # Keep session history manageable (last 20 exchanges)
    if len(messages) > 42:
        _sessions[req.session_id] = [messages[0]] + messages[-40:]

    return ChatResponse(
        response=response,
        tier=tier.value,
        timestamp=datetime.now().isoformat(),
        images=images,
        ui_action=ui_action,
    )


@app.post("/chat/stream")
@limiter.limit("30/minute")
async def chat_stream(request: Request, req: ChatRequest, user: dict = Depends(require_auth)):
    """Stream a response from Alfred."""
    messages = await get_session_messages(req.session_id)
    messages.append({"role": "user", "content": req.message})

    tier = ModelTier(req.tier) if req.tier else classify_query(req.message, session_id=req.session_id)

    async def generate():
        full_response = []
        result = await ask(req.message, messages=messages, tier=tier, stream=True)

        # Handle both streaming and non-streaming responses
        if isinstance(result, str):
            # Claude Code returns a plain string
            full_response.append(result)
            yield result
        elif isinstance(result, dict):
            # Dict response (from Claude Code or non-streaming path)
            text = result.get("response", "")
            full_response.append(text)
            yield text
        else:
            # Actual stream (from Ollama)
            async for chunk in result:
                full_response.append(chunk)
                yield chunk

        response_text = "".join(full_response)
        messages.append({"role": "assistant", "content": response_text})
        store_conversation(req.message, response_text, req.session_id)
        # Persist to SQLite conversation history
        try:
            add_message(req.session_id, "user", req.message, tier.value)
            add_message(req.session_id, "assistant", response_text, tier.value)
        except Exception as e:
            logger.warning(f"Failed to persist streamed message: {e}")

    return StreamingResponse(generate(), media_type="text/plain")


@app.get("/voice/voices")
async def get_all_voices(user: dict = Depends(require_auth)):
    """Get all available voices for both TTS backends.

    Returns voices grouped by backend for easy selection in Shortcuts.
    """
    import requests as req

    result = {
        "kokoro": KOKORO_VOICES,
        "qwen3": [],
        "current_backend": settings.tts_model,
        "current_voice": settings.tts_voice,
    }

    # Get Qwen3 voices
    resources_dir = Path("/home/aialfred/qwen3-tts/resources")
    if resources_dir.exists():
        for f in resources_dir.iterdir():
            if f.suffix in [".mp3", ".wav"]:
                voice_id = f.stem
                result["qwen3"].append({
                    "id": voice_id,
                    "name": voice_id.replace("_", " "),
                    "desc": "Cloned voice",
                    "type": "clone"
                })

    # Get designed voices from Qwen3 API
    try:
        resp = req.get("http://localhost:7860/voice_design/voices", timeout=5)
        if resp.status_code == 200:
            for v in resp.json().get("voices", []):
                result["qwen3"].append({
                    "id": f"design:{v['name']}",
                    "name": f"{v['name']} (designed)",
                    "desc": v.get("description", "Designed voice")[:50],
                    "type": "design"
                })
    except Exception:
        pass

    return result


@app.post("/voice/transcribe")
@limiter.limit("20/minute")
async def transcribe_audio(request: Request, audio: UploadFile = File(...), user: dict = Depends(require_auth)):
    """Transcribe audio to text."""
    import asyncio
    from interfaces.voice.stt import transcribe
    audio_data = await audio.read()
    loop = asyncio.get_event_loop()
    text = await loop.run_in_executor(None, lambda: transcribe(audio_data, model_size=settings.whisper_model))
    return {"text": text}


@app.post("/voice/speak")
@limiter.limit("30/minute")
async def text_to_speech(request: Request, req: ChatRequest, user: dict = Depends(require_auth)):
    """Convert text to speech audio.

    Optional parameters:
    - voice: Voice ID (e.g., "Gwen_Stacy", "design:Natalie", "bm_daniel")
    - tts_backend: TTS backend ("kokoro", "qwen3")
    """
    import asyncio
    from interfaces.voice.tts import speak_with_options
    loop = asyncio.get_event_loop()
    audio_data = await loop.run_in_executor(
        None,
        lambda: speak_with_options(req.message, voice_id=req.voice, backend=req.tts_backend)
    )
    return Response(content=audio_data, media_type="audio/wav")


@app.post("/voice/chat/fast")
@limiter.limit("20/minute")
async def fast_voice_chat(request: Request, req: ChatRequest, user: dict = Depends(require_auth)):
    """Ultra-fast voice chat - streams first sentence immediately.

    Returns audio for the first sentence ASAP while generating the rest.
    Subsequent sentences are returned in the 'more_audio' field.

    Optional parameters:
    - voice: Voice ID (e.g., "Gwen_Stacy", "design:Natalie")
    - tts_backend: TTS backend ("kokoro", "qwen3")
    """
    import asyncio
    from interfaces.voice.tts import speak_with_options

    query = req.message
    conv_id = req.session_id  # Use session_id as conversation_id
    voice_id = req.voice
    tts_backend = req.tts_backend

    # Get conversation context
    messages = None
    if conv_id:
        conv = get_conversation(conv_id)
        if conv:
            messages = [{"role": "system", "content": await get_system_prompt(query)}]
            for msg in conv["messages"][-10:]:  # Last 10 messages for context
                messages.append({"role": msg["role"], "content": msg["content"]})
            messages.append({"role": "user", "content": query})

    # Stream from LLM and capture first sentence
    tier = classify_query(query, session_id=conv_id)

    if tier == ModelTier.CLAUDE_CODE:
        # Claude Code doesn't stream well, use regular path
        result = await ask(query, messages=messages, tier=tier, stream=False)
        response_text = result["response"] if isinstance(result, dict) else result
        loop = asyncio.get_event_loop()
        audio_data = await loop.run_in_executor(
            None,
            lambda: speak_with_options(response_text, voice_id=voice_id, backend=tts_backend)
        )
        return Response(content=audio_data, media_type="audio/wav")

    # Stream from local LLM
    stream = await ask(query, messages=messages, tier=tier, stream=True)

    # Collect text until first sentence boundary
    buffer = ""
    first_sentence = ""
    remaining = []
    sentence_end = re.compile(r'[.!?]\s*$')

    async for chunk in stream:
        buffer += chunk
        # Check for sentence boundary
        if not first_sentence and sentence_end.search(buffer):
            # Found first sentence!
            first_sentence = buffer.strip()
            buffer = ""
        elif first_sentence:
            # Collecting remaining text
            remaining.append(chunk)

    # If no sentence boundary found, use whole buffer
    if not first_sentence:
        first_sentence = buffer.strip()

    full_response = first_sentence + "".join(remaining)

    # Save to conversation
    if conv_id:
        add_message(conv_id, "user", query, "local")
        add_message(conv_id, "assistant", full_response, "local")

    # Generate TTS for first sentence (fast response)
    loop = asyncio.get_event_loop()
    first_audio = await loop.run_in_executor(
        None,
        lambda: speak_with_options(first_sentence, voice_id=voice_id, backend=tts_backend)
    )

    # Return first sentence audio immediately
    # Client can request remaining audio separately if needed
    return Response(
        content=first_audio,
        media_type="audio/wav",
        headers={
            "X-Full-Response": full_response[:500],  # Truncated for header
            "X-Has-More": "true" if remaining else "false",
        }
    )


# Pre-loaded acknowledgment audio cache
_ACK_CACHE = {}

def _load_acknowledgments():
    """Load pre-generated acknowledgment audio into memory."""
    global _ACK_CACHE
    cache_dir = Path("/home/aialfred/alfred/data/audio_cache")
    for wav_file in cache_dir.glob("*.wav"):
        _ACK_CACHE[wav_file.stem] = wav_file.read_bytes()
    logger.info(f"Loaded {len(_ACK_CACHE)} acknowledgment audio clips")

# Load on module import
_load_acknowledgments()


@app.post("/voice/chat/hybrid")
@limiter.limit("20/minute")
async def hybrid_voice_chat(request: Request, req: ChatRequest, user: dict = Depends(require_auth)):
    """Hybrid voice chat - instant acknowledgment + smart response.

    Returns a multipart response:
    1. Acknowledgment audio (instant, <100ms)
    2. Full response audio (smart model, streamed)

    Optional parameters:
    - voice: Voice ID (e.g., "Gwen_Stacy", "design:Natalie")
    - tts_backend: TTS backend ("kokoro", "qwen3")
    """
    import asyncio
    import random
    from interfaces.voice.tts import speak_with_options

    query = req.message
    conv_id = req.session_id  # Use session_id as conversation_id
    voice_id = req.voice
    tts_backend = req.tts_backend

    # Pick a random acknowledgment
    ack_keys = list(_ACK_CACHE.keys()) or ["one_moment"]
    ack_audio = _ACK_CACHE.get(random.choice(ack_keys), b"")

    async def generate_audio_stream():
        # First, yield acknowledgment immediately
        if ack_audio:
            # Send acknowledgment with a marker
            yield b"--ACKNOWLEDGE--"
            yield ack_audio
            yield b"--END_ACK--"

        # Now process with smart model (mistral-large)
        messages = None
        if conv_id:
            conv = get_conversation(conv_id)
            if conv:
                messages = [{"role": "system", "content": await get_system_prompt(query)}]
                for msg in conv["messages"][-10:]:
                    messages.append({"role": msg["role"], "content": msg["content"]})
                messages.append({"role": "user", "content": query})

        # Use configured model from settings
        import ollama as ollama_client
        smart_model = settings.ollama_model

        if messages is None:
            messages = [
                {"role": "system", "content": await get_system_prompt(query)},
                {"role": "user", "content": query},
            ]

        # Get response from smart model
        response = ollama_client.chat(model=smart_model, messages=messages)
        response_text = response["message"]["content"]

        # Save to conversation
        if conv_id:
            add_message(conv_id, "user", query, "local")
            add_message(conv_id, "assistant", response_text, "local")

        # Generate TTS for response
        loop = asyncio.get_event_loop()
        response_audio = await loop.run_in_executor(
            None,
            lambda: speak_with_options(response_text, voice_id=voice_id, backend=tts_backend)
        )

        # Send response audio
        yield b"--RESPONSE--"
        yield response_audio
        yield b"--END_RESPONSE--"

    return StreamingResponse(
        generate_audio_stream(),
        media_type="application/octet-stream",
        headers={"X-Audio-Format": "wav"}
    )


def _needs_acknowledgment(query: str) -> tuple[bool, list[str]]:
    """Determine if a query needs an acknowledgment and which ones are appropriate.

    Returns (needs_ack, appropriate_ack_keys).
    - Conversational queries (greetings, how are you) → no ack
    - Task queries (search, check, look up, create) → ack needed
    - Complex queries → ack needed
    """
    q = query.lower().strip()

    # Conversational/greeting patterns - NO acknowledgment needed
    no_ack_patterns = [
        # Greetings
        "hi", "hello", "hey", "good morning", "good afternoon", "good evening",
        "what's up", "whats up", "how are you", "how's it going", "hows it going",
        "yo", "sup", "howdy",
        # Simple questions about Alfred
        "who are you", "what are you", "what can you do", "help me",
        # Thank you / pleasantries
        "thank you", "thanks", "great", "perfect", "awesome", "cool", "ok", "okay",
        "goodbye", "bye", "good night", "see you", "later",
        # Very short queries (likely conversational)
    ]

    # Check for greeting patterns
    for pattern in no_ack_patterns:
        if q == pattern or q.startswith(pattern + " ") or q.startswith(pattern + ","):
            return False, []

    # Very short queries are likely conversational
    if len(q.split()) <= 3 and not any(kw in q for kw in ["search", "find", "check", "look", "get", "send", "create"]):
        return False, []

    # Task-oriented patterns - acknowledgment needed
    task_keywords = [
        "search", "find", "look up", "lookup", "check", "fetch", "get me",
        "send", "email", "message", "create", "add", "schedule", "remind",
        "calendar", "crm", "contact", "upload", "download", "server",
        "analyze", "summarize", "remember", "recall", "what did",
    ]

    # Pick appropriate ack types based on query
    for kw in task_keywords:
        if kw in q:
            # Searching/checking queries → "checking", "let_me_check"
            if any(x in q for x in ["search", "find", "check", "look", "recall", "what did"]):
                return True, ["checking", "let_me_check", "one_moment"]
            # Action queries → "right_away", "certainly"
            if any(x in q for x in ["send", "create", "add", "schedule", "upload"]):
                return True, ["right_away", "certainly"]
            # Default task
            return True, ["one_moment", "certainly"]

    # Long queries likely need processing time
    if len(q.split()) > 10:
        return True, ["one_moment", "certainly"]

    # Default: skip ack for shorter, unclear queries
    return False, []


@app.post("/voice/chat/ack")
async def get_acknowledgment(request: Request, req: ChatRequest = None, user: dict = Depends(require_auth)):
    """Get a contextually appropriate acknowledgment audio clip.

    Pass the query in ChatRequest.message to get smart acknowledgment selection.
    Returns empty audio for conversational queries that don't need an ack.
    """
    import random
    if not _ACK_CACHE:
        _load_acknowledgments()

    # Get query from request body if provided
    query = ""
    if req and req.message:
        query = req.message

    # Determine if we need an acknowledgment
    needs_ack, appropriate_keys = _needs_acknowledgment(query)

    if not needs_ack:
        # Return empty audio for conversational queries
        return Response(content=b"", media_type="audio/wav")

    # Filter to available keys
    available_keys = [k for k in appropriate_keys if k in _ACK_CACHE]
    if not available_keys:
        available_keys = list(_ACK_CACHE.keys())

    if not available_keys:
        return Response(content=b"", media_type="audio/wav")

    return Response(content=_ACK_CACHE[random.choice(available_keys)], media_type="audio/wav")


@app.post("/memory/store")
async def store(req: MemoryRequest, user: dict = Depends(require_auth)):
    """Store information in Alfred's memory."""
    doc_id = store_memory(req.text, req.category)
    return {"stored": True, "id": doc_id}


@app.post("/memory/recall")
async def recall_memories(req: RecallRequest, user: dict = Depends(require_auth)):
    """Recall information from Alfred's memory."""
    memories = recall(req.query, req.category, req.n_results)
    return {"memories": memories, "count": len(memories)}


@app.websocket("/ws")
async def websocket_chat(ws: WebSocket):
    """WebSocket endpoint for real-time chat."""
    # Authenticate via query param or cookie
    token = ws.query_params.get("token") or ws.cookies.get("alfred_token")
    if not token:
        await ws.close(code=4001, reason="Authentication required")
        return
    from core.security.auth import decode_token
    payload = decode_token(token)
    if not payload:
        await ws.close(code=4001, reason="Invalid token")
        return
    await ws.accept()
    session_id = f"ws_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    logger.info(f"WebSocket connected: {session_id} (user: {payload.get('sub', 'unknown')})")

    try:
        while True:
            data = await ws.receive_text()
            messages = await get_session_messages(session_id)
            messages.append({"role": "user", "content": data})

            tier = classify_query(data, session_id=session_id)
            stream = await ask(data, messages=messages, tier=tier, stream=True)

            full_response = []
            async for chunk in stream:
                full_response.append(chunk)
                await ws.send_text(chunk)

            await ws.send_text("\n[END]")
            response_text = "".join(full_response)
            messages.append({"role": "assistant", "content": response_text})

    except WebSocketDisconnect:
        logger.info(f"WebSocket disconnected: {session_id}")


# ==================== Notifications WebSocket ====================

@app.websocket("/ws/notifications")
async def websocket_notifications(ws: WebSocket):
    """WebSocket endpoint for push notifications (agent completion, alerts, etc.)."""
    # Authenticate via query param or cookie
    token = ws.query_params.get("token") or ws.cookies.get("alfred_token")
    if not token:
        await ws.close(code=4001, reason="Authentication required")
        return
    from core.security.auth import decode_token
    payload = decode_token(token)
    if not payload:
        await ws.close(code=4001, reason="Invalid token")
        return

    await ws.accept()
    user_id = payload.get("sub", "anonymous")

    # Register with notification manager
    from core.notifications import get_notification_manager
    manager = get_notification_manager()
    await manager.connect(ws, user_id)

    try:
        # Send welcome message
        await ws.send_json({
            "type": "connected",
            "message": "Notification channel established",
            "user": user_id,
        })

        # Keep connection alive - just wait for disconnect
        while True:
            try:
                # Wait for ping/pong or disconnect
                data = await ws.receive_text()
                # Echo pings back as pongs
                if data == "ping":
                    await ws.send_text("pong")
            except WebSocketDisconnect:
                break

    except WebSocketDisconnect:
        pass
    finally:
        await manager.disconnect(ws, user_id)
        logger.info(f"Notification client disconnected: {user_id}")


@app.get("/notifications/status")
async def notification_status(user: dict = Depends(require_auth)):
    """Get notification system status."""
    from core.notifications import get_notification_manager
    manager = get_notification_manager()
    return {
        "connected_clients": manager.connection_count,
        "push_subscriptions": manager.push_subscription_count,
        "websocket_url": "/ws/notifications",
    }


# ==================== Web Push (VAPID) Endpoints ====================

@app.get("/push/vapid-key")
async def get_vapid_key(user: dict = Depends(require_auth)):
    """Return the VAPID public key for browser push subscription."""
    return {"publicKey": settings.vapid_public_key}


class PushSubscriptionRequest(BaseModel):
    endpoint: str
    keys: dict  # {p256dh, auth}


@app.post("/push/subscribe")
async def push_subscribe(req: PushSubscriptionRequest, user: dict = Depends(require_auth)):
    """Register a browser push subscription."""
    from core.notifications import get_notification_manager
    manager = get_notification_manager()
    manager.add_push_subscription({
        "endpoint": req.endpoint,
        "keys": req.keys,
    })
    return {"status": "subscribed"}


class PushUnsubscribeRequest(BaseModel):
    endpoint: str


@app.post("/push/unsubscribe")
async def push_unsubscribe(req: PushUnsubscribeRequest, user: dict = Depends(require_auth)):
    """Remove a browser push subscription."""
    from core.notifications import get_notification_manager
    manager = get_notification_manager()
    manager.remove_push_subscription(req.endpoint)
    return {"status": "unsubscribed"}


# ==================== Wake Word WebSocket ====================

# Wake word detection state
_wakeword_model = None
_wakeword_initialized = False

def get_wakeword_model():
    """Get or initialize the wake word model."""
    global _wakeword_model, _wakeword_initialized
    if _wakeword_initialized:
        return _wakeword_model
    try:
        from interfaces.voice.wakeword import get_model
        _wakeword_model = get_model()
        _wakeword_initialized = True
        if _wakeword_model:
            logger.info("Wake word model loaded successfully")
        return _wakeword_model
    except Exception as e:
        logger.error(f"Failed to load wake word model: {e}")
        _wakeword_initialized = True
        return None


@app.websocket("/ws/wakeword")
async def websocket_wakeword(ws: WebSocket):
    """WebSocket endpoint for wake word detection streaming."""
    import numpy as np

    # Authenticate
    token = ws.query_params.get("token") or ws.cookies.get("alfred_token")
    if not token:
        await ws.close(code=4001, reason="Authentication required")
        return
    from core.security.auth import decode_token
    payload = decode_token(token)
    if not payload:
        await ws.close(code=4001, reason="Invalid token")
        return

    await ws.accept()
    logger.info(f"Wake word WebSocket connected: {payload.get('sub', 'unknown')}")

    # Get wake word model
    model = get_wakeword_model()
    if not model:
        await ws.send_json({"type": "error", "message": "Wake word model not available"})
        await ws.close()
        return

    await ws.send_json({"type": "ready", "message": "Wake word detection active"})

    # Audio buffer for processing
    audio_buffer = np.array([], dtype=np.int16)
    chunk_size = 1280  # 80ms at 16kHz
    threshold = 0.35  # Lowered for custom trained model
    log_counter = 0  # For periodic debug logging

    try:
        while True:
            # Receive binary audio data
            data = await ws.receive_bytes()

            # Handle keep-alive pings (empty data)
            if len(data) == 0:
                continue

            # Convert to numpy array (expecting 16-bit PCM)
            audio_chunk = np.frombuffer(data, dtype=np.int16)
            audio_buffer = np.concatenate([audio_buffer, audio_chunk])

            # Process in chunks
            while len(audio_buffer) >= chunk_size:
                process_chunk = audio_buffer[:chunk_size]
                audio_buffer = audio_buffer[chunk_size:]

                # Run prediction
                try:
                    prediction = model.predict(process_chunk)
                    for model_name, score in prediction.items():
                        # Log high scores for debugging (every 50th chunk to avoid spam)
                        log_counter += 1
                        if score > 0.1 or (log_counter % 50 == 0 and score > 0.01):
                            logger.debug(f"Wake word score: {score:.3f}")

                        if score >= threshold:
                            logger.info(f"Wake word detected! Score: {score:.3f}")
                            await ws.send_json({
                                "type": "detected",
                                "score": float(score),
                                "model": model_name
                            })
                            # Reset model state
                            model.reset()
                            audio_buffer = np.array([], dtype=np.int16)
                            break
                except Exception as e:
                    logger.error(f"Wake word prediction error: {e}")

    except WebSocketDisconnect:
        logger.info("Wake word WebSocket disconnected")
    except Exception as e:
        logger.error(f"Wake word WebSocket error: {e}")


# ==================== HTML UI ====================

CHAT_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0, viewport-fit=cover">
    <title>Alfred</title>
    <link rel="manifest" href="/manifest.json">
    <link rel="apple-touch-icon" href="/static/apple-touch-icon.png">
    <meta name="theme-color" content="#0a0a0a">
    <meta name="apple-mobile-web-app-capable" content="yes">
    <meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
    <meta name="apple-mobile-web-app-title" content="Alfred">
    <script src="https://cdn.jsdelivr.net/npm/onnxruntime-web@1.14.0/dist/ort.js"></script>
    <script src="https://cdn.jsdelivr.net/npm/@ricky0123/vad-web@0.0.18/dist/bundle.min.js"></script>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', system-ui, sans-serif;
            background: #0a0a0a; color: #e0e0e0; height: 100vh;
            display: flex; flex-direction: column;
            padding-top: env(safe-area-inset-top);
            padding-bottom: env(safe-area-inset-bottom);
        }
        header {
            padding: 12px 16px; border-bottom: 1px solid #2f2f2f;
            display: flex; align-items: center; gap: 12px;
            background: #0a0a0a;
        }
        header h1 { font-size: 16px; font-weight: 500; color: #e0e0e0; display: flex; align-items: center; }
        header h1 img { height: 28px; border-radius: 4px; }
        .inline-icon { height: 18px; border-radius: 3px; vertical-align: middle; margin-right: 2px; }
        header .status { display: none; }
        .header-right { margin-left: auto; display: flex; gap: 8px; align-items: center; flex-wrap: wrap; }
        .header-btn {
            background: transparent; border: none; color: #8e8e8e; padding: 8px 12px;
            border-radius: 8px; cursor: pointer; font-size: 13px;
        }
        .header-btn:hover { background: #2f2f2f; color: #e0e0e0; }
        #hamburger-btn {
            background: none; border: none; color: #8e8e8e; font-size: 20px;
            cursor: pointer; padding: 8px; border-radius: 8px; line-height: 1;
        }
        #hamburger-btn:hover { background: #2f2f2f; color: #e0e0e0; }
        #chat {
            flex: 1; overflow-y: auto; padding: 24px 48px 120px;
            display: flex; flex-direction: column; gap: 24px;
            width: 100%;
        }
        #chat.welcome-state { display: none; }

        /* Welcome Screen */
        #welcome-screen {
            flex: 1; display: flex; flex-direction: column;
            justify-content: center; align-items: center;
            padding: 24px 48px; display: none;
            width: 100%; height: 100%;
        }
        #welcome-screen.visible { display: flex; }
        #welcome-screen h2 {
            font-size: 32px; font-weight: 500; color: #e0e0e0;
            margin-bottom: 32px;
        }
        #welcome-input-container {
            width: 100%; max-width: 100%; padding: 0 10%;
        }
        #welcome-input-box {
            background: #2f2f2f; border-radius: 24px;
            padding: 12px 16px; display: flex; flex-direction: column;
            border: 1px solid #424242;
        }
        #welcome-input {
            background: transparent; border: none; color: #e0e0e0;
            font-size: 16px; outline: none; resize: none;
            min-height: 24px; max-height: 200px; padding: 4px 0;
        }
        #welcome-input::placeholder { color: #8e8e8e; }
        #welcome-input-actions {
            display: flex; align-items: center; justify-content: space-between;
            margin-top: 8px; padding-top: 8px;
        }
        #welcome-input-left { display: flex; align-items: center; gap: 8px; }
        #welcome-input-right { display: flex; align-items: center; gap: 8px; }
        .welcome-btn {
            background: transparent; border: none; color: #8e8e8e;
            cursor: pointer; padding: 8px; border-radius: 8px;
            font-size: 18px; display: flex; align-items: center; justify-content: center;
        }
        .welcome-btn:hover { background: #424242; color: #e0e0e0; }
        .mic-logo-btn {
            padding: 0; width: 36px; height: 36px; border-radius: 50%;
            overflow: hidden; background: transparent;
        }
        .mic-logo-btn img {
            width: 100%; height: 100%; object-fit: cover; border-radius: 50%;
        }
        .mic-logo-btn:hover { background: transparent; transform: scale(1.1); }
        .mic-logo-btn.recording {
            box-shadow: 0 0 12px rgba(232, 110, 44, 0.8);
            animation: pulse-glow 1s infinite;
        }
        @keyframes pulse-glow {
            0%, 100% { box-shadow: 0 0 8px rgba(232, 110, 44, 0.6); }
            50% { box-shadow: 0 0 16px rgba(232, 110, 44, 1); }
        }
        #welcome-send-btn {
            background: #fff; color: #000; border: none;
            width: 36px; height: 36px; border-radius: 50%;
            cursor: pointer; display: flex; align-items: center; justify-content: center;
            font-size: 16px;
        }
        #welcome-send-btn:hover { background: #d1d1d1; }
        #welcome-send-btn:disabled { background: #424242; color: #8e8e8e; cursor: not-allowed; }

        /* Message Styles - ChatGPT style */
        .msg { max-width: 100%; line-height: 1.6; }
        .msg.user {
            align-self: flex-end; max-width: 60%;
            background: #2f2f2f; color: #e0e0e0;
            padding: 12px 18px; border-radius: 20px;
        }
        .msg.alfred {
            align-self: flex-start; background: transparent; border: none;
            padding: 0; width: 100%;
        }
        .msg .label { display: none; }
        .msg.user .label { display: none; }
        .msg.alfred .label { display: none; }
        .msg .content { white-space: pre-wrap; font-family: inherit; font-size: 15px; }
        .msg .content code { background: #2f2f2f; padding: 2px 6px; border-radius: 4px; font-family: 'SF Mono', Monaco, monospace; font-size: 14px; }
        .msg .content strong { color: #fff; }
        .msg .content pre { background: #1e1e1e; border-radius: 8px; padding: 12px 16px; overflow-x: auto; margin: 12px 0; }
        .msg .content pre code { background: transparent; padding: 0; }

        /* Alfred message action buttons */
        .msg-actions {
            display: flex; align-items: center; gap: 4px;
            margin-top: 8px; opacity: 0; transition: opacity 0.2s;
        }
        .msg.alfred:hover .msg-actions { opacity: 1; }
        .msg-action-btn {
            background: transparent; border: none; color: #8e8e8e;
            cursor: pointer; padding: 6px 8px; border-radius: 6px;
            font-size: 14px; display: flex; align-items: center; justify-content: center;
        }
        .msg-action-btn:hover { background: #2f2f2f; color: #e0e0e0; }
        .msg-action-btn.active { color: #10b981; }
        .msg-action-btn svg { width: 18px; height: 18px; }

        /* Thinking indicator */
        #thinking {
            display: none; align-self: flex-start; padding: 8px 0;
            background: transparent; border: none; border-radius: 0;
            margin: 0;
        }
        #thinking.visible { display: flex; align-items: center; gap: 12px; }
        #thinking .label { font-size: 14px; color: #8e8e8e; }
        .morph-shape {
            width: 24px; height: 24px;
            background: linear-gradient(135deg, #4a9eff, #e86e2c);
            animation: morph 2s ease-in-out infinite, spin 3s linear infinite, colorShift 4s ease-in-out infinite;
        }
        @keyframes morph {
            0%, 100% { border-radius: 50%; }           /* circle */
            25% { border-radius: 0; }                  /* square */
            50% { border-radius: 50% 0 50% 0; }        /* diamond */
            75% { border-radius: 50% 50% 0 50%; }      /* leaf */
        }
        @keyframes spin {
            from { transform: rotate(0deg); }
            to { transform: rotate(360deg); }
        }
        @keyframes colorShift {
            0%, 100% { background: linear-gradient(135deg, #4a9eff, #e86e2c); }
            50% { background: linear-gradient(135deg, #e86e2c, #4ade80); }
        }
        #input-area {
            position: fixed; bottom: 0; left: 0; right: 0;
            padding: 16px 24px 24px;
            display: flex; justify-content: center;
            background: linear-gradient(transparent, #0a0a0a 20%);
            padding-bottom: max(24px, env(safe-area-inset-bottom));
        }
        #input-area.hidden { display: none; }
        #input-box {
            background: #2f2f2f; border-radius: 24px;
            padding: 12px 20px; display: flex; flex-direction: column;
            border: 1px solid #424242; width: 100%;
            margin: 0 10%;
        }
        #input {
            background: transparent; border: none; color: #e0e0e0;
            font-size: 15px; outline: none; resize: none;
            min-height: 24px; max-height: 200px; padding: 4px 0;
            flex: 1;
        }
        #input::placeholder { color: #8e8e8e; }
        #input:focus { border-color: transparent; }
        #input-actions {
            display: flex; align-items: center; justify-content: space-between;
            margin-top: 8px;
        }
        #input-left { display: flex; align-items: center; gap: 4px; }
        #input-right { display: flex; align-items: center; gap: 8px; }
        button {
            padding: 10px 20px; border-radius: 8px; border: none;
            background: #2563eb; color: white; font-size: 15px;
            cursor: pointer; font-weight: 500;
        }
        button:hover { background: #1d4ed8; }
        button:disabled { opacity: 0.5; cursor: not-allowed; }
        #send-btn {
            background: #fff; color: #000; border: none;
            width: 36px; height: 36px; border-radius: 50%;
            cursor: pointer; display: flex; align-items: center; justify-content: center;
            font-size: 16px; padding: 0;
        }
        #send-btn:hover { background: #d1d1d1; }
        #send-btn:disabled { background: #424242; color: #8e8e8e; cursor: not-allowed; }
        #mic-btn {
            background: transparent; width: 36px; height: 36px; border-radius: 50%;
            display: flex; align-items: center; justify-content: center;
            padding: 0; overflow: hidden; border: none; transition: all 0.2s;
            color: #8e8e8e; font-size: 18px;
        }
        #mic-btn img, #welcome-mic-btn img { width: 28px; height: 28px; object-fit: cover; border-radius: 50%; transition: opacity 0.2s; }
        #mic-btn:hover { background: #424242; color: #e0e0e0; }
        #mic-btn.recording, #welcome-mic-btn.recording { background: transparent; color: #fff; animation: pulse 1s infinite; }
        @keyframes pulse { 0%,100% { opacity:1; } 50% { opacity:0.6; } }
        @keyframes glow-listen { 0%,100% { box-shadow: 0 0 8px rgba(74,222,128,0.4); } 50% { box-shadow: 0 0 16px rgba(74,222,128,0.7); } }
        @keyframes glow-hear { 0%,100% { box-shadow: 0 0 10px rgba(232,110,44,0.5); } 50% { box-shadow: 0 0 20px rgba(232,110,44,0.9); } }
        @keyframes glow-process { 0% { box-shadow: 0 0 8px rgba(96,165,250,0.4); } 50% { box-shadow: 0 0 18px rgba(96,165,250,0.8); } 100% { box-shadow: 0 0 8px rgba(96,165,250,0.4); } }
        @keyframes glow-speak { 0%,100% { box-shadow: 0 0 8px rgba(167,139,250,0.4); } 50% { box-shadow: 0 0 16px rgba(167,139,250,0.7); } }
        #mic-btn.vad-listen { border-color: #4ade80; animation: glow-listen 2s ease-in-out infinite; }
        #mic-btn.vad-hear { border-color: #e86e2c; animation: glow-hear 0.6s ease-in-out infinite; }
        #mic-btn.vad-process { border-color: #60a5fa; animation: glow-process 1s linear infinite; }
        #mic-btn.vad-speak { border-color: #a78bfa; animation: glow-speak 1.5s ease-in-out infinite; }
        .tier-badge {
            font-size: 10px; padding: 2px 6px; border-radius: 4px;
            display: inline-block; margin-top: 4px;
        }
        .tier-local { background: #166534; color: #4ade80; }
        .tier-cloud { background: #1e3a5f; color: #60a5fa; }
        .tier-claude-code { background: #5b21b6; color: #c4b5fd; }
        .speak-btn {
            background: none; border: none; color: #666; cursor: pointer;
            font-size: 16px; padding: 2px 6px; border-radius: 4px;
            margin-left: 8px; vertical-align: middle;
        }
        .speak-btn:hover { color: #4a9eff; background: #1a1a2e; }
        .speak-btn.speaking { color: #4ade80; animation: pulse 1s infinite; }
        .stop-btn {
            display: none; background: none; border: none; color: #666; cursor: pointer;
            font-size: 14px; padding: 2px 6px; border-radius: 4px;
            margin-left: 4px; vertical-align: middle;
        }
        .stop-btn.visible { display: inline-block; }
        .stop-btn:hover { color: #f87171; background: #1a1a1a; }
        .upload-btn.has-file { color: #10b981; }
        #file-input { display: none; }
        #pending-file {
            display: none; padding: 8px 16px; background: #2f2f2f;
            border-radius: 12px; font-size: 13px;
            align-items: center; gap: 10px;
            position: fixed; bottom: 100px; left: 10%; right: 10%;
            z-index: 10; border: 1px solid #424242;
        }
        #pending-file.visible { display: flex; }
        #pending-file img { max-width: 80px; max-height: 60px; border-radius: 4px; }
        #pending-file .file-info { flex: 1; color: #ccc; }
        #pending-file .file-name { font-weight: 500; }
        #pending-file .file-meta { font-size: 11px; color: #666; margin-top: 2px; }
        #pending-file .remove-file {
            background: none; border: none; color: #666; cursor: pointer;
            font-size: 18px; padding: 4px;
        }
        #pending-file .remove-file:hover { color: #f87171; }
        .msg .download-btn {
            display: inline-block; margin-top: 8px; padding: 6px 12px;
            background: #2563eb; color: white; border-radius: 6px;
            text-decoration: none; font-size: 12px;
        }
        .msg .download-btn:hover { background: #1d4ed8; }
        .msg .inline-image { max-width: 100%; max-height: 300px; border-radius: 8px; margin-top: 8px; }
        .mode-toggle {
            display: inline-flex; align-items: center; gap: 6px;
            padding: 7px 14px; border-radius: 20px;
            border: 1px solid #2a2a2a; background: #161616;
            color: #555; cursor: pointer; font-size: 12px;
            font-weight: 500; transition: all 0.3s ease;
            -webkit-tap-highlight-color: transparent;
            letter-spacing: 0.2px;
        }
        .mode-toggle:hover { border-color: #444; color: #999; }
        .mode-toggle .dot {
            width: 6px; height: 6px; border-radius: 50%;
            background: #333; transition: all 0.3s ease;
        }
        #auto-speak-btn.active {
            border-color: rgba(74,222,128,0.3); background: #0a1f0e; color: #4ade80;
        }
        #auto-speak-btn.active .dot { background: #4ade80; box-shadow: 0 0 6px #4ade80; }
        #handsfree-btn.active {
            border-color: rgba(232,110,44,0.3); background: #1a0f08; color: #e86e2c;
        }
        #handsfree-btn.active .dot { background: #e86e2c; box-shadow: 0 0 6px #e86e2c; }
        #wakeword-btn.active {
            border-color: rgba(96,165,250,0.3); background: #0a1020; color: #60a5fa;
        }
        #wakeword-btn.active .dot { background: #60a5fa; box-shadow: 0 0 6px #60a5fa; animation: pulse 2s infinite; }
        #wakeword-status {
            display: none; padding: 4px 16px; text-align: center;
            font-size: 11px; color: #60a5fa; background: rgba(96,165,250,0.1);
            border-radius: 4px; margin: 0 8px;
        }
        #wakeword-status.visible { display: inline-block; }
        #vad-status {
            display: none; padding: 8px 16px; text-align: center;
            font-size: 12px; color: #8e8e8e;
            letter-spacing: 0.3px;
            position: fixed; bottom: 90px; left: 50%;
            transform: translateX(-50%);
            background: #2f2f2f; border-radius: 20px;
            border: 1px solid #424242;
        }
        #vad-status.visible { display: block; }
        #vad-status .vad-dot {
            display: inline-block; width: 6px; height: 6px; border-radius: 50%;
            margin-right: 6px; vertical-align: middle;
        }
        #vad-status.state-listen { color: #4ade80; }
        #vad-status.state-listen .vad-dot { background: #4ade80; animation: pulse 2s infinite; }
        #vad-status.state-hear { color: #e86e2c; }
        #vad-status.state-hear .vad-dot { background: #e86e2c; animation: pulse 0.5s infinite; }
        #vad-status.state-process { color: #60a5fa; }
        #vad-status.state-process .vad-dot { background: #60a5fa; animation: pulse 0.8s infinite; }
        #vad-status.state-speak { color: #a78bfa; }
        #vad-status.state-speak .vad-dot { background: #a78bfa; animation: pulse 1.2s infinite; }

        /* Login overlay */
        #login-overlay {
            position: fixed; top:0; left:0; width:100%; height:100%;
            background: #000;
            z-index: 100;
            display: flex; align-items: center; justify-content: center;
            overflow: hidden;
        }
        #matrix-canvas {
            position: absolute; top: 0; left: 0; width: 100%; height: 100%;
            z-index: 1;
        }
        #login-overlay.hidden { display: none; }
        #login-overlay.hidden #matrix-canvas { display: none; }
        .login-container {
            position: relative;
            display: flex; flex-direction: column; align-items: center;
            z-index: 2;
        }
        .login-box {
            background: rgba(0,0,0,0.85);
            border: 1px solid rgba(0,255,65,0.3);
            border-radius: 16px;
            padding: 40px 36px;
            width: 380px;
            backdrop-filter: blur(10px);
            box-shadow: 0 0 40px rgba(0,255,65,0.15), 0 20px 60px rgba(0,0,0,0.8), inset 0 0 60px rgba(0,255,65,0.03);
            animation: matrixBoxGlow 3s ease-in-out infinite;
        }
        @keyframes matrixBoxGlow {
            0%, 100% { box-shadow: 0 0 40px rgba(0,255,65,0.15), 0 20px 60px rgba(0,0,0,0.8), inset 0 0 60px rgba(0,255,65,0.03); }
            50% { box-shadow: 0 0 60px rgba(0,255,65,0.25), 0 20px 60px rgba(0,0,0,0.8), inset 0 0 80px rgba(0,255,65,0.05); }
        }
        .login-logo {
            display: block;
            width: 120px; height: 120px;
            margin: 0 auto 24px auto;
            border-radius: 50%;
            object-fit: cover;
            filter: drop-shadow(0 0 20px rgba(232,110,44,0.5)) drop-shadow(0 0 40px rgba(232,110,44,0.3));
            animation: logoGlow 3s ease-in-out infinite;
        }
        @keyframes logoGlow {
            0%, 100% { filter: drop-shadow(0 0 20px rgba(232,110,44,0.5)) drop-shadow(0 0 40px rgba(232,110,44,0.3)); }
            50% { filter: drop-shadow(0 0 30px rgba(232,110,44,0.7)) drop-shadow(0 0 60px rgba(232,110,44,0.4)); }
        }
        .login-box h2 {
            margin-bottom: 8px; color: #00ff41; font-size: 32px; text-align: center;
            font-weight: 600; letter-spacing: 4px;
            font-family: 'Courier New', monospace;
            text-transform: uppercase;
            text-shadow: 0 0 10px rgba(0,255,65,0.5), 0 0 20px rgba(0,255,65,0.3);
        }
        .login-subtitle {
            color: rgba(0,255,65,0.5); font-size: 12px; text-align: center;
            margin-bottom: 28px; font-weight: 400;
            font-family: 'Courier New', monospace;
            letter-spacing: 2px;
        }
        .login-box input {
            width: 100%; padding: 14px 16px; margin-bottom: 14px;
            border-radius: 8px; border: 1px solid rgba(0,255,65,0.2);
            background: rgba(0,255,65,0.03);
            color: #00ff41; font-size: 15px; outline: none;
            font-family: 'Courier New', monospace;
            transition: all 0.2s ease;
        }
        .login-box input:focus {
            border-color: rgba(0,255,65,0.5);
            background: rgba(0,255,65,0.08);
            box-shadow: 0 0 20px rgba(0,255,65,0.15);
        }
        .login-box input::placeholder { color: rgba(0,255,65,0.4); }
        .login-box button {
            width: 100%; margin-top: 8px;
            padding: 14px;
            background: linear-gradient(135deg, rgba(0,255,65,0.2) 0%, rgba(0,255,65,0.1) 100%);
            border: 1px solid rgba(0,255,65,0.4);
            border-radius: 8px;
            color: #00ff41;
            font-weight: 600; font-size: 15px;
            font-family: 'Courier New', monospace;
            text-transform: uppercase;
            letter-spacing: 2px;
            transition: all 0.2s ease;
            box-shadow: 0 4px 15px rgba(0,255,65,0.2);
        }
        .login-box button:hover {
            transform: translateY(-2px);
            background: linear-gradient(135deg, rgba(0,255,65,0.3) 0%, rgba(0,255,65,0.15) 100%);
            box-shadow: 0 6px 25px rgba(0,255,65,0.3);
        }
        .login-box button:active { transform: translateY(0); }
        .login-error { color: #f87171; font-size: 13px; margin-bottom: 8px; }
        .login-divider {
            display: flex; align-items: center; margin: 16px 0; color: #666;
        }
        .login-divider::before, .login-divider::after {
            content: ''; flex: 1; border-bottom: 1px solid #333;
        }
        .login-divider span { padding: 0 12px; font-size: 12px; }
        .passkey-btn {
            background: #1a1a1a; border: 1px solid #333; color: #e0e0e0;
            display: flex; align-items: center; justify-content: center; gap: 8px;
        }
        .passkey-btn:hover { background: #2a2a2a; border-color: #4a9eff; }
        .passkey-btn svg { width: 18px; height: 18px; }
        .totp-section { display: none; }
        .totp-section.visible { display: block; }
        .totp-input { text-align: center; font-size: 24px; letter-spacing: 8px; font-family: monospace; }
        .login-back { color: #4a9eff; cursor: pointer; font-size: 13px; margin-top: 12px; text-align: center; }
        .login-back:hover { text-decoration: underline; }

        /* Settings panel */
        #settings-panel {
            position: fixed; top:0; right:-420px; width:400px; height:100%;
            background: #111; border-left: 1px solid #333; z-index: 90;
            transition: right 0.3s; padding: 24px; overflow-y: auto;
        }
        #settings-panel.open { right: 0; }
        #settings-panel h2 { font-size: 18px; margin-bottom: 20px; color: #fff; }
        .setting-section { margin-bottom: 24px; }
        .setting-section h3 {
            font-size: 14px; color: #888; text-transform: uppercase;
            letter-spacing: 0.5px; margin-bottom: 10px;
        }
        .status-item {
            display: flex; justify-content: space-between; align-items: center;
            padding: 8px 0; border-bottom: 1px solid #1a1a1a;
        }
        .status-dot {
            width: 8px; height: 8px; border-radius: 50%; display: inline-block; margin-right: 6px;
        }
        .status-dot.green { background: #4ade80; }
        .status-dot.red { background: #f87171; }
        .status-dot.yellow { background: #facc15; }
        .connect-btn {
            padding: 4px 10px; font-size: 12px; border-radius: 4px;
            background: #2563eb; border: none; color: white; cursor: pointer;
        }
        .connect-btn:hover { background: #1d4ed8; }
        .tts-option {
            flex: 1; min-width: 120px; padding: 12px; border-radius: 8px;
            background: #1a1a1a; border: 2px solid #333; cursor: pointer;
            text-align: left; transition: all 0.2s;
        }
        .tts-option:hover { border-color: #555; }
        .tts-option.active { border-color: #3b82f6; background: #1e3a5f; }
        .tts-option .tts-name { display: block; font-weight: 600; color: #e0e0e0; font-size: 14px; }
        .tts-option .tts-desc { display: block; font-size: 11px; color: #888; margin-top: 4px; }
        .close-btn {
            position: absolute; top: 16px; right: 16px; background: none;
            border: none; color: #888; font-size: 24px; cursor: pointer;
        }
        .close-btn:hover { color: #fff; }
        #settings-overlay {
            position: fixed; top:0; left:0; width:100%; height:100%;
            background: rgba(0,0,0,0.5); z-index: 85; display: none;
        }
        #settings-overlay.open { display: block; }

        /* History drawer */
        #history-overlay {
            position: fixed; top:0; left:0; width:100%; height:100%;
            background: rgba(0,0,0,0.5); z-index: 85; display: none;
        }
        #history-overlay.open { display: block; }
        #history-panel {
            position: fixed; top:0; left:-320px; width:300px; height:100%;
            background: #111; border-right: 1px solid #333; z-index: 90;
            transition: left 0.3s; display: flex; flex-direction: column;
        }
        #history-panel.open { left: 0; }
        .history-header {
            padding: 12px 16px; border-bottom: 1px solid #222;
            display: flex; align-items: center; justify-content: space-between;
        }
        .history-header h2 { font-size: 16px; color: #fff; }
        #new-chat-btn {
            padding: 6px 14px; font-size: 13px; border-radius: 6px;
            background: #2563eb; border: none; color: white; cursor: pointer;
        }
        #new-chat-btn:hover { background: #1d4ed8; }

        /* Search bar */
        #search-container { padding: 8px 12px; border-bottom: 1px solid #222; }
        #search-input {
            width: 100%; padding: 8px 12px; border-radius: 6px;
            border: 1px solid #333; background: #1a1a1a; color: #e0e0e0;
            font-size: 13px; outline: none;
        }
        #search-input:focus { border-color: #4a9eff; }
        #search-input::placeholder { color: #666; }
        #search-results { display: none; padding: 8px; border-bottom: 1px solid #222; max-height: 200px; overflow-y: auto; }
        #search-results.visible { display: block; }
        .search-result { padding: 8px 10px; border-radius: 6px; cursor: pointer; margin-bottom: 4px; }
        .search-result:hover { background: #1a1a1a; }
        .search-result-title { font-size: 13px; color: #e0e0e0; }
        .search-result-snippet { font-size: 11px; color: #888; margin-top: 2px; }
        .search-result-snippet mark { background: #3b5998; color: #fff; padding: 0 2px; border-radius: 2px; }

        /* Sidebar sections */
        .sidebar-content { flex: 1; overflow-y: auto; }
        #sidebar-footer {
            padding: 12px 16px; border-top: 1px solid #2f2f2f;
            background: #0a0a0a;
        }
        #sidebar-settings-btn {
            display: flex; align-items: center; gap: 10px;
            width: 100%; padding: 10px 12px; border-radius: 8px;
            background: transparent; border: none; color: #8e8e8e;
            font-size: 14px; cursor: pointer; text-align: left;
        }
        #sidebar-settings-btn:hover { background: #2f2f2f; color: #e0e0e0; }
        #sidebar-settings-btn .settings-icon { font-size: 18px; }
        .sidebar-section { border-bottom: 1px solid #1a1a1a; }
        .section-header {
            padding: 10px 12px; display: flex; align-items: center; gap: 8px;
            cursor: pointer; user-select: none; color: #888; font-size: 12px;
            text-transform: uppercase; letter-spacing: 0.5px;
        }
        .section-header:hover { color: #ccc; }
        .section-header .arrow { transition: transform 0.2s; font-size: 10px; }
        .section-header.collapsed .arrow { transform: rotate(-90deg); }
        .section-content { display: block; }
        .section-header.collapsed + .section-content { display: none; }
        .section-badge {
            background: #3b82f6; color: #fff; font-size: 11px;
            padding: 2px 6px; border-radius: 10px; margin-left: auto;
            min-width: 18px; text-align: center;
        }

        /* Projects */
        .project-item {
            padding: 8px 12px; border-radius: 6px; cursor: pointer;
            margin: 2px 8px; display: flex; align-items: center; gap: 8px;
        }
        .project-item:hover { background: #1a1a1a; }
        .project-item.active { background: #1e3a5f; }
        .project-color { width: 8px; height: 8px; border-radius: 2px; flex-shrink: 0; }
        .project-name { font-size: 13px; color: #e0e0e0; flex: 1; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }
        .project-count { font-size: 11px; color: #666; }
        #add-project-btn {
            margin: 8px 12px; padding: 6px 12px; font-size: 12px;
            background: transparent; border: 1px dashed #333; color: #666;
            border-radius: 6px; cursor: pointer; width: calc(100% - 24px);
        }
        #add-project-btn:hover { border-color: #4a9eff; color: #4a9eff; }

        /* Conversation list */
        #conversation-list { padding: 4px 0; }
        .conv-item {
            padding: 10px 12px; border-radius: 8px; cursor: pointer;
            margin: 2px 8px; transition: background 0.15s;
        }
        .conv-item:hover { background: #1a1a1a; }
        .conv-item.active { background: #1e3a5f; }
        .conv-item-title {
            font-size: 14px; color: #e0e0e0; white-space: nowrap;
            overflow: hidden; text-overflow: ellipsis;
        }
        .conv-item-meta {
            font-size: 11px; color: #666; margin-top: 3px;
            display: flex; justify-content: space-between; align-items: center;
        }
        .conv-item-preview {
            white-space: nowrap; overflow: hidden; text-overflow: ellipsis;
            flex: 1; margin-right: 8px;
        }
        .conv-actions { display: flex; gap: 2px; visibility: hidden; }
        .conv-item:hover .conv-actions { visibility: visible; }
        .conv-action-btn {
            background: none; border: none; color: #555; cursor: pointer;
            font-size: 12px; padding: 2px 4px; border-radius: 3px;
        }
        .conv-action-btn:hover { color: #fff; background: #333; }
        .conv-action-btn.delete:hover { color: #f87171; }
        @media (pointer: coarse) { .conv-actions { visibility: visible !important; } }

        /* Project view header */
        #project-view-header {
            display: none; padding: 12px; border-bottom: 1px solid #222;
            background: #0f0f0f;
        }
        #project-view-header.visible { display: block; }
        .project-view-title { display: flex; align-items: center; gap: 8px; margin-bottom: 8px; }
        .project-view-name { font-size: 15px; font-weight: 500; color: #fff; }
        .project-view-desc { font-size: 12px; color: #888; margin-bottom: 8px; }
        .project-tabs { display: flex; gap: 4px; }
        .project-tab {
            padding: 6px 12px; font-size: 12px; border-radius: 4px;
            background: transparent; border: none; color: #888; cursor: pointer;
        }
        .project-tab:hover { color: #ccc; }
        .project-tab.active { background: #222; color: #fff; }

        /* References list */
        #references-list { display: none; padding: 8px; }
        #references-list.visible { display: block; }
        .ref-item {
            padding: 10px 12px; border-radius: 6px; margin-bottom: 4px;
            background: #1a1a1a; border: 1px solid #222;
        }
        .ref-item:hover { border-color: #333; }
        .ref-item-header { display: flex; align-items: center; gap: 8px; }
        .ref-type-icon { font-size: 14px; }
        .ref-title { font-size: 13px; color: #e0e0e0; flex: 1; }
        .ref-actions { display: flex; gap: 4px; }
        .ref-action-btn {
            background: none; border: none; color: #666; cursor: pointer;
            font-size: 12px; padding: 2px 4px;
        }
        .ref-action-btn:hover { color: #fff; }
        .ref-preview { font-size: 11px; color: #888; margin-top: 4px; max-height: 40px; overflow: hidden; }
        .ref-meta { font-size: 10px; color: #555; margin-top: 4px; }
        #add-ref-btn {
            margin-top: 8px; padding: 8px; font-size: 12px; width: 100%;
            background: transparent; border: 1px dashed #333; color: #666;
            border-radius: 6px; cursor: pointer;
        }
        #add-ref-btn:hover { border-color: #4a9eff; color: #4a9eff; }

        /* Modal */
        .modal-overlay {
            position: fixed; top: 0; left: 0; width: 100%; height: 100%;
            background: rgba(0,0,0,0.7); z-index: 100; display: none;
            align-items: center; justify-content: center;
        }
        .modal-overlay.visible { display: flex; }
        .modal {
            background: #111; border: 1px solid #333; border-radius: 12px;
            padding: 24px; width: 90%; max-width: 500px; max-height: 80vh;
            overflow-y: auto;
        }
        .modal h3 { font-size: 16px; color: #fff; margin-bottom: 16px; }
        .modal-input {
            width: 100%; padding: 10px 12px; margin-bottom: 12px;
            border-radius: 6px; border: 1px solid #333; background: #1a1a1a;
            color: #e0e0e0; font-size: 14px; outline: none;
        }
        .modal-input:focus { border-color: #4a9eff; }
        .modal-textarea {
            width: 100%; padding: 10px 12px; margin-bottom: 12px;
            border-radius: 6px; border: 1px solid #333; background: #1a1a1a;
            color: #e0e0e0; font-size: 14px; outline: none; resize: vertical;
            min-height: 100px; font-family: inherit;
        }
        .modal-textarea:focus { border-color: #4a9eff; }
        .modal-actions { display: flex; gap: 8px; justify-content: flex-end; margin-top: 16px; }
        .modal-btn {
            padding: 8px 16px; border-radius: 6px; font-size: 13px;
            cursor: pointer; border: none;
        }
        .modal-btn-primary { background: #2563eb; color: white; }
        .modal-btn-primary:hover { background: #1d4ed8; }
        .modal-btn-secondary { background: #333; color: #ccc; }
        .modal-btn-secondary:hover { background: #444; }
        .color-picker { display: flex; gap: 8px; margin-bottom: 12px; }
        .color-option {
            width: 24px; height: 24px; border-radius: 4px; cursor: pointer;
            border: 2px solid transparent;
        }
        .color-option:hover { opacity: 0.8; }
        .color-option.selected { border-color: #fff; }

        /* File drop zone */
        .file-drop-zone {
            border: 2px dashed #333; border-radius: 8px; padding: 24px;
            text-align: center; color: #666; margin-bottom: 12px;
            transition: all 0.2s;
        }
        .file-drop-zone.dragover { border-color: #4a9eff; background: rgba(74,158,255,0.1); }
        .file-drop-zone input { display: none; }

        /* Responsive: tablet / landscape phone */
        @media (max-width: 768px) {
            .msg { max-width: 90%; }
            #settings-panel { width: 100%; right: -100%; }
            #history-panel { width: 85%; left: -85%; }
            .login-box { width: 90%; max-width: 380px; padding: 32px 28px; }
            header { padding: 12px 16px; gap: 8px; }
            #input-area { padding: 12px 16px; }
        }

        /* Responsive: portrait phone */
        @media (max-width: 480px) {
            .msg { max-width: 85%; font-size: 14px; }
            .msg.user { padding: 10px 14px; max-width: 85%; }
            header { padding: 10px 12px; gap: 6px; }
            header h1 { font-size: 16px; }
            .header-right { gap: 4px; }
            .header-btn { padding: 6px 10px; font-size: 12px; }
            .mode-toggle { padding: 5px 10px; font-size: 11px; gap: 5px; }
            .mode-toggle .dot { width: 5px; height: 5px; }
            #vad-status { padding: 6px 12px; font-size: 11px; bottom: 80px; left: 5%; right: 5%; transform: none; }
            #input-area { padding: 12px 12px 16px; }
            #input-box { padding: 10px 14px; margin: 0 2%; }
            #input { font-size: 14px; }
            #welcome-screen { padding: 16px; }
            #welcome-screen h2 { font-size: 22px; }
            #welcome-input-container { padding: 0 2%; }
            #welcome-input-box { padding: 10px 14px; }
            #chat { padding: 16px 12px 100px; gap: 16px; }
            #history-panel { width: 100%; left: -100%; }
            .login-box { padding: 28px 20px; }
            .login-box h2 { font-size: 24px; letter-spacing: 2px; }
            #pending-file { bottom: 85px; left: 5%; right: 5%; }
            .msg-actions { opacity: 1; }
        }
    </style>
</head>
<body>
    <!-- Login Overlay -->
    <div id="login-overlay">
        <canvas id="matrix-canvas"></canvas>
        <div class="login-container">
            <div class="login-box">
                <img src="/static/logo.jpeg" alt="GR" class="login-logo">
                <div id="login-error" class="login-error" style="display:none"></div>

                <!-- Password login section -->
                <div id="password-section">
                    <input id="login-user" type="text" placeholder="Username" autocomplete="username">
                    <input id="login-pass" type="password" placeholder="Password" autocomplete="current-password"
                        onkeydown="if(event.key==='Enter')doLogin()">
                    <button onclick="doLogin()">Sign In</button>
                    <div class="login-divider"><span>or</span></div>
                    <button class="passkey-btn" onclick="doPasskeyLogin()">
                        <svg viewBox="0 0 24 24" fill="currentColor"><path d="M12.65 10A5.99 5.99 0 0 0 7 6c-3.31 0-6 2.69-6 6s2.69 6 6 6a5.99 5.99 0 0 0 5.65-4H17v4h4v-4h2v-4H12.65zM7 14c-1.1 0-2-.9-2-2s.9-2 2-2 2 .9 2 2-.9 2-2 2z"/></svg>
                        Sign in with Passkey
                    </button>
                </div>

            <!-- 2FA code section (hidden by default) -->
            <div id="totp-section" class="totp-section">
                <p style="color:#888;font-size:13px;margin-bottom:16px;text-align:center;">Enter your 2FA code from your authenticator app</p>
                <input id="totp-code" type="text" class="totp-input" placeholder="000000" maxlength="6"
                    autocomplete="one-time-code" inputmode="numeric"
                    onkeydown="if(event.key==='Enter')doTOTPLogin()">
                <button onclick="doTOTPLogin()">Verify</button>
                <div class="login-back" onclick="backToPassword()">Back to login</div>
            </div>
        </div>
        </div>
    </div>

    <!-- History Drawer -->
    <div id="history-overlay" onclick="toggleHistory()"></div>
    <div id="history-panel">
        <div class="history-header">
            <h2 id="sidebar-title">Conversations</h2>
            <button id="new-chat-btn" onclick="newConversation()">New Chat</button>
        </div>

        <!-- Search -->
        <div id="search-container">
            <input type="text" id="search-input" placeholder="Search conversations..." oninput="debounceSearch()">
        </div>
        <div id="search-results"></div>

        <!-- Project View Header (shown when viewing a project) -->
        <div id="project-view-header">
            <div class="project-view-title">
                <span class="project-color" id="project-view-color"></span>
                <span class="project-view-name" id="project-view-name"></span>
                <button class="conv-action-btn" onclick="exitProjectView()" title="Back">&#8592;</button>
                <button class="conv-action-btn" onclick="editCurrentProject()" title="Edit">&#9998;</button>
            </div>
            <div class="project-view-desc" id="project-view-desc"></div>
            <div class="project-tabs">
                <button class="project-tab active" id="tab-chats" onclick="showProjectTab('chats')">Chats</button>
                <button class="project-tab" id="tab-refs" onclick="showProjectTab('refs')">References</button>
            </div>
        </div>

        <!-- References List (in project view) -->
        <div id="references-list"></div>

        <div class="sidebar-content">
            <!-- Projects Section -->
            <div class="sidebar-section" id="projects-section">
                <div class="section-header" onclick="toggleSection('projects')">
                    <span class="arrow">&#9660;</span>
                    <span>Projects</span>
                </div>
                <div class="section-content" id="projects-content">
                    <div id="project-list"></div>
                    <button id="add-project-btn" onclick="showProjectModal()">+ New Project</button>
                </div>
            </div>

            <!-- Recent Chats Section -->
            <div class="sidebar-section" id="chats-section">
                <div class="section-header" onclick="toggleSection('chats')">
                    <span class="arrow">&#9660;</span>
                    <span>Recent Chats</span>
                </div>
                <div class="section-content" id="chats-content">
                    <div id="conversation-list"></div>
                </div>
            </div>

            <!-- Archived Section -->
            <div class="sidebar-section" id="archived-section">
                <div class="section-header collapsed" onclick="toggleSection('archived')">
                    <span class="arrow">&#9660;</span>
                    <span>Archived</span>
                    <span id="archived-count" class="section-badge"></span>
                </div>
                <div class="section-content" id="archived-content">
                    <div id="archived-list"></div>
                </div>
            </div>
        </div>

        <!-- Settings Button at Bottom -->
        <div id="sidebar-footer">
            <button id="sidebar-settings-btn" onclick="toggleSettings(); toggleHistory();">
                <span class="settings-icon">&#9881;</span>
                <span>Settings</span>
            </button>
        </div>
    </div>

    <!-- Project Modal -->
    <div class="modal-overlay" id="project-modal">
        <div class="modal">
            <h3 id="project-modal-title">New Project</h3>
            <input type="text" class="modal-input" id="project-name-input" placeholder="Project name">
            <textarea class="modal-textarea" id="project-desc-input" placeholder="Description (optional)"></textarea>
            <div class="color-picker" id="color-picker">
                <div class="color-option selected" style="background:#3b82f6" data-color="#3b82f6"></div>
                <div class="color-option" style="background:#10b981" data-color="#10b981"></div>
                <div class="color-option" style="background:#f59e0b" data-color="#f59e0b"></div>
                <div class="color-option" style="background:#ef4444" data-color="#ef4444"></div>
                <div class="color-option" style="background:#8b5cf6" data-color="#8b5cf6"></div>
                <div class="color-option" style="background:#ec4899" data-color="#ec4899"></div>
            </div>
            <div class="modal-actions">
                <button class="modal-btn modal-btn-secondary" onclick="hideProjectModal()">Cancel</button>
                <button class="modal-btn modal-btn-primary" onclick="saveProject()">Save</button>
            </div>
        </div>
    </div>

    <!-- Reference Modal -->
    <div class="modal-overlay" id="reference-modal">
        <div class="modal">
            <h3 id="reference-modal-title">Add Note</h3>
            <input type="text" class="modal-input" id="ref-title-input" placeholder="Title">
            <textarea class="modal-textarea" id="ref-content-input" placeholder="Note content (Markdown supported)" style="min-height:150px"></textarea>
            <div class="file-drop-zone" id="file-drop-zone" style="display:none">
                <div>Drop file here or click to upload</div>
                <div style="font-size:11px;margin-top:8px">PDF, TXT, MD, DOCX, PNG, JPG (max 10MB)</div>
                <input type="file" id="ref-file-input" accept=".pdf,.txt,.md,.docx,.png,.jpg,.jpeg,.gif,.webp">
            </div>
            <div class="modal-actions">
                <button class="modal-btn modal-btn-secondary" onclick="hideReferenceModal()">Cancel</button>
                <button class="modal-btn modal-btn-primary" onclick="saveReference()">Save</button>
            </div>
        </div>
    </div>

    <!-- Move to Project Modal -->
    <div class="modal-overlay" id="move-modal">
        <div class="modal">
            <h3>Move to Project</h3>
            <div id="move-project-list"></div>
            <div class="modal-actions">
                <button class="modal-btn modal-btn-secondary" onclick="hideMoveModal()">Cancel</button>
            </div>
        </div>
    </div>

    <!-- Settings Panel -->
    <div id="settings-overlay" onclick="toggleSettings()"></div>
    <div id="settings-panel">
        <button class="close-btn" onclick="toggleSettings()">&times;</button>
        <h2>Settings</h2>

        <div class="setting-section">
            <h3>Integrations</h3>
            <div id="integration-list">Loading...</div>
        </div>

        <div class="setting-section">
            <h3>LLM</h3>
            <div id="llm-info">Loading...</div>
        </div>

        <div class="setting-section">
            <h3>Text-to-Speech</h3>
            <div style="display:flex;gap:8px;flex-wrap:wrap;margin-bottom:12px">
                <button id="tts-kokoro" class="tts-option" onclick="setTTS('kokoro')">
                    <span class="tts-name">Kokoro</span>
                    <span class="tts-desc">Fast, local</span>
                </button>
                <button id="tts-qwen3" class="tts-option" onclick="setTTS('qwen3')">
                    <span class="tts-name">Qwen3</span>
                    <span class="tts-desc">High quality, cloning</span>
                </button>
            </div>
            <div style="margin-bottom:8px">
                <label style="font-size:12px;color:#888;display:block;margin-bottom:4px">Voice</label>
                <select id="tts-voice" onchange="setVoice(this.value)" style="width:100%;padding:8px 12px;border-radius:6px;border:1px solid #333;background:#1a1a1a;color:#e0e0e0;font-size:13px">
                    <option value="">Loading voices...</option>
                </select>
            </div>
            <div id="tts-status" style="font-size:12px;color:#888;margin-top:8px"></div>
        </div>

        <div class="setting-section">
            <h3>Account</h3>
            <div id="account-info"></div>
            <div style="margin-top:12px">
                <input id="new-pass" type="password" placeholder="New password" style="width:100%;padding:8px 12px;border-radius:6px;border:1px solid #333;background:#1a1a1a;color:#e0e0e0;font-size:13px;margin-bottom:8px;outline:none">
                <button class="header-btn" style="width:100%" onclick="changePassword()">Change Password</button>
                <div id="pass-msg" style="font-size:12px;margin-top:6px;display:none"></div>
            </div>
            <button class="header-btn" style="margin-top:12px;width:100%" onclick="doLogout()">Sign Out</button>
        </div>

        <div class="setting-section">
            <h3>Security</h3>

            <!-- 2FA Section -->
            <div style="margin-bottom:16px;padding:12px;background:#1a1a1a;border-radius:8px;border:1px solid #333">
                <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:8px">
                    <span style="font-size:14px;color:#e0e0e0">Two-Factor Authentication</span>
                    <span id="2fa-status" style="font-size:12px;padding:2px 8px;border-radius:4px"></span>
                </div>
                <p style="font-size:12px;color:#888;margin-bottom:8px">Add an extra layer of security with TOTP authentication</p>
                <div id="2fa-setup-area"></div>
                <button id="2fa-toggle-btn" class="header-btn" style="width:100%" onclick="toggle2FA()">Loading...</button>
            </div>

            <!-- Passkeys Section -->
            <div style="padding:12px;background:#1a1a1a;border-radius:8px;border:1px solid #333">
                <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:8px">
                    <span style="font-size:14px;color:#e0e0e0">Passkeys</span>
                    <span id="passkey-count" style="font-size:12px;color:#888"></span>
                </div>
                <p style="font-size:12px;color:#888;margin-bottom:8px">Sign in securely with Face ID, Touch ID, or security key</p>
                <div id="passkeys-list" style="margin-bottom:8px"></div>
                <button class="header-btn" style="width:100%" onclick="addPasskey()">Add Passkey</button>
            </div>
        </div>
    </div>

    <header>
        <button id="hamburger-btn" onclick="toggleHistory()" title="Conversation history">&#9776;</button>
        <h1><img src="/static/alfred-icon.jpg" alt="Alfred"></h1>
        <span class="status" id="header-status">Online</span>
        <div class="header-right">
            <button id="wakeword-btn" class="mode-toggle" onclick="toggleWakeWord()" title="'Hey Alfred' wake word detection"><span class="dot"></span>Hey Alfred</button>
            <span id="wakeword-status">Listening for "Hey Alfred"...</span>
            <button id="handsfree-btn" class="mode-toggle" onclick="toggleHandsFree()" title="Hands-free voice conversation"><span class="dot"></span>Hands-free</button>
            <button id="auto-speak-btn" class="mode-toggle" onclick="toggleAutoSpeak()" title="Auto-speak responses"><span class="dot"></span>Auto-speak</button>
            <button class="header-btn" onclick="toggleSettings()">&#9881;</button>
        </div>
    </header>

    <!-- Welcome Screen (shown when no messages) -->
    <div id="welcome-screen" class="visible">
        <h2>What can I help with?</h2>
        <div id="welcome-input-container">
            <div id="welcome-input-box">
                <textarea id="welcome-input" placeholder="Ask anything" rows="1"
                    onkeydown="if(event.key==='Enter'&&!event.shiftKey){event.preventDefault();sendFromWelcome()}"
                    oninput="autoResizeWelcome(this)"></textarea>
                <div id="welcome-input-actions">
                    <div id="welcome-input-left">
                        <button class="welcome-btn upload-btn" onclick="document.getElementById('file-input').click()" title="Attach">&#43;</button>
                    </div>
                    <div id="welcome-input-right">
                        <button class="welcome-btn mic-logo-btn" id="welcome-mic-btn" onclick="toggleMic()" title="Voice input"><img src="/static/logo-white.png" alt="Mic"></button>
                        <button id="welcome-send-btn" onclick="sendFromWelcome()" title="Send" disabled>&#9650;</button>
                    </div>
                </div>
            </div>
        </div>
    </div>

    <!-- Chat Area (shown after first message) -->
    <div id="chat" class="welcome-state">
        <div id="thinking">
            <div class="morph-shape"></div>
            <div class="label"><img src="/static/alfred-icon.jpg" alt="Alfred" class="inline-icon"> is thinking...</div>
        </div>
    </div>

    <div id="vad-status"><span class="vad-dot"></span><span id="vad-status-text">Listening...</span></div>
    <div id="pending-file">
        <img id="pending-preview" src="" alt="">
        <div class="file-info">
            <div class="file-name" id="pending-name"></div>
            <div class="file-meta" id="pending-meta"></div>
        </div>
        <button class="remove-file" onclick="clearPendingFile()" title="Remove">&times;</button>
    </div>

    <!-- Bottom Input (shown after first message) -->
    <div id="input-area" class="hidden">
        <div id="input-box">
            <textarea id="input" placeholder="Ask anything" rows="1"
                onkeydown="if(event.key==='Enter'&&!event.shiftKey){event.preventDefault();send()}"
                oninput="autoResizeInput(this)"></textarea>
            <div id="input-actions">
                <div id="input-left">
                    <button class="welcome-btn upload-btn" onclick="document.getElementById('file-input').click()" title="Attach">&#43;</button>
                </div>
                <div id="input-right">
                    <button class="welcome-btn mic-logo-btn" id="mic-btn" onclick="toggleMic()" title="Voice input"><img src="/static/logo-white.png" alt="Mic"></button>
                    <button id="send-btn" onclick="send()" title="Send">&#9650;</button>
                </div>
            </div>
        </div>
        <input type="file" id="file-input" accept=".pdf,.doc,.docx,.xls,.xlsx,.csv,.txt,.md,.json,.jpg,.jpeg,.png,.gif,.webp" onchange="handleFileSelect(event)">
    </div>
    <script>
        const chat = document.getElementById('chat');
        const input = document.getElementById('input');
        let mediaRecorder = null;
        let isRecording = false;
        let authToken = localStorage.getItem('alfred_token') || '';
        let currentConversationId = null;
        let loadingHistory = false;

        // ==================== State ====================
        let currentProjectId = null;
        let editingProjectId = null;
        let editingRefId = null;
        let projects = [];
        let searchTimeout = null;
        let isWelcomeState = true;

        // ==================== Welcome Screen ====================

        function showWelcomeState() {
            isWelcomeState = true;
            document.getElementById('welcome-screen').classList.add('visible');
            document.getElementById('chat').classList.add('welcome-state');
            document.getElementById('input-area').classList.add('hidden');
        }

        function showChatState() {
            isWelcomeState = false;
            document.getElementById('welcome-screen').classList.remove('visible');
            document.getElementById('chat').classList.remove('welcome-state');
            document.getElementById('input-area').classList.remove('hidden');
        }

        function autoResizeWelcome(el) {
            el.style.height = 'auto';
            el.style.height = Math.min(el.scrollHeight, 200) + 'px';
            // Enable/disable send button
            const btn = document.getElementById('welcome-send-btn');
            btn.disabled = !el.value.trim();
        }

        function autoResizeInput(el) {
            el.style.height = 'auto';
            el.style.height = Math.min(el.scrollHeight, 200) + 'px';
        }

        function sendFromWelcome() {
            const welcomeInput = document.getElementById('welcome-input');
            const text = welcomeInput.value.trim();
            if (!text && !pendingFile) return;
            // Transfer to main input and send
            document.getElementById('input').value = text;
            welcomeInput.value = '';
            showChatState();
            send();
        }

        // Enable welcome send button on input or pending file
        function updateWelcomeSendBtn() {
            const welcomeInput = document.getElementById('welcome-input');
            const btn = document.getElementById('welcome-send-btn');
            btn.disabled = !welcomeInput?.value.trim() && !pendingFile;
        }
        document.getElementById('welcome-input')?.addEventListener('input', updateWelcomeSendBtn);

        // ==================== Conversation History ====================

        function toggleHistory() {
            document.getElementById('history-panel').classList.toggle('open');
            document.getElementById('history-overlay').classList.toggle('open');
        }

        function toggleSection(section) {
            const header = document.querySelector(`#${section}-section .section-header`);
            header.classList.toggle('collapsed');
        }

        function formatRelativeDate(isoStr) {
            const d = new Date(isoStr);
            const now = new Date();
            const diffMs = now - d;
            const diffMin = Math.floor(diffMs / 60000);
            if (diffMin < 1) return 'just now';
            if (diffMin < 60) return diffMin + 'm ago';
            const diffHr = Math.floor(diffMin / 60);
            if (diffHr < 24) return diffHr + 'h ago';
            const diffDay = Math.floor(diffHr / 24);
            if (diffDay < 7) return diffDay + 'd ago';
            return d.toLocaleDateString();
        }

        function authHeaders(extra) {
            const h = authToken ? {'Authorization': 'Bearer ' + authToken} : {};
            return extra ? Object.assign(h, extra) : h;
        }

        async function loadConversations() {
            try {
                const url = currentProjectId ? `/conversations?project_id=${currentProjectId}` : '/conversations';
                const resp = await fetch(url, {headers: authHeaders()});
                const convs = await resp.json();
                const listEl = document.getElementById('conversation-list');
                if (!convs.length) {
                    listEl.innerHTML = '<div style="padding:16px;color:#666;font-size:13px;text-align:center">No conversations yet</div>';
                    return;
                }
                listEl.innerHTML = convs.map(c => {
                    const active = c.id === currentConversationId ? ' active' : '';
                    const title = c.title || 'New conversation';
                    const preview = c.last_message || '';
                    return `<div class="conv-item${active}" data-id="${c.id}" onclick="switchConversation('${c.id}')">
                        <div class="conv-item-title">${escapeHtml(title)}</div>
                        <div class="conv-item-meta">
                            <span class="conv-item-preview">${escapeHtml(preview)}</span>
                            <span>${formatRelativeDate(c.updated_at)}</span>
                            <div class="conv-actions">
                                <button class="conv-action-btn" onclick="event.stopPropagation();showMoveModal('${c.id}')" title="Move to project">&#128193;</button>
                                <button class="conv-action-btn" onclick="event.stopPropagation();archiveConversation('${c.id}')" title="Archive">&#128451;</button>
                                <button class="conv-action-btn delete" onclick="event.stopPropagation();deleteConversationPermanently('${c.id}')" title="Delete">&#10005;</button>
                            </div>
                        </div>
                    </div>`;
                }).join('');
            } catch(e) {
                console.error('Failed to load conversations:', e);
            }
        }

        async function loadArchivedConversations() {
            try {
                const resp = await fetch('/conversations/archived', {headers: authHeaders()});
                const convs = await resp.json();
                const listEl = document.getElementById('archived-list');

                // Update archived count badge
                const badge = document.getElementById('archived-count');
                if (badge) {
                    badge.textContent = convs.length || '';
                    badge.style.display = convs.length ? 'inline-block' : 'none';
                }

                if (!convs.length) {
                    listEl.innerHTML = '<div style="padding:16px;color:#666;font-size:13px;text-align:center">No archived chats</div>';
                    return;
                }
                listEl.innerHTML = convs.map(c => {
                    const title = c.title || 'Untitled conversation';
                    return `<div class="conv-item" data-id="${c.id}">
                        <div class="conv-item-title" style="color:#888">${escapeHtml(title)}</div>
                        <div class="conv-item-meta">
                            <span>${formatRelativeDate(c.updated_at)}</span>
                            <div class="conv-actions" style="visibility:visible">
                                <button class="conv-action-btn" onclick="event.stopPropagation();restoreConversation('${c.id}')" title="Restore">&#8634;</button>
                                <button class="conv-action-btn delete" onclick="event.stopPropagation();deleteConversationPermanently('${c.id}')" title="Delete permanently">&#10005;</button>
                            </div>
                        </div>
                    </div>`;
                }).join('');
            } catch(e) {
                console.error('Failed to load archived:', e);
            }
        }

        async function newConversation() {
            // Don't create conversation until user sends a message
            // This prevents empty conversations from cluttering the sidebar
            currentConversationId = null;
            clearChat();
            toggleHistory();
        }

        async function switchConversation(convId) {
            if (convId === currentConversationId) {
                toggleHistory();
                return;
            }
            try {
                const resp = await fetch('/conversations/' + convId, {headers: authHeaders()});
                if (!resp.ok) return;
                const data = await resp.json();
                currentConversationId = convId;
                clearChat();
                loadingHistory = true;
                if (data.messages && data.messages.length > 0) {
                    showChatState();
                    data.messages.forEach(m => {
                        if (m.role === 'user') addMsg(m.content, 'user', null, true);
                        else if (m.role === 'assistant') addMsg(m.content, 'alfred', m.tier, true);
                    });
                }
                loadingHistory = false;
                await loadConversations();
                toggleHistory();
            } catch(e) {
                console.error('Failed to load conversation:', e);
            }
        }

        async function archiveConversation(convId) {
            try {
                await fetch('/conversations/' + convId, {method: 'DELETE', headers: authHeaders()});
                if (convId === currentConversationId) {
                    currentConversationId = null;
                    clearChat();
                    const cr = await fetch('/conversations', {method: 'POST', headers: authHeaders()});
                    currentConversationId = (await cr.json()).id;
                }
                await loadConversations();
                await loadArchivedConversations();
            } catch(e) {
                console.error('Failed to archive:', e);
            }
        }

        async function restoreConversation(convId) {
            try {
                await fetch(`/conversations/${convId}/restore`, {method: 'POST', headers: authHeaders()});
                await loadConversations();
                await loadArchivedConversations();
            } catch(e) {
                console.error('Failed to restore:', e);
            }
        }

        async function deleteConversationPermanently(convId) {
            if (!confirm('Permanently delete this conversation? This cannot be undone.')) return;
            try {
                await fetch(`/conversations/${convId}/permanent`, {method: 'DELETE', headers: authHeaders()});
                await loadArchivedConversations();
            } catch(e) {
                console.error('Failed to delete permanently:', e);
            }
        }

        function clearChat() {
            chat.innerHTML = '<div id="thinking"><div class="morph-shape"></div><div class="label"><img src="/static/alfred-icon.jpg" alt="Alfred" class="inline-icon"> is thinking...</div></div>';
            msgTexts = {};
            msgCounter = 0;
            showWelcomeState();
        }

        async function initConversations() {
            await loadProjects();
            await loadConversations();
            await loadArchivedConversations();
            // Don't create a conversation until user actually sends a message
            // This prevents empty conversations from cluttering the sidebar
            currentConversationId = null;
            document.getElementById('messages').innerHTML = '';
            // Ensure menu stays closed on load
            document.getElementById('history-panel').classList.remove('open');
            document.getElementById('history-overlay').classList.remove('open');
        }

        // ==================== Search ====================

        function debounceSearch() {
            clearTimeout(searchTimeout);
            searchTimeout = setTimeout(doSearch, 300);
        }

        async function doSearch() {
            const query = document.getElementById('search-input').value.trim();
            const resultsEl = document.getElementById('search-results');
            if (query.length < 2) {
                resultsEl.classList.remove('visible');
                return;
            }
            try {
                const resp = await fetch(`/conversations/search?q=${encodeURIComponent(query)}`, {headers: authHeaders()});
                const results = await resp.json();
                if (!results.length) {
                    resultsEl.innerHTML = '<div style="padding:12px;color:#666;font-size:12px;text-align:center">No results found</div>';
                } else {
                    resultsEl.innerHTML = results.map(r => `
                        <div class="search-result" onclick="switchConversation('${r.id}')">
                            <div class="search-result-title">${escapeHtml(r.title || 'Untitled')}</div>
                            <div class="search-result-snippet">${r.snippet || ''}</div>
                        </div>
                    `).join('');
                }
                resultsEl.classList.add('visible');
            } catch(e) {
                console.error('Search failed:', e);
            }
        }

        // ==================== Projects ====================

        async function loadProjects() {
            try {
                const resp = await fetch('/projects', {headers: authHeaders()});
                projects = await resp.json();
                const listEl = document.getElementById('project-list');
                if (!projects.length) {
                    listEl.innerHTML = '';
                    return;
                }
                listEl.innerHTML = projects.map(p => `
                    <div class="project-item${currentProjectId === p.id ? ' active' : ''}" onclick="viewProject('${p.id}')">
                        <span class="project-color" style="background:${p.color}"></span>
                        <span class="project-name">${escapeHtml(p.name)}</span>
                        <span class="project-count">${p.conversation_count || 0}</span>
                    </div>
                `).join('');
            } catch(e) {
                console.error('Failed to load projects:', e);
            }
        }

        async function viewProject(projectId) {
            currentProjectId = projectId;
            const project = projects.find(p => p.id === projectId);
            if (!project) return;

            // Show project view header
            document.getElementById('project-view-header').classList.add('visible');
            document.getElementById('project-view-color').style.background = project.color;
            document.getElementById('project-view-name').textContent = project.name;
            document.getElementById('project-view-desc').textContent = project.description || '';

            // Hide sidebar sections, show project chats
            document.getElementById('projects-section').style.display = 'none';
            document.getElementById('archived-section').style.display = 'none';
            document.querySelector('#chats-section .section-header').style.display = 'none';

            showProjectTab('chats');
            await loadConversations();
            await loadProjects();
        }

        function exitProjectView() {
            currentProjectId = null;
            document.getElementById('project-view-header').classList.remove('visible');
            document.getElementById('references-list').classList.remove('visible');
            document.getElementById('projects-section').style.display = '';
            document.getElementById('archived-section').style.display = '';
            document.querySelector('#chats-section .section-header').style.display = '';
            loadConversations();
            loadProjects();
        }

        function showProjectTab(tab) {
            document.getElementById('tab-chats').classList.toggle('active', tab === 'chats');
            document.getElementById('tab-refs').classList.toggle('active', tab === 'refs');
            document.getElementById('conversation-list').parentElement.style.display = tab === 'chats' ? '' : 'none';
            document.getElementById('references-list').classList.toggle('visible', tab === 'refs');
            if (tab === 'refs') loadReferences();
        }

        function showProjectModal(projectId = null) {
            editingProjectId = projectId;
            document.getElementById('project-modal-title').textContent = projectId ? 'Edit Project' : 'New Project';
            if (projectId) {
                const p = projects.find(pr => pr.id === projectId);
                if (p) {
                    document.getElementById('project-name-input').value = p.name;
                    document.getElementById('project-desc-input').value = p.description || '';
                    selectColor(p.color);
                }
            } else {
                document.getElementById('project-name-input').value = '';
                document.getElementById('project-desc-input').value = '';
                selectColor('#3b82f6');
            }
            document.getElementById('project-modal').classList.add('visible');
        }

        function hideProjectModal() {
            document.getElementById('project-modal').classList.remove('visible');
            editingProjectId = null;
        }

        function selectColor(color) {
            document.querySelectorAll('.color-option').forEach(el => {
                el.classList.toggle('selected', el.dataset.color === color);
            });
        }

        document.getElementById('color-picker').addEventListener('click', e => {
            if (e.target.classList.contains('color-option')) {
                selectColor(e.target.dataset.color);
            }
        });

        async function saveProject() {
            const name = document.getElementById('project-name-input').value.trim();
            if (!name) return alert('Please enter a project name');
            const description = document.getElementById('project-desc-input').value.trim();
            const color = document.querySelector('.color-option.selected')?.dataset.color || '#3b82f6';

            try {
                if (editingProjectId) {
                    await fetch(`/projects/${editingProjectId}`, {
                        method: 'PUT',
                        headers: authHeaders({'Content-Type': 'application/json'}),
                        body: JSON.stringify({name, description, color})
                    });
                } else {
                    await fetch('/projects', {
                        method: 'POST',
                        headers: authHeaders({'Content-Type': 'application/json'}),
                        body: JSON.stringify({name, description, color})
                    });
                }
                hideProjectModal();
                await loadProjects();
                if (currentProjectId === editingProjectId) {
                    viewProject(currentProjectId);
                }
            } catch(e) {
                console.error('Failed to save project:', e);
            }
        }

        function editCurrentProject() {
            if (currentProjectId) showProjectModal(currentProjectId);
        }

        async function deleteProject(projectId) {
            if (!confirm('Delete this project and all its references? Conversations will be kept.')) return;
            try {
                await fetch(`/projects/${projectId}`, {method: 'DELETE', headers: authHeaders()});
                if (currentProjectId === projectId) exitProjectView();
                await loadProjects();
            } catch(e) {
                console.error('Failed to delete project:', e);
            }
        }

        // ==================== References ====================

        async function loadReferences() {
            if (!currentProjectId) return;
            try {
                const resp = await fetch(`/projects/${currentProjectId}/references`, {headers: authHeaders()});
                const refs = await resp.json();
                const listEl = document.getElementById('references-list');
                let html = '';
                if (!refs.length) {
                    html = '<div style="padding:16px;color:#666;font-size:13px;text-align:center">No references yet</div>';
                } else {
                    html = refs.map(r => {
                        const icon = r.type === 'note' ? '&#128221;' : '&#128196;';
                        const preview = r.content ? r.content.substring(0, 100) + (r.content.length > 100 ? '...' : '') : '';
                        const meta = r.type === 'file' ? `${r.file_type || 'file'} - ${formatFileSize(r.file_size)}` : '';
                        return `<div class="ref-item">
                            <div class="ref-item-header">
                                <span class="ref-type-icon">${icon}</span>
                                <span class="ref-title">${escapeHtml(r.title)}</span>
                                <div class="ref-actions">
                                    ${r.type === 'note' ? `<button class="ref-action-btn" onclick="editReference(${r.id})" title="Edit">&#9998;</button>` : ''}
                                    ${r.type === 'file' ? `<button class="ref-action-btn" onclick="downloadReference(${r.id})" title="Download">&#8681;</button>` : ''}
                                    <button class="ref-action-btn" onclick="deleteReference(${r.id})" title="Delete">&#10005;</button>
                                </div>
                            </div>
                            ${preview ? `<div class="ref-preview">${escapeHtml(preview)}</div>` : ''}
                            ${meta ? `<div class="ref-meta">${meta}</div>` : ''}
                        </div>`;
                    }).join('');
                }
                html += `<button id="add-ref-btn" onclick="showReferenceModal()">+ Add Reference</button>`;
                listEl.innerHTML = html;
            } catch(e) {
                console.error('Failed to load references:', e);
            }
        }

        function formatFileSize(bytes) {
            if (!bytes) return '';
            if (bytes < 1024) return bytes + ' B';
            if (bytes < 1024*1024) return (bytes/1024).toFixed(1) + ' KB';
            return (bytes/1024/1024).toFixed(1) + ' MB';
        }

        function showReferenceModal(refId = null, isFile = false) {
            editingRefId = refId;
            document.getElementById('reference-modal-title').textContent = refId ? 'Edit Note' : (isFile ? 'Upload File' : 'Add Note');
            document.getElementById('ref-title-input').value = '';
            document.getElementById('ref-content-input').value = '';
            document.getElementById('ref-content-input').style.display = isFile ? 'none' : '';
            document.getElementById('file-drop-zone').style.display = isFile ? 'block' : 'none';
            if (refId) {
                fetch(`/references/${refId}`, {headers: authHeaders()})
                    .then(r => r.json())
                    .then(ref => {
                        document.getElementById('ref-title-input').value = ref.title;
                        document.getElementById('ref-content-input').value = ref.content || '';
                    });
            }
            document.getElementById('reference-modal').classList.add('visible');
        }

        function hideReferenceModal() {
            document.getElementById('reference-modal').classList.remove('visible');
            editingRefId = null;
        }

        async function saveReference() {
            const title = document.getElementById('ref-title-input').value.trim();
            const content = document.getElementById('ref-content-input').value;
            const fileInput = document.getElementById('ref-file-input');

            if (fileInput.files.length > 0) {
                // File upload
                const form = new FormData();
                form.append('file', fileInput.files[0]);
                try {
                    await fetch(`/projects/${currentProjectId}/references/upload`, {
                        method: 'POST',
                        headers: authHeaders(),
                        body: form
                    });
                    hideReferenceModal();
                    fileInput.value = '';
                    await loadReferences();
                } catch(e) {
                    console.error('Upload failed:', e);
                    alert('Upload failed');
                }
            } else if (title) {
                // Note
                try {
                    if (editingRefId) {
                        await fetch(`/references/${editingRefId}`, {
                            method: 'PUT',
                            headers: authHeaders({'Content-Type': 'application/json'}),
                            body: JSON.stringify({title, content})
                        });
                    } else {
                        await fetch(`/projects/${currentProjectId}/references`, {
                            method: 'POST',
                            headers: authHeaders({'Content-Type': 'application/json'}),
                            body: JSON.stringify({title, content})
                        });
                    }
                    hideReferenceModal();
                    await loadReferences();
                } catch(e) {
                    console.error('Save failed:', e);
                }
            }
        }

        async function editReference(refId) {
            showReferenceModal(refId, false);
        }

        async function deleteReference(refId) {
            if (!confirm('Delete this reference?')) return;
            try {
                await fetch(`/references/${refId}`, {method: 'DELETE', headers: authHeaders()});
                await loadReferences();
            } catch(e) {
                console.error('Delete failed:', e);
            }
        }

        function downloadReference(refId) {
            window.open(`/references/${refId}/download`, '_blank');
        }

        // File drop zone
        const dropZone = document.getElementById('file-drop-zone');
        dropZone.addEventListener('click', () => document.getElementById('ref-file-input').click());
        dropZone.addEventListener('dragover', e => { e.preventDefault(); dropZone.classList.add('dragover'); });
        dropZone.addEventListener('dragleave', () => dropZone.classList.remove('dragover'));
        dropZone.addEventListener('drop', e => {
            e.preventDefault();
            dropZone.classList.remove('dragover');
            if (e.dataTransfer.files.length) {
                document.getElementById('ref-file-input').files = e.dataTransfer.files;
                dropZone.querySelector('div').textContent = e.dataTransfer.files[0].name;
            }
        });

        // ==================== Move to Project ====================

        let movingConvId = null;

        function showMoveModal(convId) {
            movingConvId = convId;
            const listEl = document.getElementById('move-project-list');
            let html = `<div class="project-item" onclick="moveToProject(null)">
                <span class="project-name">No Project</span>
            </div>`;
            html += projects.map(p => `
                <div class="project-item" onclick="moveToProject('${p.id}')">
                    <span class="project-color" style="background:${p.color}"></span>
                    <span class="project-name">${escapeHtml(p.name)}</span>
                </div>
            `).join('');
            listEl.innerHTML = html;
            document.getElementById('move-modal').classList.add('visible');
        }

        function hideMoveModal() {
            document.getElementById('move-modal').classList.remove('visible');
            movingConvId = null;
        }

        async function moveToProject(projectId) {
            if (!movingConvId) return;
            try {
                await fetch(`/conversations/${movingConvId}/project`, {
                    method: 'PUT',
                    headers: authHeaders({'Content-Type': 'application/json'}),
                    body: JSON.stringify({project_id: projectId})
                });
                hideMoveModal();
                await loadConversations();
                await loadProjects();
            } catch(e) {
                console.error('Move failed:', e);
            }
        }

        // ==================== Auth ====================

        (async function checkAuth() {
            try {
                const resp = await fetch('/auth/me', {headers: authHeaders()});
                const data = await resp.json();
                if (data.authenticated) {
                    document.getElementById('login-overlay').classList.add('hidden');
                    loadIntegrations();
                    initConversations();
                    // Auto-enable Hey Alfred wake word on startup
                    setTimeout(() => {
                        toggleWakeWord();
                        // Force menu closed after everything initializes
                        document.getElementById('history-panel').classList.remove('open');
                        document.getElementById('history-overlay').classList.remove('open');
                    }, 500);
                    // Extra delayed close in case async operations open it
                    setTimeout(() => {
                        document.getElementById('history-panel').classList.remove('open');
                        document.getElementById('history-overlay').classList.remove('open');
                    }, 1200);
                }
            } catch(e) {}
        })();

        let pendingTOTPUsername = '';

        async function doLogin() {
            const user = document.getElementById('login-user').value;
            const pass = document.getElementById('login-pass').value;
            const errEl = document.getElementById('login-error');
            errEl.style.display = 'none';

            try {
                const resp = await fetch('/auth/login', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({username: user, password: pass})
                });
                if (!resp.ok) {
                    errEl.textContent = 'Invalid credentials';
                    errEl.style.display = 'block';
                    return;
                }
                const data = await resp.json();

                // Check if 2FA is required
                if (data.requires_2fa) {
                    pendingTOTPUsername = data.username;
                    document.getElementById('password-section').style.display = 'none';
                    document.getElementById('totp-section').classList.add('visible');
                    document.getElementById('totp-code').focus();
                    return;
                }

                completeLogin(data);
            } catch(e) {
                errEl.textContent = 'Connection error';
                errEl.style.display = 'block';
            }
        }

        async function doTOTPLogin() {
            const code = document.getElementById('totp-code').value.trim();
            const errEl = document.getElementById('login-error');
            errEl.style.display = 'none';

            if (code.length !== 6) {
                errEl.textContent = 'Please enter a 6-digit code';
                errEl.style.display = 'block';
                return;
            }

            try {
                const resp = await fetch('/auth/2fa/login', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({username: pendingTOTPUsername, code: code})
                });
                if (!resp.ok) {
                    errEl.textContent = 'Invalid 2FA code';
                    errEl.style.display = 'block';
                    document.getElementById('totp-code').value = '';
                    return;
                }
                const data = await resp.json();
                completeLogin(data);
            } catch(e) {
                errEl.textContent = 'Connection error';
                errEl.style.display = 'block';
            }
        }

        function backToPassword() {
            pendingTOTPUsername = '';
            document.getElementById('totp-section').classList.remove('visible');
            document.getElementById('password-section').style.display = 'block';
            document.getElementById('totp-code').value = '';
            document.getElementById('login-error').style.display = 'none';
        }

        async function doPasskeyLogin() {
            const errEl = document.getElementById('login-error');
            errEl.style.display = 'none';

            if (!window.PublicKeyCredential) {
                errEl.textContent = 'Passkeys not supported in this browser';
                errEl.style.display = 'block';
                return;
            }

            try {
                // Get authentication options from server
                const optResp = await fetch('/auth/passkey/login/begin', {method: 'POST'});
                const options = await optResp.json();

                // Decode challenge
                options.challenge = base64urlToBuffer(options.challenge);
                if (options.allowCredentials) {
                    options.allowCredentials = options.allowCredentials.map(c => ({
                        ...c,
                        id: base64urlToBuffer(c.id)
                    }));
                }

                // Get credential from authenticator
                const credential = await navigator.credentials.get({publicKey: options});

                // Send to server for verification
                const credData = {
                    id: credential.id,
                    rawId: bufferToBase64url(credential.rawId),
                    type: credential.type,
                    response: {
                        clientDataJSON: bufferToBase64url(credential.response.clientDataJSON),
                        authenticatorData: bufferToBase64url(credential.response.authenticatorData),
                        signature: bufferToBase64url(credential.response.signature),
                        userHandle: credential.response.userHandle ? bufferToBase64url(credential.response.userHandle) : null
                    }
                };

                const verifyResp = await fetch('/auth/passkey/login/complete', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({credential: credData})
                });

                if (!verifyResp.ok) {
                    const err = await verifyResp.json();
                    errEl.textContent = err.detail || 'Passkey login failed';
                    errEl.style.display = 'block';
                    return;
                }

                const data = await verifyResp.json();
                completeLogin(data);
            } catch(e) {
                if (e.name === 'NotAllowedError') {
                    errEl.textContent = 'Passkey request cancelled';
                } else {
                    errEl.textContent = 'Passkey login failed: ' + e.message;
                }
                errEl.style.display = 'block';
            }
        }

        function completeLogin(data) {
            authToken = data.token;
            localStorage.setItem('alfred_token', authToken);
            document.getElementById('login-overlay').classList.add('hidden');
            loadIntegrations();
            initConversations();
            // Auto-enable Hey Alfred wake word on startup
            setTimeout(() => {
                toggleWakeWord();
                // Force menu closed after everything initializes
                document.getElementById('history-panel').classList.remove('open');
                document.getElementById('history-overlay').classList.remove('open');
            }, 500);
            // Extra delayed close in case async operations open it
            setTimeout(() => {
                document.getElementById('history-panel').classList.remove('open');
                document.getElementById('history-overlay').classList.remove('open');
            }, 1200);
        }

        // WebAuthn helper functions
        function base64urlToBuffer(base64url) {
            const base64 = base64url.replace(/-/g, '+').replace(/_/g, '/');
            const pad = base64.length % 4;
            const padded = pad ? base64 + '='.repeat(4 - pad) : base64;
            const binary = atob(padded);
            const bytes = new Uint8Array(binary.length);
            for (let i = 0; i < binary.length; i++) {
                bytes[i] = binary.charCodeAt(i);
            }
            return bytes.buffer;
        }

        function bufferToBase64url(buffer) {
            const bytes = new Uint8Array(buffer);
            let binary = '';
            for (let i = 0; i < bytes.length; i++) {
                binary += String.fromCharCode(bytes[i]);
            }
            return btoa(binary).replace(/\+/g, '-').replace(/\//g, '_').replace(/=/g, '');
        }

        async function doLogout() {
            await fetch('/auth/logout', {method: 'POST'});
            localStorage.removeItem('alfred_token');
            authToken = '';
            location.reload();
        }

        async function changePassword() {
            const newPass = document.getElementById('new-pass').value;
            const msgEl = document.getElementById('pass-msg');
            if (!newPass || newPass.length < 6) {
                msgEl.textContent = 'Password must be at least 6 characters';
                msgEl.style.color = '#f87171';
                msgEl.style.display = 'block';
                return;
            }
            try {
                const resp = await fetch('/auth/change-password', {
                    method: 'POST',
                    headers: authHeaders({'Content-Type': 'application/json'}),
                    body: JSON.stringify({username: 'bruce', password: newPass})
                });
                if (resp.ok) {
                    msgEl.textContent = 'Password changed';
                    msgEl.style.color = '#4ade80';
                    document.getElementById('new-pass').value = '';
                } else {
                    msgEl.textContent = 'Failed to change password';
                    msgEl.style.color = '#f87171';
                }
            } catch(e) {
                msgEl.textContent = 'Connection error';
                msgEl.style.color = '#f87171';
            }
            msgEl.style.display = 'block';
        }

        function toggleSettings() {
            document.getElementById('settings-panel').classList.toggle('open');
            document.getElementById('settings-overlay').classList.toggle('open');
            if (document.getElementById('settings-panel').classList.contains('open')) {
                loadSecuritySettings();
            }
        }

        // ==================== 2FA Management ====================
        let totpSecret = null;

        async function loadSecuritySettings() {
            try {
                // Load 2FA status
                const resp = await fetch('/auth/2fa/status', {headers: authHeaders()});
                const data = await resp.json();
                const statusEl = document.getElementById('2fa-status');
                const btnEl = document.getElementById('2fa-toggle-btn');
                const setupEl = document.getElementById('2fa-setup-area');

                if (data.enabled) {
                    statusEl.textContent = 'Enabled';
                    statusEl.style.background = '#166534';
                    statusEl.style.color = '#4ade80';
                    btnEl.textContent = 'Disable 2FA';
                    btnEl.onclick = disable2FA;
                    setupEl.innerHTML = '';
                } else {
                    statusEl.textContent = 'Disabled';
                    statusEl.style.background = '#7f1d1d';
                    statusEl.style.color = '#f87171';
                    btnEl.textContent = 'Enable 2FA';
                    btnEl.onclick = setup2FA;
                    setupEl.innerHTML = '';
                }

                // Load passkeys
                await loadPasskeys();
            } catch(e) {
                console.error('Failed to load security settings:', e);
            }
        }

        async function setup2FA() {
            try {
                const resp = await fetch('/auth/2fa/setup', {
                    method: 'POST',
                    headers: authHeaders()
                });
                const data = await resp.json();
                if (data.error) {
                    alert('Failed to setup 2FA: ' + data.error);
                    return;
                }
                totpSecret = data.secret;
                const setupEl = document.getElementById('2fa-setup-area');
                setupEl.innerHTML = `
                    <div style="text-align:center;margin-bottom:12px">
                        <img src="${data.qr_code}" alt="QR Code" style="max-width:180px;border-radius:8px;background:#fff;padding:8px">
                    </div>
                    <p style="font-size:11px;color:#888;text-align:center;margin-bottom:8px">Scan with your authenticator app</p>
                    <p style="font-size:10px;color:#666;text-align:center;margin-bottom:12px;word-break:break-all">Secret: ${data.secret}</p>
                    <input type="text" id="2fa-verify-code" placeholder="Enter 6-digit code" maxlength="6"
                        style="width:100%;padding:10px;border-radius:6px;border:1px solid #333;background:#0a0a0a;color:#e0e0e0;font-size:16px;text-align:center;letter-spacing:4px;margin-bottom:8px"
                        onkeydown="if(event.key==='Enter')verify2FASetup()">
                `;
                document.getElementById('2fa-toggle-btn').textContent = 'Verify & Enable';
                document.getElementById('2fa-toggle-btn').onclick = verify2FASetup;
            } catch(e) {
                alert('Error setting up 2FA: ' + e.message);
            }
        }

        async function verify2FASetup() {
            const code = document.getElementById('2fa-verify-code').value.trim();
            if (code.length !== 6) {
                alert('Please enter a 6-digit code');
                return;
            }
            try {
                const resp = await fetch('/auth/2fa/enable', {
                    method: 'POST',
                    headers: authHeaders({'Content-Type': 'application/json'}),
                    body: JSON.stringify({code: code})
                });
                const data = await resp.json();
                if (data.error) {
                    alert('Invalid code. Please try again.');
                    return;
                }
                alert('2FA enabled successfully!');
                loadSecuritySettings();
            } catch(e) {
                alert('Error enabling 2FA: ' + e.message);
            }
        }

        async function disable2FA() {
            const code = prompt('Enter your 2FA code to disable:');
            if (!code) return;
            try {
                const resp = await fetch('/auth/2fa/disable', {
                    method: 'POST',
                    headers: authHeaders({'Content-Type': 'application/json'}),
                    body: JSON.stringify({code: code})
                });
                const data = await resp.json();
                if (data.error) {
                    alert('Invalid code. 2FA was not disabled.');
                    return;
                }
                alert('2FA disabled.');
                loadSecuritySettings();
            } catch(e) {
                alert('Error disabling 2FA: ' + e.message);
            }
        }

        function toggle2FA() {
            // Placeholder - replaced by setup2FA or disable2FA
        }

        // ==================== Passkey Management ====================
        async function loadPasskeys() {
            try {
                const resp = await fetch('/auth/passkey/list', {headers: authHeaders()});
                const data = await resp.json();
                const listEl = document.getElementById('passkeys-list');
                const countEl = document.getElementById('passkey-count');

                countEl.textContent = data.passkeys.length + ' registered';

                if (data.passkeys.length === 0) {
                    listEl.innerHTML = '<p style="font-size:12px;color:#666;text-align:center;padding:8px">No passkeys registered</p>';
                } else {
                    listEl.innerHTML = data.passkeys.map(p => `
                        <div style="display:flex;justify-content:space-between;align-items:center;padding:8px;background:#0a0a0a;border-radius:6px;margin-bottom:4px">
                            <div>
                                <span style="font-size:13px;color:#e0e0e0">${p.name}</span>
                                <span style="font-size:11px;color:#666;margin-left:8px">${new Date(p.created).toLocaleDateString()}</span>
                            </div>
                            <button onclick="deletePasskey('${p.id}')" style="background:none;border:none;color:#f87171;cursor:pointer;font-size:16px" title="Delete">&times;</button>
                        </div>
                    `).join('');
                }
            } catch(e) {
                console.error('Failed to load passkeys:', e);
            }
        }

        async function addPasskey() {
            if (!window.PublicKeyCredential) {
                alert('Passkeys are not supported in this browser');
                return;
            }

            const name = prompt('Name this passkey (e.g., "MacBook Touch ID"):', 'My Passkey');
            if (!name) return;

            try {
                // Get registration options
                const optResp = await fetch('/auth/passkey/register/begin', {
                    method: 'POST',
                    headers: authHeaders()
                });
                const options = await optResp.json();
                if (options.error) {
                    alert('Failed to start passkey registration: ' + options.error);
                    return;
                }

                // Decode challenge and user ID
                options.challenge = base64urlToBuffer(options.challenge);
                options.user.id = base64urlToBuffer(options.user.id);

                // Create credential
                const credential = await navigator.credentials.create({publicKey: options});

                // Send to server
                const credData = {
                    id: credential.id,
                    rawId: bufferToBase64url(credential.rawId),
                    type: credential.type,
                    response: {
                        clientDataJSON: bufferToBase64url(credential.response.clientDataJSON),
                        attestationObject: bufferToBase64url(credential.response.attestationObject)
                    }
                };

                const verifyResp = await fetch('/auth/passkey/register/complete', {
                    method: 'POST',
                    headers: authHeaders({'Content-Type': 'application/json'}),
                    body: JSON.stringify({credential: credData, name: name})
                });
                const result = await verifyResp.json();
                if (result.error) {
                    alert('Failed to register passkey: ' + result.error);
                    return;
                }

                alert('Passkey registered successfully!');
                loadPasskeys();
            } catch(e) {
                if (e.name === 'NotAllowedError') {
                    alert('Passkey registration was cancelled');
                } else {
                    alert('Error registering passkey: ' + e.message);
                }
            }
        }

        async function deletePasskey(id) {
            if (!confirm('Delete this passkey?')) return;
            try {
                const resp = await fetch('/auth/passkey/' + encodeURIComponent(id), {
                    method: 'DELETE',
                    headers: authHeaders()
                });
                const data = await resp.json();
                if (data.error) {
                    alert('Failed to delete passkey: ' + data.error);
                    return;
                }
                loadPasskeys();
            } catch(e) {
                alert('Error deleting passkey: ' + e.message);
            }
        }

        async function loadIntegrations() {
            try {
                const resp = await fetch('/integrations/status', {headers: authHeaders()});
                const data = await resp.json();

                // Integration list
                let html = '';
                // Google
                const gConnected = data.google?.connected;
                html += `<div class="status-item">
                    <span><span class="status-dot ${gConnected?'green':'red'}"></span>Google (Gmail + Calendar)</span>
                    ${gConnected ? '<span style="color:#4ade80;font-size:12px">Connected</span>' :
                    '<a href="/auth/google" class="connect-btn">Connect</a>'}
                </div>`;
                // Servers
                const sCount = data.servers?.count || 0;
                html += `<div class="status-item">
                    <span><span class="status-dot ${sCount>0?'green':'yellow'}"></span>Servers</span>
                    <span style="color:#888;font-size:12px">${sCount} registered</span>
                </div>`;
                // CRM
                const crmConnected = data.base_crm?.connected;
                const crmLabel = data.base_crm?.label || 'CRM';
                html += `<div class="status-item">
                    <span><span class="status-dot ${crmConnected?'green':'red'}"></span>${crmLabel}</span>
                    <span style="color:#888;font-size:12px">${crmConnected?'Connected':'Not connected'}</span>
                </div>`;
                // n8n
                const n8nConnected = data.n8n?.connected;
                const n8nCount = data.n8n?.workflow_count || 0;
                html += `<div class="status-item">
                    <span><span class="status-dot ${n8nConnected?'green':'red'}"></span>n8n Automations</span>
                    <span style="color:#888;font-size:12px">${n8nConnected ? n8nCount + ' workflows' : 'Not connected'}</span>
                </div>`;
                // Nextcloud
                const ncConnected = data.nextcloud?.connected;
                html += `<div class="status-item">
                    <span><span class="status-dot ${ncConnected?'green':'red'}"></span>Nextcloud</span>
                    <span style="color:#888;font-size:12px">${ncConnected?'Connected':'Not connected'}</span>
                </div>`;
                // Stripe
                const stripeConnected = data.stripe?.connected;
                html += `<div class="status-item">
                    <span><span class="status-dot ${stripeConnected?'green':'red'}"></span>Stripe Payments</span>
                    <span style="color:#888;font-size:12px">${stripeConnected?'Connected':'Not connected'}</span>
                </div>`;

                document.getElementById('integration-list').innerHTML = html;

                // LLM info
                let llmHtml = `<div class="status-item">
                    <span><span class="status-dot green"></span>Local: ${data.ollama?.model || 'N/A'}</span>
                </div>`;
                // Check Claude Code CLI first (Max subscription), then API key
                const claudeCodeConfigured = data.claude_code?.configured;
                const apiConfigured = data.anthropic?.configured;
                if (claudeCodeConfigured) {
                    llmHtml += `<div class="status-item">
                        <span><span class="status-dot green"></span>Cloud: Claude Code</span>
                        <span style="color:#888;font-size:12px">Max subscription</span>
                    </div>`;
                } else if (apiConfigured) {
                    llmHtml += `<div class="status-item">
                        <span><span class="status-dot green"></span>Cloud: ${data.anthropic?.model || 'N/A'}</span>
                        <span style="color:#888;font-size:12px">API key</span>
                    </div>`;
                } else {
                    llmHtml += `<div class="status-item">
                        <span><span class="status-dot red"></span>Cloud: Not configured</span>
                        <span style="color:#888;font-size:12px">No CLI or API key</span>
                    </div>`;
                }
                document.getElementById('llm-info').innerHTML = llmHtml;

                // Account
                const me = await (await fetch('/auth/me', {headers: authHeaders()})).json();
                document.getElementById('account-info').innerHTML =
                    `<div class="status-item"><span>User: ${me.username || 'N/A'}</span><span style="color:#888;font-size:12px">${me.role || ''}</span></div>`;

                // TTS settings
                const tts = data.tts || {};
                currentTtsBackend = tts.backend || 'kokoro';
                document.querySelectorAll('.tts-option').forEach(b => b.classList.remove('active'));
                const activeBtn = document.getElementById('tts-' + currentTtsBackend);
                if (activeBtn) activeBtn.classList.add('active');
                const ttsStatus = document.getElementById('tts-status');
                if (ttsStatus) {
                    if (tts.backend === 'qwen3' && !tts.qwen3_available) {
                        ttsStatus.innerHTML = '<span style="color:#f87171">⚠ Qwen3 server not available</span>';
                    } else if (tts.backend === 'qwen3') {
                        ttsStatus.innerHTML = '<span style="color:#4ade80">✓ Qwen3 server running</span>';
                    } else {
                        ttsStatus.textContent = '';
                    }
                }
                // Load available voices
                await loadVoices();
            } catch(e) {
                document.getElementById('integration-list').innerHTML = '<span style="color:#f87171">Failed to load</span>';
            }
        }

        async function setTTS(backend) {
            try {
                const resp = await fetch('/settings/tts?backend=' + backend, {
                    method: 'PUT',
                    headers: authHeaders()
                });
                if (resp.ok) {
                    currentTtsBackend = backend;
                    document.querySelectorAll('.tts-option').forEach(b => b.classList.remove('active'));
                    document.getElementById('tts-' + backend)?.classList.add('active');
                    const ttsStatus = document.getElementById('tts-status');
                    if (backend === 'qwen3') {
                        ttsStatus.innerHTML = '<span style="color:#888">Checking Qwen3 server...</span>';
                        const check = await fetch('/settings/tts', {headers: authHeaders()});
                        const data = await check.json();
                        if (data.qwen3_available) {
                            ttsStatus.innerHTML = '<span style="color:#4ade80">✓ Qwen3 server running</span>';
                        } else {
                            ttsStatus.innerHTML = '<span style="color:#f87171">⚠ Qwen3 server not available</span>';
                        }
                    } else {
                        ttsStatus.textContent = '';
                    }
                    // Reload voices for new backend
                    await loadVoices();
                }
            } catch(e) {
                console.error('Failed to set TTS:', e);
            }
        }

        async function loadVoices() {
            try {
                const resp = await fetch('/settings/voices', {headers: authHeaders()});
                const data = await resp.json();
                const select = document.getElementById('tts-voice');
                if (select && data.voices) {
                    select.innerHTML = data.voices.map(v =>
                        `<option value="${v.id}" ${v.id === data.current ? 'selected' : ''}>${v.name} (${v.desc})</option>`
                    ).join('');
                }
            } catch(e) {
                console.error('Failed to load voices:', e);
            }
        }

        async function setVoice(voice) {
            try {
                await fetch('/settings/voice?voice=' + encodeURIComponent(voice), {
                    method: 'PUT',
                    headers: authHeaders()
                });
            } catch(e) {
                console.error('Failed to set voice:', e);
            }
        }

        let autoSpeak = localStorage.getItem('alfred_auto_speak') === 'true';
        let currentAudio = null;
        let currentTtsBackend = 'kokoro';  // Track TTS backend for acknowledgment logic
        let msgTexts = {};
        let msgCounter = 0;
        if (autoSpeak) document.getElementById('auto-speak-btn')?.classList.add('active');

        function addMsg(text, role, tier, noAutoSpeak = false) {
            // Show chat state when adding messages
            if (isWelcomeState) showChatState();

            const div = document.createElement('div');
            div.className = `msg ${role}`;
            const mid = ++msgCounter;
            msgTexts[mid] = text;

            if (role === 'alfred') {
                // Alfred message: no bubble, just text + action buttons
                div.innerHTML = `
                    <div class="content">${renderText(text)}</div>
                    <div class="msg-actions">
                        <button class="msg-action-btn copy-btn" data-mid="${mid}" title="Copy">
                            <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><rect x="9" y="9" width="13" height="13" rx="2"/><path d="M5 15H4a2 2 0 01-2-2V4a2 2 0 012-2h9a2 2 0 012 2v1"/></svg>
                        </button>
                        <button class="msg-action-btn thumbs-up-btn" data-mid="${mid}" title="Good response">
                            <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M14 9V5a3 3 0 00-3-3l-4 9v11h11.28a2 2 0 002-1.7l1.38-9a2 2 0 00-2-2.3zM7 22H4a2 2 0 01-2-2v-7a2 2 0 012-2h3"/></svg>
                        </button>
                        <button class="msg-action-btn thumbs-down-btn" data-mid="${mid}" title="Bad response">
                            <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M10 15v4a3 3 0 003 3l4-9V2H5.72a2 2 0 00-2 1.7l-1.38 9a2 2 0 002 2.3zm7-13h2.67A2.31 2.31 0 0122 4v7a2.31 2.31 0 01-2.33 2H17"/></svg>
                        </button>
                        <button class="msg-action-btn speak-btn" data-mid="${mid}" title="Read aloud">
                            <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><polygon points="11 5 6 9 2 9 2 15 6 15 11 19 11 5"/><path d="M15.54 8.46a5 5 0 010 7.07"/></svg>
                        </button>
                        <button class="msg-action-btn regenerate-btn" data-mid="${mid}" title="Regenerate">
                            <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M23 4v6h-6M1 20v-6h6"/><path d="M3.51 9a9 9 0 0114.85-3.36L23 10M1 14l4.64 4.36A9 9 0 0020.49 15"/></svg>
                        </button>
                    </div>
                `;
            } else {
                // User message: bubble style
                div.innerHTML = `<div class="content">${renderText(text)}</div>`;
            }

            chat.appendChild(div);
            chat.scrollTop = chat.scrollHeight;
            if (role === 'alfred' && autoSpeak && !loadingHistory && !noAutoSpeak) {
                const btn = div.querySelector('.speak-btn');
                if (btn) speakText(btn);
            }
            return div;
        }

        function addMsgHtml(text, role, tier, noAutoSpeak = false, extraHtml = '') {
            // Show chat state when adding messages
            if (isWelcomeState) showChatState();

            const div = document.createElement('div');
            div.className = `msg ${role}`;
            const mid = ++msgCounter;
            msgTexts[mid] = text;

            if (role === 'alfred') {
                div.innerHTML = `
                    <div class="content">${renderText(text)}${extraHtml}</div>
                    <div class="msg-actions">
                        <button class="msg-action-btn copy-btn" data-mid="${mid}" title="Copy">
                            <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><rect x="9" y="9" width="13" height="13" rx="2"/><path d="M5 15H4a2 2 0 01-2-2V4a2 2 0 012-2h9a2 2 0 012 2v1"/></svg>
                        </button>
                        <button class="msg-action-btn thumbs-up-btn" data-mid="${mid}" title="Good response">
                            <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M14 9V5a3 3 0 00-3-3l-4 9v11h11.28a2 2 0 002-1.7l1.38-9a2 2 0 00-2-2.3zM7 22H4a2 2 0 01-2-2v-7a2 2 0 012-2h3"/></svg>
                        </button>
                        <button class="msg-action-btn thumbs-down-btn" data-mid="${mid}" title="Bad response">
                            <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M10 15v4a3 3 0 003 3l4-9V2H5.72a2 2 0 00-2 1.7l-1.38 9a2 2 0 002 2.3zm7-13h2.67A2.31 2.31 0 0122 4v7a2.31 2.31 0 01-2.33 2H17"/></svg>
                        </button>
                        <button class="msg-action-btn speak-btn" data-mid="${mid}" title="Read aloud">
                            <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><polygon points="11 5 6 9 2 9 2 15 6 15 11 19 11 5"/><path d="M15.54 8.46a5 5 0 010 7.07"/></svg>
                        </button>
                        <button class="msg-action-btn regenerate-btn" data-mid="${mid}" title="Regenerate">
                            <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M23 4v6h-6M1 20v-6h6"/><path d="M3.51 9a9 9 0 0114.85-3.36L23 10M1 14l4.64 4.36A9 9 0 0020.49 15"/></svg>
                        </button>
                    </div>
                `;
            } else {
                div.innerHTML = `<div class="content">${renderText(text)}${extraHtml}</div>`;
            }

            chat.appendChild(div);
            chat.scrollTop = chat.scrollHeight;
            if (role === 'alfred' && autoSpeak && !loadingHistory && !noAutoSpeak) {
                const btn = div.querySelector('.speak-btn');
                if (btn) speakText(btn);
            }
            return div;
        }

        document.addEventListener('click', function(e) {
            const speakBtn = e.target.closest('.speak-btn');
            if (speakBtn) { speakText(speakBtn); return; }
            const stopBtn = e.target.closest('.stop-btn');
            if (stopBtn) { stopAudio(); return; }
            const copyBtn = e.target.closest('.copy-btn');
            if (copyBtn) { copyMessage(copyBtn); return; }
            const thumbsUpBtn = e.target.closest('.thumbs-up-btn');
            if (thumbsUpBtn) { thumbsUpBtn.classList.toggle('active'); return; }
            const thumbsDownBtn = e.target.closest('.thumbs-down-btn');
            if (thumbsDownBtn) { thumbsDownBtn.classList.toggle('active'); return; }
        });

        function copyMessage(btn) {
            const mid = btn.getAttribute('data-mid');
            const text = msgTexts[mid];
            if (!text) return;
            navigator.clipboard.writeText(text).then(() => {
                btn.classList.add('active');
                setTimeout(() => btn.classList.remove('active'), 2000);
            });
        }

        function stopAudio() {
            if (currentAudio) { currentAudio.pause(); currentAudio = null; }
            document.querySelectorAll('.speak-btn.speaking').forEach(b => b.classList.remove('speaking'));
            document.querySelectorAll('.stop-btn.visible').forEach(b => b.classList.remove('visible'));
        }

        function toggleAutoSpeak() {
            autoSpeak = !autoSpeak;
            localStorage.setItem('alfred_auto_speak', autoSpeak);
            document.getElementById('auto-speak-btn').classList.toggle('active', autoSpeak);
        }

        function executeUiAction(uiAction) {
            // Execute a UI action from the API response
            if (!uiAction || !uiAction.action) return;
            console.log('Executing UI action:', uiAction);

            if (uiAction.action === 'set_auto_speak') {
                autoSpeak = uiAction.value;
                localStorage.setItem('alfred_auto_speak', String(autoSpeak));
                document.getElementById('auto-speak-btn').classList.toggle('active', autoSpeak);
                console.log('UI Action: auto-speak set to', autoSpeak);
            }
            if (uiAction.action === 'set_hands_free') {
                if (uiAction.value) {
                    if (!handsFreeActive) toggleHandsFree();
                } else {
                    if (handsFreeActive) {
                        handsFreeActive = false;
                        document.getElementById('handsfree-btn').classList.remove('active');
                        if (vadInstance) vadInstance.pause();
                        setVadState('idle');
                    }
                }
                console.log('UI Action: hands-free set to', uiAction.value);
            }
        }

        async function speakText(btn) {
            const mid = btn.getAttribute('data-mid');
            const text = msgTexts[mid];
            if (!text) return;
            const stopBtn = btn.parentElement.querySelector('.stop-btn');
            // Toggle off if this button is already speaking
            if (btn.classList.contains('speaking')) {
                stopAudio();
                return;
            }
            // Stop any other audio first
            stopAudio();
            btn.classList.add('speaking');
            if (stopBtn) stopBtn.classList.add('visible');
            try {
                const resp = await fetch('/voice/speak', {
                    method: 'POST',
                    headers: authHeaders({'Content-Type': 'application/json'}),
                    body: JSON.stringify({message: text})
                });
                if (!resp.ok) throw new Error('TTS failed');
                const blob = await resp.blob();
                const url = URL.createObjectURL(blob);
                currentAudio = new Audio(url);
                currentAudio.onended = () => {
                    btn.classList.remove('speaking');
                    if (stopBtn) stopBtn.classList.remove('visible');
                    currentAudio = null;
                    URL.revokeObjectURL(url);
                };
                currentAudio.onerror = () => {
                    btn.classList.remove('speaking');
                    if (stopBtn) stopBtn.classList.remove('visible');
                    currentAudio = null;
                };
                currentAudio.play();
            } catch(e) {
                console.error('TTS error:', e);
                btn.classList.remove('speaking');
                if (stopBtn) stopBtn.classList.remove('visible');
            }
        }

        // ==================== File Upload ====================

        let pendingFile = null;  // {type: 'image'|'document', path, base64?, media_type?, filename, text_preview?}

        function handleFileSelect(e) {
            const file = e.target.files[0];
            if (!file) return;
            const ext = file.name.split('.').pop().toLowerCase();
            const imageExts = ['jpg','jpeg','png','gif','webp'];
            const docExts = ['pdf','doc','docx','xls','xlsx','csv','txt','md','json'];

            if (imageExts.includes(ext)) {
                uploadImage(file);
            } else if (docExts.includes(ext)) {
                uploadDocument(file);
            } else {
                alert('Unsupported file type');
            }
            e.target.value = '';
        }

        async function uploadImage(file) {
            const form = new FormData();
            form.append('file', file);
            try {
                const resp = await fetch('/upload/image', {method:'POST', headers:authHeaders(), body:form});
                if (!resp.ok) throw new Error(await resp.text());
                const data = await resp.json();
                pendingFile = {type:'image', path:data.path, base64:data.base64, media_type:data.media_type, filename:data.filename};
                showPendingFile(data.filename, `Image (${Math.round(data.size_bytes/1024)}KB)`, `data:${data.media_type};base64,${data.base64}`);
            } catch(e) { alert('Upload failed: ' + e.message); }
        }

        async function uploadDocument(file) {
            const form = new FormData();
            form.append('file', file);
            try {
                const resp = await fetch('/upload/document', {method:'POST', headers:authHeaders(), body:form});
                if (!resp.ok) throw new Error(await resp.text());
                const data = await resp.json();
                pendingFile = {type:'document', path:data.path, filename:data.filename, text_preview:data.text_preview, metadata:data.metadata};
                const meta = data.metadata;
                let info = meta.extension.toUpperCase();
                if (meta.pages) info += `, ${meta.pages} pages`;
                if (meta.rows) info += `, ${meta.rows} rows`;
                showPendingFile(data.filename, info, null);
            } catch(e) { alert('Upload failed: ' + e.message); }
        }

        function showPendingFile(name, meta, previewUrl) {
            const container = document.getElementById('pending-file');
            const preview = document.getElementById('pending-preview');
            document.getElementById('pending-name').textContent = name;
            document.getElementById('pending-meta').textContent = meta;
            if (previewUrl) {
                preview.src = previewUrl;
                preview.style.display = 'block';
            } else {
                preview.style.display = 'none';
            }
            container.classList.add('visible');
            // Mark both upload buttons as having a file
            document.querySelectorAll('.upload-btn').forEach(btn => btn.classList.add('has-file'));
            // Enable send buttons when file is pending
            if (typeof updateWelcomeSendBtn === 'function') updateWelcomeSendBtn();
        }

        function clearPendingFile() {
            pendingFile = null;
            document.getElementById('pending-file').classList.remove('visible');
            document.querySelectorAll('.upload-btn').forEach(btn => btn.classList.remove('has-file'));
            // Update send button state
            if (typeof updateWelcomeSendBtn === 'function') updateWelcomeSendBtn();
        }

        // ==================== Image Paste Handler ====================

        document.addEventListener('paste', async function(e) {
            // Check if we have image data in clipboard
            const items = e.clipboardData?.items;
            if (!items) return;

            for (const item of items) {
                if (item.type.startsWith('image/')) {
                    e.preventDefault();
                    const blob = item.getAsFile();
                    if (blob) {
                        await handlePastedImage(blob, item.type);
                    }
                    return;
                }
            }
        });

        async function handlePastedImage(blob, mimeType) {
            // Convert blob to base64
            const reader = new FileReader();
            reader.onload = async function() {
                const base64 = reader.result.split(',')[1];
                const ext = mimeType.split('/')[1] || 'png';
                const filename = `pasted-image-${Date.now()}.${ext}`;
                const sizeKB = Math.round(blob.size / 1024);

                // Store as pending file
                pendingFile = {
                    type: 'image',
                    path: null,  // No server path yet
                    base64: base64,
                    media_type: mimeType,
                    filename: filename
                };

                // Show preview
                showPendingFile(filename, `Pasted image (${sizeKB}KB)`, `data:${mimeType};base64,${base64}`);

                // Focus input so user can add a message
                if (isWelcomeState) {
                    document.getElementById('welcome-input').focus();
                } else {
                    document.getElementById('input').focus();
                }
            };
            reader.readAsDataURL(blob);
        }

        function escapeHtml(t) {
            const d = document.createElement('div');
            d.textContent = t;
            return d.innerHTML;
        }

        function renderText(t) {
            let s = escapeHtml(t);
            s = s.replace(/`([^`]+)`/g, '<code>$1</code>');
            s = s.replace(/\*\*([^*]+)\*\*/g, '<strong>$1</strong>');
            // Handle download links: [Download: filename](/download/filename)
            s = s.replace(/\[Download: ([^\]]+)\]\(([^)]+)\)/g, '<a class="download-btn" href="$2" download>📥 $1</a>');
            return s;
        }

        async function send() {
            const msg = input.value.trim();
            if (!msg && !pendingFile) return;
            input.value = '';
            input.style.height = 'auto';

            // Build user message with file context
            let userMsgDisplay = msg;
            let payload = {message: msg, session_id: currentConversationId || 'default'};

            if (pendingFile) {
                if (pendingFile.type === 'image') {
                    userMsgDisplay = msg ? `[Image: ${pendingFile.filename}] ${msg}` : `[Image: ${pendingFile.filename}]`;
                    payload.image_path = pendingFile.path;
                    payload.image_base64 = pendingFile.base64;
                    payload.image_media_type = pendingFile.media_type;
                    // Show image in chat
                    const imgHtml = `<img class="inline-image" src="data:${pendingFile.media_type};base64,${pendingFile.base64}" alt="${pendingFile.filename}">`;
                    addMsgHtml(userMsgDisplay, 'user', null, true, imgHtml);
                } else {
                    userMsgDisplay = msg ? `[Document: ${pendingFile.filename}] ${msg}` : `Analyze this document: ${pendingFile.filename}`;
                    payload.document_path = pendingFile.path;
                    payload.message = userMsgDisplay;
                    addMsg(userMsgDisplay, 'user', null, true);
                }
                clearPendingFile();
            } else {
                addMsg(msg, 'user');
            }

            document.getElementById('send-btn').disabled = true;

            // Show thinking indicator (move to end of chat, after user's message)
            const thinking = document.getElementById('thinking');
            chat.appendChild(thinking);  // Move to end
            thinking.classList.add('visible');
            chat.scrollTop = chat.scrollHeight;

            // Ensure a conversation exists - create on first message only
            if (!currentConversationId) {
                try {
                    const cr = await fetch('/conversations', {method: 'POST', headers: authHeaders()});
                    const cd = await cr.json();
                    currentConversationId = cd.id;
                    // If in project view, assign to project
                    if (currentProjectId) {
                        await fetch(`/conversations/${cd.id}/project`, {
                            method: 'PUT',
                            headers: authHeaders({'Content-Type': 'application/json'}),
                            body: JSON.stringify({project_id: currentProjectId})
                        });
                    }
                } catch(e) {}
            }

            try {
                payload.session_id = currentConversationId || 'default';
                const resp = await fetch('/chat', {
                    method: 'POST',
                    headers: authHeaders({'Content-Type': 'application/json'}),
                    body: JSON.stringify(payload)
                });
                const data = await resp.json();
                // Check for UI actions in response
                if (data.ui_action) {
                    executeUiAction(data.ui_action);
                }
                const displayResponse = data.response;
                // Check for generated images
                if (data.images && data.images.length > 0) {
                    let imagesHtml = '';
                    data.images.forEach(img => {
                        imagesHtml += `<img class="inline-image" src="data:image/png;base64,${img.base64}" alt="${img.filename}">`;
                        if (img.download_url) {
                            imagesHtml += `<a class="download-btn" href="${img.download_url}" download>📥 ${img.filename}</a>`;
                        }
                    });
                    addMsgHtml(displayResponse, 'alfred', data.tier, false, imagesHtml);
                } else {
                    addMsg(displayResponse, 'alfred', data.tier);
                }
                // Refresh conversation list to update title/preview
                loadConversations();
            } catch(e) {
                addMsg('Connection error. Please try again.', 'alfred');
            }
            // Hide thinking indicator
            thinking.classList.remove('visible');
            document.getElementById('send-btn').disabled = false;
            input.focus();
        }

        async function toggleMic() {
            // Cut off any audio Alfred is currently speaking
            if (currentAudio) {
                currentAudio.pause();
                currentAudio = null;
                const speaking = document.querySelector('.speak-btn.speaking');
                if (speaking) speaking.classList.remove('speaking');
            }
            const mainBtn = document.getElementById('mic-btn');
            const welcomeBtn = document.getElementById('welcome-mic-btn');
            const mainImg = mainBtn?.querySelector('img');
            const welcomeImg = welcomeBtn?.querySelector('img');
            if (isRecording) {
                mediaRecorder.stop();
                mainBtn?.classList.remove('recording');
                welcomeBtn?.classList.remove('recording');
                if (mainImg) mainImg.src = '/static/logo-white.png';
                if (welcomeImg) welcomeImg.src = '/static/logo-white.png';
                isRecording = false;
            } else {
                const stream = await navigator.mediaDevices.getUserMedia({audio: true});
                mediaRecorder = new MediaRecorder(stream);
                const chunks = [];
                mediaRecorder.ondataavailable = e => chunks.push(e.data);
                mediaRecorder.onstop = async () => {
                    stream.getTracks().forEach(t => t.stop());
                    const blob = new Blob(chunks, {type: 'audio/webm'});
                    const form = new FormData();
                    form.append('audio', blob, 'recording.webm');
                    try {
                        const resp = await fetch('/voice/transcribe', {method:'POST', headers: authHeaders(), body: form});
                        const data = await resp.json();
                        if (data.text) {
                            // Put transcription in the appropriate input
                            if (isWelcomeState) {
                                document.getElementById('welcome-input').value = data.text;
                                sendFromWelcome();
                            } else {
                                input.value = data.text;
                                send();
                            }
                        }
                    } catch(e) { console.error('Transcription failed:', e); }
                };
                mediaRecorder.start();
                mainBtn?.classList.add('recording');
                welcomeBtn?.classList.add('recording');
                if (mainImg) mainImg.src = '/static/glabs.jpeg';
                if (welcomeImg) welcomeImg.src = '/static/glabs.jpeg';
                isRecording = true;
            }
        }

        input.addEventListener('input', function() {
            this.style.height = 'auto';
            this.style.height = Math.min(this.scrollHeight, 120) + 'px';
        });

        // ==================== Screen Wake Lock ====================

        let wakeLock = null;

        async function acquireWakeLock() {
            if (wakeLock) return; // Already have it
            try {
                if ('wakeLock' in navigator) {
                    wakeLock = await navigator.wakeLock.request('screen');
                    console.log('Screen wake lock acquired');
                    wakeLock.addEventListener('release', () => {
                        console.log('Screen wake lock released');
                        wakeLock = null;
                    });
                }
            } catch (err) {
                console.log('Wake lock request failed:', err.message);
            }
        }

        function releaseWakeLock() {
            // Only release if neither wake word nor hands-free is active
            if (wakeWordActive || handsFreeActive) return;
            if (wakeLock) {
                wakeLock.release();
                wakeLock = null;
                console.log('Screen wake lock released manually');
            }
        }

        // Re-acquire wake lock when page becomes visible again
        document.addEventListener('visibilitychange', async () => {
            if (document.visibilityState === 'visible' && (wakeWordActive || handsFreeActive)) {
                await acquireWakeLock();
            }
        });

        // ==================== Wake Word Detection ====================

        let wakeWordActive = false;
        let wakeWordWs = null;
        let wakeWordStream = null;
        let wakeWordContext = null;
        let wakeWordProcessor = null;

        async function toggleWakeWord() {
            const btn = document.getElementById('wakeword-btn');
            const status = document.getElementById('wakeword-status');

            if (wakeWordActive) {
                // Stop wake word detection
                wakeWordActive = false;
                btn.classList.remove('active');
                status.classList.remove('visible');
                if (wakeWordWs) {
                    wakeWordWs.close();
                    wakeWordWs = null;
                }
                if (wakeWordStream) {
                    wakeWordStream.getTracks().forEach(t => t.stop());
                    wakeWordStream = null;
                }
                if (wakeWordContext) {
                    wakeWordContext.close();
                    wakeWordContext = null;
                }
                releaseWakeLock(); // Release if hands-free also inactive
                return;
            }

            // Start wake word detection
            try {
                btn.classList.add('active');
                status.classList.add('visible');
                status.textContent = 'Starting...';
                await acquireWakeLock(); // Keep screen on

                // Get microphone access
                wakeWordStream = await navigator.mediaDevices.getUserMedia({
                    audio: {
                        sampleRate: 16000,
                        channelCount: 1,
                        echoCancellation: true,
                        noiseSuppression: true
                    }
                });

                // Create audio context for resampling
                wakeWordContext = new (window.AudioContext || window.webkitAudioContext)({
                    sampleRate: 16000
                });

                const source = wakeWordContext.createMediaStreamSource(wakeWordStream);
                wakeWordProcessor = wakeWordContext.createScriptProcessor(1024, 1, 1);

                // Connect to WebSocket
                const wsProtocol = window.location.protocol === 'https:' ? 'wss:' : 'ws:';
                const token = localStorage.getItem('alfred_token');
                wakeWordWs = new WebSocket(`${wsProtocol}//${window.location.host}/ws/wakeword?token=${token}`);

                wakeWordWs.onopen = () => {
                    console.log('Wake word WebSocket connected');
                    status.textContent = 'Listening for "Hey Alfred"...';
                    wakeWordActive = true;
                    // Keep-alive ping every 30 seconds to prevent timeout
                    wakeWordWs._pingInterval = setInterval(() => {
                        if (wakeWordWs && wakeWordWs.readyState === WebSocket.OPEN) {
                            try { wakeWordWs.send(new ArrayBuffer(0)); } catch(e) {}
                        }
                    }, 30000);
                };

                wakeWordWs.onmessage = (event) => {
                    const data = JSON.parse(event.data);
                    if (data.type === 'detected') {
                        console.log('Wake word detected!', data.score);
                        status.textContent = 'Wake word detected!';

                        // Visual feedback
                        btn.style.animation = 'pulse 0.3s ease-in-out 3';
                        setTimeout(() => btn.style.animation = '', 1000);

                        // Trigger hands-free mode if not already active
                        if (!handsFreeActive) {
                            toggleHandsFree();
                        }

                        // Reset status after a moment
                        setTimeout(() => {
                            if (wakeWordActive) {
                                status.textContent = 'Listening for "Hey Alfred"...';
                            }
                        }, 2000);
                    } else if (data.type === 'error') {
                        console.error('Wake word error:', data.message);
                        status.textContent = 'Error: ' + data.message;
                    }
                };

                wakeWordWs.onclose = () => {
                    console.log('Wake word WebSocket closed');
                    if (wakeWordWs && wakeWordWs._pingInterval) {
                        clearInterval(wakeWordWs._pingInterval);
                    }
                    wakeWordWs = null;  // Clear reference immediately
                    if (wakeWordActive) {
                        // Reconnect after a short delay
                        status.textContent = 'Reconnecting...';
                        setTimeout(() => {
                            if (wakeWordActive) {
                                // Stop and restart wake word
                                wakeWordActive = false;
                                toggleWakeWord();  // This will restart it
                            }
                        }, 1000);
                    }
                };

                wakeWordWs.onerror = (err) => {
                    console.error('Wake word WebSocket error:', err);
                    status.textContent = 'Connection error';
                };

                // Process audio and send to WebSocket
                wakeWordProcessor.onaudioprocess = (e) => {
                    if (!wakeWordActive || !wakeWordWs || wakeWordWs.readyState !== WebSocket.OPEN) return;

                    const inputData = e.inputBuffer.getChannelData(0);

                    // Convert float32 to int16
                    const int16Data = new Int16Array(inputData.length);
                    for (let i = 0; i < inputData.length; i++) {
                        const s = Math.max(-1, Math.min(1, inputData[i]));
                        int16Data[i] = s < 0 ? s * 0x8000 : s * 0x7FFF;
                    }

                    // Send to WebSocket
                    wakeWordWs.send(int16Data.buffer);
                };

                source.connect(wakeWordProcessor);
                wakeWordProcessor.connect(wakeWordContext.destination);

            } catch (err) {
                console.error('Wake word init failed:', err);
                btn.classList.remove('active');
                status.classList.remove('visible');
                wakeWordActive = false;
                alert('Wake word detection failed to start. Check microphone permissions.');
            }
        }

        // ==================== Hands-free VAD Mode ====================

        let handsFreeActive = false;
        let vadInstance = null;
        let vadState = 'idle'; // idle, listen, hear, process, speak
        let vadProcessing = false;

        function cutOffAlfred() {
            if (currentAudio) {
                console.log('Interrupting Alfred');
                currentAudio.pause();
                currentAudio = null;
                const s = document.querySelector('.speak-btn.speaking');
                if (s) s.classList.remove('speaking');
                // Reset to listen state after interruption
                if (handsFreeActive) setVadState('listen');
            }
        }

        // Allow interrupting Alfred by pressing Escape or clicking the chat area
        document.addEventListener('keydown', (e) => {
            if (e.key === 'Escape' && currentAudio) {
                cutOffAlfred();
            }
        });
        document.getElementById('chat')?.addEventListener('click', (e) => {
            // Don't interrupt if clicking buttons
            if (e.target.closest('button')) return;
            if (currentAudio && vadState === 'speak') {
                cutOffAlfred();
            }
        });

        function setVadState(state) {
            vadState = state;
            const statusEl = document.getElementById('vad-status');
            const textEl = document.getElementById('vad-status-text');
            const micBtn = document.getElementById('mic-btn');
            micBtn.classList.remove('vad-listen','vad-hear','vad-process','vad-speak');
            statusEl.classList.remove('state-listen','state-hear','state-process','state-speak');
            if (state === 'idle' || !handsFreeActive) {
                statusEl.classList.remove('visible');
                return;
            }
            statusEl.classList.add('visible', 'state-' + state);
            micBtn.classList.add('vad-' + state);
            const labels = {listen:'Listening...', hear:'Hearing you...', process:'Processing...', speak:'<img src="/static/alfred-icon.jpg" alt="Alfred" class="inline-icon"> speaking...'};
            textEl.innerHTML = labels[state] || '';
        }

        function float32ToWav(samples, sampleRate) {
            const buf = new ArrayBuffer(44 + samples.length * 2);
            const v = new DataView(buf);
            function w(off, s) { for(let i=0;i<s.length;i++) v.setUint8(off+i,s.charCodeAt(i)); }
            w(0,'RIFF'); v.setUint32(4,36+samples.length*2,true);
            w(8,'WAVE'); w(12,'fmt '); v.setUint32(16,16,true);
            v.setUint16(20,1,true); v.setUint16(22,1,true);
            v.setUint32(24,sampleRate,true); v.setUint32(28,sampleRate*2,true);
            v.setUint16(32,2,true); v.setUint16(34,16,true);
            w(36,'data'); v.setUint32(40,samples.length*2,true);
            for(let i=0;i<samples.length;i++){
                const s=Math.max(-1,Math.min(1,samples[i]));
                v.setInt16(44+i*2, s<0?s*0x8000:s*0x7FFF, true);
            }
            return new Blob([buf],{type:'audio/wav'});
        }

        function playVadTTS(text) {
            setVadState('speak');
            // Explicitly ensure VAD is running so user can interrupt Alfred
            if (vadInstance) {
                try {
                    vadInstance.start();
                    console.log('VAD started for interrupt detection during speech');
                } catch(e) {
                    console.log('VAD start failed:', e);
                }
            } else {
                console.log('No VAD instance available');
            }
            fetch('/voice/speak', {
                method:'POST',
                headers: authHeaders({'Content-Type':'application/json'}),
                body: JSON.stringify({message: text})
            }).then(r => r.ok ? r.blob() : null).then(blob => {
                if (!blob || !handsFreeActive) {
                    if (handsFreeActive && vadInstance) vadInstance.start();
                    setVadState(handsFreeActive?'listen':'idle');
                    return;
                }
                const url = URL.createObjectURL(blob);
                currentAudio = new Audio(url);
                currentAudio.onplay = () => {
                    console.log('Alfred speaking - VAD state:', vadState, 'VAD running:', !!vadInstance);
                };
                currentAudio.onended = () => {
                    console.log('Alfred finished speaking');
                    currentAudio = null; URL.revokeObjectURL(url);
                    if (handsFreeActive) {
                        // Resume VAD after speech ends
                        if (vadInstance) vadInstance.start();
                        setVadState('listen');
                    }
                };
                currentAudio.onerror = () => {
                    currentAudio = null;
                    if (handsFreeActive) {
                        if (vadInstance) vadInstance.start();
                        setVadState('listen');
                    }
                };
                currentAudio.play();
            }).catch(() => {
                if (handsFreeActive) {
                    if (vadInstance) vadInstance.start();
                    setVadState('listen');
                }
            });
        }

        async function processVadAudio(audioFloat32) {
            if (vadProcessing) return;
            vadProcessing = true;
            setVadState('process');
            try {
                const wavBlob = float32ToWav(audioFloat32, 16000);
                const form = new FormData();
                form.append('audio', wavBlob, 'vad_recording.wav');
                const trResp = await fetch('/voice/transcribe', {method:'POST', headers:authHeaders(), body:form});
                const trData = await trResp.json();
                const text = (trData.text || '').trim();
                if (!text || text.length < 2) {
                    vadProcessing = false;
                    setVadState('listen');
                    return;
                }

                // Check for dismissal phrases
                const dismissalPhrases = [
                    "that's all for now", "thats all for now", "that is all for now",
                    "that's all", "thats all", "that is all",
                    "goodbye alfred", "goodbye", "good bye", "bye",
                    "thanks alfred", "thank you alfred", "thanks that's all",
                    "that will be all", "that'll be all", "thatll be all",
                    "i'm done", "im done", "i am done", "done for now",
                    "go to sleep", "you can go", "dismiss", "dismissed",
                    "see you later", "talk to you later", "later alfred",
                    "all for now", "nothing else", "no that's it", "no thats it"
                ];
                const textLower = text.toLowerCase().trim();
                const isDismissal = dismissalPhrases.some(phrase => textLower.includes(phrase));
                console.log('Transcribed:', textLower, '| Dismissal detected:', isDismissal);

                if (isDismissal) {
                    console.log('Dismissal phrase detected, ending hands-free mode');
                    addMsg(text, 'user');
                    const farewellMsg = "Very good, sir. I'm here if you need me.";
                    addMsg(farewellMsg, 'alfred', 'local', true);  // noAutoSpeak - we play it explicitly below
                    vadProcessing = false;

                    // Function to disable hands-free
                    function disableHandsFree() {
                        console.log('Disabling hands-free mode');
                        handsFreeActive = false;
                        document.getElementById('handsfree-btn').classList.remove('active');
                        if (vadInstance) vadInstance.pause();
                        setVadState('idle');
                        releaseWakeLock(); // Release if wake word also inactive
                        // Update wake word status if active
                        const wwStatus = document.getElementById('wakeword-status');
                        if (wwStatus && wakeWordActive) {
                            wwStatus.textContent = 'Listening for "Hey Alfred"...';
                            wwStatus.classList.add('visible');
                        }
                    }

                    // Play farewell TTS then disable hands-free
                    setVadState('speak');
                    fetch('/voice/speak', {
                        method:'POST',
                        headers: authHeaders({'Content-Type':'application/json'}),
                        body: JSON.stringify({message: farewellMsg})
                    }).then(r => r.ok ? r.blob() : null).then(blob => {
                        if (blob) {
                            const url = URL.createObjectURL(blob);
                            currentAudio = new Audio(url);
                            currentAudio.onended = () => {
                                currentAudio = null;
                                URL.revokeObjectURL(url);
                                disableHandsFree();
                            };
                            currentAudio.onerror = () => {
                                disableHandsFree();
                            };
                            currentAudio.play().catch(() => disableHandsFree());
                        } else {
                            disableHandsFree();
                        }
                    }).catch(() => {
                        disableHandsFree();
                    });
                    return;
                }

                // Check for voice control commands
                const autoSpeakOffPhrases = [
                    "turn off auto speak", "turn off autospeak", "turn off auto-speak",
                    "disable auto speak", "disable auto-speak", "disable autospeak",
                    "stop auto speak", "stop auto-speak", "mute", "mute yourself",
                    "stop speaking", "be quiet", "quiet mode", "silent mode",
                    "no more speaking", "stop talking"
                ];
                const autoSpeakOnPhrases = [
                    "turn on auto speak", "turn on autospeak", "turn on auto-speak",
                    "enable auto speak", "enable auto-speak", "enable autospeak",
                    "start auto speak", "unmute", "unmute yourself", "start speaking",
                    "speak mode", "voice mode", "talk to me", "start talking"
                ];
                const handsFreeOffPhrases = [
                    "turn off hands free", "turn off handsfree", "turn off hands-free",
                    "disable hands free", "disable hands-free", "disable handsfree",
                    "stop hands free", "stop listening", "pause listening"
                ];

                // Handle auto-speak toggle
                if (autoSpeakOffPhrases.some(phrase => textLower.includes(phrase))) {
                    addMsg(text, 'user');
                    autoSpeak = false;
                    localStorage.setItem('alfred_auto_speak', 'false');
                    document.getElementById('auto-speak-btn').classList.remove('active');
                    addMsg("Auto-speak disabled, sir. I'll stay quiet unless you ask me to speak.", 'alfred', 'local');
                    vadProcessing = false;
                    setVadState('listen');
                    return;
                }
                if (autoSpeakOnPhrases.some(phrase => textLower.includes(phrase))) {
                    addMsg(text, 'user');
                    autoSpeak = true;
                    localStorage.setItem('alfred_auto_speak', 'true');
                    document.getElementById('auto-speak-btn').classList.add('active');
                    const confirmMsg = "Auto-speak enabled, sir. I'll read my responses aloud.";
                    addMsg(confirmMsg, 'alfred', 'local');
                    playVadTTS(confirmMsg);
                    vadProcessing = false;
                    return;
                }
                // Handle hands-free off (similar to dismissal but explicit)
                if (handsFreeOffPhrases.some(phrase => textLower.includes(phrase))) {
                    addMsg(text, 'user');
                    const confirmMsg = "Hands-free mode disabled, sir.";
                    addMsg(confirmMsg, 'alfred', 'local');
                    vadProcessing = false;
                    handsFreeActive = false;
                    document.getElementById('handsfree-btn').classList.remove('active');
                    if (vadInstance) vadInstance.pause();
                    setVadState('idle');
                    releaseWakeLock(); // Release if wake word also inactive
                    return;
                }

                addMsg(text, 'user');
                if (!currentConversationId) {
                    try {
                        const cr = await fetch('/conversations', {method:'POST', headers:authHeaders()});
                        const cd = await cr.json();
                        currentConversationId = cd.id;
                        // If in project view, assign to project
                        if (currentProjectId) {
                            await fetch(`/conversations/${cd.id}/project`, {
                                method: 'PUT',
                                headers: authHeaders({'Content-Type': 'application/json'}),
                                body: JSON.stringify({project_id: currentProjectId})
                            });
                        }
                    } catch(e){}
                }

                // HYBRID APPROACH: Smart acknowledgment while processing
                // - Skips ack for conversational queries ("how are you", greetings)
                // - Plays ack for task queries ("search email", "check calendar")
                // - Skip for Qwen3 (different voice, adds latency)
                setVadState('speak');
                // Explicitly ensure VAD keeps running so user can interrupt Alfred
                if (vadInstance) {
                    try { vadInstance.start(); } catch(e) {}
                }
                let ackPromise = Promise.resolve();
                if (currentTtsBackend === 'kokoro') {
                    // Pass the query so backend can decide if ack is appropriate
                    ackPromise = fetch('/voice/chat/ack', {
                        method:'POST',
                        headers: authHeaders({'Content-Type':'application/json'}),
                        body: JSON.stringify({message: text})
                    }).then(r => r.ok ? r.blob() : null).then(blob => {
                        // Only play if we got audio (empty blob = no ack needed)
                        if (blob && blob.size > 100 && handsFreeActive) {
                            const url = URL.createObjectURL(blob);
                            const ackAudio = new Audio(url);
                            ackAudio.onended = () => URL.revokeObjectURL(url);
                            ackAudio.play().catch(() => {});
                        }
                    }).catch(() => {});
                }

                // Process with smart model in parallel
                const chatResp = await fetch('/chat', {
                    method:'POST',
                    headers: authHeaders({'Content-Type':'application/json'}),
                    body: JSON.stringify({message:text, session_id: currentConversationId||'default'})
                });
                const chatData = await chatResp.json();
                loadingHistory = true;
                addMsg(chatData.response, 'alfred', chatData.tier);
                loadingHistory = false;
                loadConversations();
                vadProcessing = false;

                // Wait a moment for ack to finish, then play full response
                await ackPromise;
                setTimeout(() => playVadTTS(chatData.response), 300);
            } catch(e) {
                console.error('Hands-free error:', e);
                vadProcessing = false;
                if (handsFreeActive) {
                    // Resume VAD on error
                    if (vadInstance) vadInstance.start();
                    setVadState('listen');
                }
            }
        }

        async function toggleHandsFree() {
            const btn = document.getElementById('handsfree-btn');
            if (handsFreeActive) {
                handsFreeActive = false;
                btn.classList.remove('active');
                if (vadInstance) vadInstance.pause();
                cutOffAlfred();
                setVadState('idle');
                releaseWakeLock(); // Release if wake word also inactive
                return;
            }
            handsFreeActive = true;
            btn.classList.add('active');
            await acquireWakeLock(); // Keep screen on
            if (!autoSpeak) {
                autoSpeak = true;
                localStorage.setItem('alfred_auto_speak', 'true');
                document.getElementById('auto-speak-btn').classList.add('active');
            }
            cutOffAlfred();
            try {
                if (!vadInstance) {
                    setVadState('process');
                    vadInstance = await vad.MicVAD.new({
                        positiveSpeechThreshold: 0.5,  // Lower threshold for easier interruption
                        negativeSpeechThreshold: 0.35,
                        minSpeechFrames: 2,  // Faster detection
                        preSpeechPadFrames: 6,
                        redemptionFrames: 4,
                        // Disable echo cancellation with headphones - it can filter user voice
                        additionalAudioConstraints: {
                            echoCancellation: false,
                            noiseSuppression: false,
                            autoGainControl: true
                        },
                        onSpeechStart: () => {
                            console.log('VAD onSpeechStart - state:', vadState, 'hasAudio:', !!currentAudio);
                            if (!handsFreeActive) return;
                            // Interrupt Alfred if he's speaking
                            if (currentAudio && vadState === 'speak') {
                                console.log('Interrupting Alfred via voice!');
                                cutOffAlfred();
                                return; // Don't process this speech (likely just interruption)
                            }
                            if (!vadProcessing) setVadState('hear');
                        },
                        onSpeechEnd: (audio) => {
                            if (!handsFreeActive || vadProcessing) return;
                            processVadAudio(audio);
                        },
                        onVADMisfire: () => {
                            if (handsFreeActive && vadState === 'hear') setVadState('listen');
                        }
                    });
                }
                vadInstance.start();
                setVadState('listen');
            } catch(e) {
                console.error('VAD init failed:', e);
                handsFreeActive = false;
                btn.classList.remove('active');
                setVadState('idle');
                alert('Hands-free failed to start. Check mic permissions.');
            }
        }

        // Matrix Rain Effect
        (function initMatrix() {
            const canvas = document.getElementById('matrix-canvas');
            if (!canvas) return;
            const ctx = canvas.getContext('2d');

            let width, height, columns, drops;
            const chars = '01';
            const fontSize = 14;

            function resize() {
                width = canvas.width = window.innerWidth;
                height = canvas.height = window.innerHeight;
                columns = Math.floor(width / fontSize);
                drops = Array(columns).fill(1);
            }

            function draw() {
                ctx.fillStyle = 'rgba(0, 0, 0, 0.05)';
                ctx.fillRect(0, 0, width, height);

                ctx.fillStyle = '#00ff41';
                ctx.font = fontSize + 'px monospace';

                for (let i = 0; i < drops.length; i++) {
                    const char = chars[Math.floor(Math.random() * chars.length)];
                    const x = i * fontSize;
                    const y = drops[i] * fontSize;

                    // Varying brightness
                    const brightness = Math.random();
                    if (brightness > 0.9) {
                        ctx.fillStyle = '#fff';
                    } else if (brightness > 0.7) {
                        ctx.fillStyle = '#00ff41';
                    } else {
                        ctx.fillStyle = 'rgba(0, 255, 65, 0.5)';
                    }

                    ctx.fillText(char, x, y);

                    if (y > height && Math.random() > 0.975) {
                        drops[i] = 0;
                    }
                    drops[i]++;
                }
            }

            resize();
            window.addEventListener('resize', resize);
            setInterval(draw, 50);
        })();

        // Register service worker for PWA with update handling
        if ('serviceWorker' in navigator) {
            navigator.serviceWorker.register('/sw.js').then(reg => {
                console.log('SW registered:', reg.scope);

                // Check for updates on load and when user returns to tab (minimal CPU impact)
                reg.update();
                document.addEventListener('visibilitychange', () => {
                    if (document.visibilityState === 'visible') reg.update();
                });

                // Handle updates
                reg.addEventListener('updatefound', () => {
                    const newWorker = reg.installing;
                    console.log('SW update found, installing...');

                    newWorker.addEventListener('statechange', () => {
                        if (newWorker.state === 'installed' && navigator.serviceWorker.controller) {
                            // New version available, notify user or auto-reload
                            console.log('New version available, reloading...');
                            newWorker.postMessage('SKIP_WAITING');
                        }
                    });
                });
            }).catch(err => console.log('SW registration failed:', err));

            // Listen for SW_UPDATED message and reload
            navigator.serviceWorker.addEventListener('message', event => {
                if (event.data && event.data.type === 'SW_UPDATED') {
                    console.log('SW updated, reloading page...');
                    window.location.reload();
                }
            });

            // Handle controller change (new SW took over)
            navigator.serviceWorker.addEventListener('controllerchange', () => {
                console.log('SW controller changed, reloading...');
                window.location.reload();
            });
        }

        // ==================== Push Notifications WebSocket ====================
        let notificationWs = null;
        let notificationReconnectAttempts = 0;
        const MAX_RECONNECT_ATTEMPTS = 5;

        function connectNotifications() {
            const token = localStorage.getItem('alfred_token');
            if (!token) return;

            const wsProtocol = window.location.protocol === 'https:' ? 'wss:' : 'ws:';
            notificationWs = new WebSocket(`${wsProtocol}//${window.location.host}/ws/notifications?token=${token}`);

            notificationWs.onopen = () => {
                console.log('Notification WebSocket connected');
                notificationReconnectAttempts = 0;
                // Keep-alive ping every 30 seconds
                notificationWs._pingInterval = setInterval(() => {
                    if (notificationWs && notificationWs.readyState === WebSocket.OPEN) {
                        notificationWs.send('ping');
                    }
                }, 30000);
            };

            notificationWs.onmessage = (event) => {
                try {
                    const data = JSON.parse(event.data);
                    handleNotification(data);
                } catch (e) {
                    // Ignore pong responses
                    if (event.data !== 'pong') {
                        console.log('Notification:', event.data);
                    }
                }
            };

            notificationWs.onclose = () => {
                console.log('Notification WebSocket closed');
                if (notificationWs && notificationWs._pingInterval) {
                    clearInterval(notificationWs._pingInterval);
                }
                notificationWs = null;
                // Reconnect with backoff
                if (notificationReconnectAttempts < MAX_RECONNECT_ATTEMPTS) {
                    const delay = Math.min(1000 * Math.pow(2, notificationReconnectAttempts), 30000);
                    notificationReconnectAttempts++;
                    setTimeout(connectNotifications, delay);
                }
            };

            notificationWs.onerror = (err) => {
                console.error('Notification WebSocket error:', err);
            };
        }

        function handleNotification(data) {
            console.log('Received notification:', data);

            // Handle different notification types
            switch (data.type) {
                case 'agent_completed':
                    showAgentNotification(data.data, 'completed');
                    break;
                case 'agent_failed':
                    showAgentNotification(data.data, 'failed');
                    break;
                case 'agent_started':
                    // Optional: show subtle indicator
                    console.log(`Agent started: ${data.data.agent_type} - ${data.data.goal}`);
                    break;
                case 'long_processing':
                    showLongProcessingNotification(data.data);
                    break;
                case 'connected':
                    console.log('Notification channel ready');
                    break;
                default:
                    console.log('Unknown notification type:', data.type);
            }
        }

        function showAgentNotification(data, status) {
            // Create toast notification
            const toast = document.createElement('div');
            toast.className = 'agent-notification';
            toast.style.cssText = `
                position: fixed;
                top: 20px;
                right: 20px;
                background: ${status === 'completed' ? '#1a472a' : '#4a1a1a'};
                border: 1px solid ${status === 'completed' ? '#2d6a4f' : '#6a2d2d'};
                border-radius: 12px;
                padding: 16px 20px;
                max-width: 400px;
                z-index: 10000;
                box-shadow: 0 4px 20px rgba(0,0,0,0.4);
                animation: slideIn 0.3s ease-out;
            `;

            const icon = status === 'completed' ? '✓' : '✗';
            const title = status === 'completed' ? 'Agent Completed' : 'Agent Failed';
            const color = status === 'completed' ? '#4ade80' : '#f87171';

            toast.innerHTML = `
                <div style="display: flex; align-items: flex-start; gap: 12px;">
                    <span style="font-size: 24px; color: ${color};">${icon}</span>
                    <div style="flex: 1;">
                        <div style="font-weight: 600; color: ${color}; margin-bottom: 4px;">${title}</div>
                        <div style="font-size: 13px; color: #9ca3af; margin-bottom: 8px;">
                            ${data.agent_type} agent ${status === 'completed' ? `finished in ${data.duration_seconds?.toFixed(1) || '?'}s` : 'encountered an error'}
                        </div>
                        <div style="font-size: 14px; color: #e0e0e0; line-height: 1.4;">
                            ${status === 'completed' ? (data.result_preview || '').substring(0, 200) : data.error}
                            ${(data.result_preview || '').length > 200 ? '...' : ''}
                        </div>
                        <div style="margin-top: 8px; font-size: 12px; color: #6b7280;">
                            Task ID: ${data.task_id}
                        </div>
                    </div>
                    <button onclick="this.parentElement.parentElement.remove()" style="
                        background: none; border: none; color: #9ca3af; cursor: pointer;
                        font-size: 18px; padding: 0; line-height: 1;
                    ">×</button>
                </div>
            `;

            // Add animation keyframes if not already added
            if (!document.getElementById('notification-styles')) {
                const style = document.createElement('style');
                style.id = 'notification-styles';
                style.textContent = `
                    @keyframes slideIn {
                        from { transform: translateX(100%); opacity: 0; }
                        to { transform: translateX(0); opacity: 1; }
                    }
                    @keyframes slideOut {
                        from { transform: translateX(0); opacity: 1; }
                        to { transform: translateX(100%); opacity: 0; }
                    }
                `;
                document.head.appendChild(style);
            }

            document.body.appendChild(toast);

            // Auto-remove after 15 seconds
            setTimeout(() => {
                toast.style.animation = 'slideOut 0.3s ease-in forwards';
                setTimeout(() => toast.remove(), 300);
            }, 15000);
        }

        function showLongProcessingNotification(data) {
            // In-page toast notification (blue theme)
            const toast = document.createElement('div');
            toast.className = 'long-processing-notification';
            toast.style.cssText = `
                position: fixed;
                top: 20px;
                right: 20px;
                background: #1a2744;
                border: 1px solid #2d4a6f;
                border-radius: 12px;
                padding: 16px 20px;
                max-width: 400px;
                z-index: 10000;
                box-shadow: 0 4px 20px rgba(0,0,0,0.4);
                animation: slideIn 0.3s ease-out;
            `;

            const elapsed = Math.round(data.elapsed_seconds || 60);
            toast.innerHTML = `
                <div style="display: flex; align-items: flex-start; gap: 12px;">
                    <span style="font-size: 24px; color: #60a5fa;">&#9202;</span>
                    <div style="flex: 1;">
                        <div style="font-weight: 600; color: #60a5fa; margin-bottom: 4px;">Still Working...</div>
                        <div style="font-size: 13px; color: #9ca3af; margin-bottom: 8px;">
                            Alfred has been processing for ${elapsed}s
                        </div>
                        <div style="font-size: 14px; color: #e0e0e0; line-height: 1.4;">
                            ${data.query_preview || 'Your request is still being processed.'}
                        </div>
                    </div>
                    <button onclick="this.parentElement.parentElement.remove()" style="
                        background: none; border: none; color: #9ca3af; cursor: pointer;
                        font-size: 18px; padding: 0; line-height: 1;
                    ">&times;</button>
                </div>
            `;

            // Add animation keyframes if not already added
            if (!document.getElementById('notification-styles')) {
                const style = document.createElement('style');
                style.id = 'notification-styles';
                style.textContent = `
                    @keyframes slideIn {
                        from { transform: translateX(100%); opacity: 0; }
                        to { transform: translateX(0); opacity: 1; }
                    }
                    @keyframes slideOut {
                        from { transform: translateX(0); opacity: 1; }
                        to { transform: translateX(100%); opacity: 0; }
                    }
                `;
                document.head.appendChild(style);
            }

            document.body.appendChild(toast);

            // Also fire a native Notification if tab is hidden (and push not available)
            if (document.hidden && Notification.permission === 'granted') {
                try {
                    new Notification('Alfred is still working...', {
                        body: data.query_preview || 'Your request is taking longer than usual.',
                        icon: '/static/icon-192.png',
                        tag: 'alfred-long-processing',
                    });
                } catch (e) { /* Native notifications not supported */ }
            }

            // Auto-remove after 30 seconds
            setTimeout(() => {
                toast.style.animation = 'slideOut 0.3s ease-in forwards';
                setTimeout(() => toast.remove(), 300);
            }, 30000);
        }

        // ==================== Push Notification Subscription ====================
        async function subscribeToPush() {
            if (!('serviceWorker' in navigator) || !('PushManager' in window)) {
                console.log('Push notifications not supported');
                return;
            }

            const permission = await Notification.requestPermission();
            if (permission !== 'granted') {
                console.log('Push notification permission denied');
                return;
            }

            try {
                const reg = await navigator.serviceWorker.ready;
                const token = localStorage.getItem('alfred_token');

                // Get VAPID public key from server
                const keyResp = await fetch('/push/vapid-key', {
                    headers: { 'Authorization': 'Bearer ' + token }
                });
                const { publicKey } = await keyResp.json();
                if (!publicKey) {
                    console.log('No VAPID key configured on server');
                    return;
                }

                // Convert URL-safe base64 to Uint8Array
                const padding = '='.repeat((4 - publicKey.length % 4) % 4);
                const base64 = (publicKey + padding).replace(/-/g, '+').replace(/_/g, '/');
                const rawKey = Uint8Array.from(atob(base64), c => c.charCodeAt(0));

                // Subscribe to push
                const subscription = await reg.pushManager.subscribe({
                    userVisibleOnly: true,
                    applicationServerKey: rawKey,
                });

                // Send subscription to server
                const subJson = subscription.toJSON();
                await fetch('/push/subscribe', {
                    method: 'POST',
                    headers: {
                        'Content-Type': 'application/json',
                        'Authorization': 'Bearer ' + token,
                    },
                    body: JSON.stringify({
                        endpoint: subJson.endpoint,
                        keys: subJson.keys,
                    }),
                });

                console.log('Push notifications subscribed successfully');
            } catch (err) {
                console.error('Push subscription failed:', err);
            }
        }

        // Connect to notifications when page loads (if authenticated)
        if (localStorage.getItem('alfred_token')) {
            connectNotifications();
            subscribeToPush();
        }
    </script>
</body>
</html>"""


# Generate dynamic cache version based on app version and startup time
import time as _time
_BUILD_TIMESTAMP = str(int(_time.time()))

def _get_service_worker_js():
    """Generate service worker with dynamic cache version."""
    return f"""
const CACHE_NAME = 'alfred-v030-{_BUILD_TIMESTAMP}';
const PRECACHE_URLS = ['/', '/manifest.json', '/static/icon-192.png', '/static/icon-512.png', '/static/logo-white.png', '/static/glabs.jpeg', '/static/alfred-icon.jpg'];

// Force install new service worker immediately
self.addEventListener('install', event => {{
    console.log('[SW] Installing new version:', CACHE_NAME);
    event.waitUntil(
        caches.open(CACHE_NAME)
            .then(cache => cache.addAll(PRECACHE_URLS))
            .then(() => self.skipWaiting())  // Skip waiting, activate immediately
    );
}});

// Clean up old caches and take control of all pages
self.addEventListener('activate', event => {{
    console.log('[SW] Activating new version:', CACHE_NAME);
    event.waitUntil(
        caches.keys().then(keys => {{
            // Delete ALL old caches (any cache not matching current version)
            return Promise.all(
                keys.filter(k => k !== CACHE_NAME).map(k => {{
                    console.log('[SW] Deleting old cache:', k);
                    return caches.delete(k);
                }})
            );
        }}).then(() => {{
            // Take control of all pages immediately
            return self.clients.claim();
        }}).then(() => {{
            // Notify all clients to reload
            return self.clients.matchAll().then(clients => {{
                clients.forEach(client => client.postMessage({{type: 'SW_UPDATED'}}));
            }});
        }})
    );
}});

// Message handler for manual cache clearing
self.addEventListener('message', event => {{
    if (event.data === 'SKIP_WAITING') {{
        self.skipWaiting();
    }}
    if (event.data === 'CLEAR_CACHE') {{
        caches.keys().then(keys => Promise.all(keys.map(k => caches.delete(k))));
    }}
}});

self.addEventListener('fetch', event => {{
    const url = new URL(event.request.url);

    // NEVER cache API calls - always network first
    if (url.pathname.startsWith('/chat') || url.pathname.startsWith('/auth') ||
        url.pathname.startsWith('/conversations') || url.pathname.startsWith('/voice') ||
        url.pathname.startsWith('/integrations') || url.pathname.startsWith('/memory') ||
        url.pathname.startsWith('/projects') || url.pathname.startsWith('/references') ||
        url.pathname.startsWith('/health') || url.pathname.startsWith('/settings') ||
        url.pathname.startsWith('/push') || url.pathname.startsWith('/notifications') ||
        url.pathname.startsWith('/api')) {{
        event.respondWith(fetch(event.request).catch(() => caches.match(event.request)));
    }}
    // For main HTML page - network first with cache fallback (always get latest)
    else if (url.pathname === '/' || url.pathname.endsWith('.html')) {{
        event.respondWith(
            fetch(event.request)
                .then(resp => {{
                    const clone = resp.clone();
                    caches.open(CACHE_NAME).then(cache => cache.put(event.request, clone));
                    return resp;
                }})
                .catch(() => caches.match(event.request))
        );
    }}
    // Static assets can be cache-first
    else {{
        event.respondWith(
            caches.match(event.request).then(cached => cached || fetch(event.request).then(resp => {{
                if (resp.ok) {{
                    const clone = resp.clone();
                    caches.open(CACHE_NAME).then(cache => cache.put(event.request, clone));
                }}
                return resp;
            }}))
        );
    }}
}});

// ==================== Web Push Handlers ====================

self.addEventListener('push', event => {{
    let data = {{}};
    try {{
        data = event.data ? event.data.json() : {{}};
    }} catch (e) {{
        data = {{ title: 'Alfred', body: event.data ? event.data.text() : 'New notification' }};
    }}

    const title = data.title || 'Alfred';
    const options = {{
        body: data.body || 'Alfred has a notification for you.',
        icon: '/static/icon-192.png',
        badge: '/static/icon-192.png',
        tag: 'alfred-long-processing',
        renotify: false,
        data: {{ url: data.url || '/' }},
    }};

    event.waitUntil(self.registration.showNotification(title, options));
}});

self.addEventListener('notificationclick', event => {{
    event.notification.close();
    const url = event.notification.data?.url || '/';

    event.waitUntil(
        clients.matchAll({{ type: 'window', includeUncontrolled: true }}).then(windowClients => {{
            // Focus existing Alfred tab if open
            for (const client of windowClients) {{
                if (client.url.includes(self.location.origin) && 'focus' in client) {{
                    return client.focus();
                }}
            }}
            // Otherwise open a new tab
            return clients.openWindow(url);
        }})
    );
}});
""".strip()

SERVICE_WORKER_JS = _get_service_worker_js()


# Mount frontend build assets (must be after all API routes)
if _frontend_dist.exists():
    _frontend_assets = _frontend_dist / "assets"
    if _frontend_assets.exists():
        app.mount("/assets", StaticFiles(directory=str(_frontend_assets)), name="frontend_assets")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host=settings.host, port=settings.port)

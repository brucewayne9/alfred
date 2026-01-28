"""File handling for document parsing and creation."""

import csv
import io
import json
import logging
import os
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)

# Upload directory
UPLOAD_DIR = Path(__file__).parent.parent.parent / "data" / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

GENERATED_DIR = Path(__file__).parent.parent.parent / "data" / "generated"
GENERATED_DIR.mkdir(parents=True, exist_ok=True)


def parse_document(file_path: str) -> dict:
    """Parse a document and extract its text content.

    Supports: PDF, TXT, MD, CSV, DOC/DOCX, XLS/XLSX, JSON
    Returns: {"text": str, "metadata": dict, "error": str|None}
    """
    path = Path(file_path)
    if not path.exists():
        return {"text": "", "metadata": {}, "error": f"File not found: {file_path}"}

    ext = path.suffix.lower()
    metadata = {
        "filename": path.name,
        "extension": ext,
        "size_bytes": path.stat().st_size,
    }

    try:
        if ext == ".pdf":
            return _parse_pdf(path, metadata)
        elif ext in (".txt", ".md", ".markdown"):
            return _parse_text(path, metadata)
        elif ext == ".csv":
            return _parse_csv(path, metadata)
        elif ext in (".doc", ".docx"):
            return _parse_docx(path, metadata)
        elif ext in (".xls", ".xlsx"):
            return _parse_xlsx(path, metadata)
        elif ext == ".json":
            return _parse_json(path, metadata)
        else:
            return {"text": "", "metadata": metadata, "error": f"Unsupported format: {ext}"}
    except Exception as e:
        logger.error(f"Error parsing {file_path}: {e}")
        return {"text": "", "metadata": metadata, "error": str(e)}


def _parse_pdf(path: Path, metadata: dict) -> dict:
    """Parse PDF file."""
    from PyPDF2 import PdfReader
    reader = PdfReader(str(path))
    metadata["pages"] = len(reader.pages)
    text_parts = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            text_parts.append(f"[Page {i+1}]\n{text}")
    return {"text": "\n\n".join(text_parts), "metadata": metadata, "error": None}


def _parse_text(path: Path, metadata: dict) -> dict:
    """Parse plain text or markdown file."""
    text = path.read_text(encoding="utf-8", errors="replace")
    metadata["lines"] = text.count("\n") + 1
    return {"text": text, "metadata": metadata, "error": None}


def _parse_csv(path: Path, metadata: dict) -> dict:
    """Parse CSV file."""
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        rows = list(reader)
    metadata["rows"] = len(rows)
    metadata["columns"] = len(rows[0]) if rows else 0
    # Format as readable text
    if rows:
        header = rows[0] if rows else []
        text_parts = [f"Columns: {', '.join(header)}"]
        for i, row in enumerate(rows[1:21], 1):  # First 20 data rows
            text_parts.append(f"Row {i}: {', '.join(row)}")
        if len(rows) > 21:
            text_parts.append(f"... and {len(rows) - 21} more rows")
        text = "\n".join(text_parts)
    else:
        text = "(empty CSV)"
    return {"text": text, "metadata": metadata, "error": None}


def _parse_docx(path: Path, metadata: dict) -> dict:
    """Parse Word document."""
    from docx import Document
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    metadata["paragraphs"] = len(paragraphs)
    # Also extract tables
    tables_text = []
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            tables_text.append(f"[Table {i+1}]\n" + "\n".join(rows))
    text = "\n\n".join(paragraphs)
    if tables_text:
        text += "\n\n" + "\n\n".join(tables_text)
    return {"text": text, "metadata": metadata, "error": None}


def _parse_xlsx(path: Path, metadata: dict) -> dict:
    """Parse Excel spreadsheet."""
    from openpyxl import load_workbook
    wb = load_workbook(str(path), read_only=True, data_only=True)
    metadata["sheets"] = wb.sheetnames
    text_parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(max_row=50, values_only=True))  # First 50 rows
        if rows:
            text_parts.append(f"[Sheet: {sheet_name}]")
            for i, row in enumerate(rows[:25], 1):
                cells = [str(c) if c is not None else "" for c in row]
                text_parts.append(f"Row {i}: {', '.join(cells)}")
            if len(rows) > 25:
                text_parts.append(f"... and more rows")
    wb.close()
    return {"text": "\n".join(text_parts), "metadata": metadata, "error": None}


def _parse_json(path: Path, metadata: dict) -> dict:
    """Parse JSON file."""
    data = json.loads(path.read_text(encoding="utf-8"))
    text = json.dumps(data, indent=2)[:10000]  # Limit to 10k chars
    if len(json.dumps(data)) > 10000:
        text += "\n... (truncated)"
    return {"text": text, "metadata": metadata, "error": None}


def create_document(content: str, filename: str, format: str = "txt") -> dict:
    """Create a document from content.

    Args:
        content: The text content to put in the document
        filename: Base filename (without extension)
        format: Output format (txt, md, csv, pdf, docx, xlsx, json)

    Returns: {"path": str, "filename": str, "error": str|None}
    """
    format = format.lower().lstrip(".")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = "".join(c for c in filename if c.isalnum() or c in "._- ")[:50]
    out_filename = f"{safe_name}_{timestamp}.{format}"
    out_path = GENERATED_DIR / out_filename

    try:
        if format == "txt":
            out_path.write_text(content, encoding="utf-8")
        elif format == "md":
            out_path.write_text(content, encoding="utf-8")
        elif format == "json":
            # Try to parse as JSON, otherwise wrap as string
            try:
                data = json.loads(content)
            except json.JSONDecodeError:
                data = {"content": content}
            out_path.write_text(json.dumps(data, indent=2), encoding="utf-8")
        elif format == "csv":
            _create_csv(content, out_path)
        elif format == "pdf":
            _create_pdf(content, out_path)
        elif format == "docx":
            _create_docx(content, out_path)
        elif format == "xlsx":
            _create_xlsx(content, out_path)
        else:
            return {"path": "", "filename": "", "error": f"Unsupported format: {format}"}

        return {"path": str(out_path), "filename": out_filename, "error": None}
    except Exception as e:
        logger.error(f"Error creating {format} document: {e}")
        return {"path": "", "filename": "", "error": str(e)}


def _create_csv(content: str, out_path: Path):
    """Create CSV from content (expects rows separated by newlines, columns by comma or tab)."""
    lines = content.strip().split("\n")
    with open(out_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        for line in lines:
            if "\t" in line:
                row = line.split("\t")
            elif "," in line:
                row = [c.strip() for c in line.split(",")]
            else:
                row = [line]
            writer.writerow(row)


def _create_pdf(content: str, out_path: Path):
    """Create PDF from content."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    doc = SimpleDocTemplate(str(out_path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    for para in content.split("\n\n"):
        if para.strip():
            # Handle line breaks within paragraphs
            text = para.replace("\n", "<br/>")
            story.append(Paragraph(text, styles["Normal"]))
            story.append(Spacer(1, 12))

    if not story:
        story.append(Paragraph(content.replace("\n", "<br/>"), styles["Normal"]))

    doc.build(story)


def _create_docx(content: str, out_path: Path):
    """Create Word document from content."""
    from docx import Document
    doc = Document()
    for para in content.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para)
    doc.save(str(out_path))


def _create_xlsx(content: str, out_path: Path):
    """Create Excel spreadsheet from content."""
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"

    lines = content.strip().split("\n")
    for i, line in enumerate(lines, 1):
        if "\t" in line:
            cells = line.split("\t")
        elif "," in line:
            cells = [c.strip() for c in line.split(",")]
        else:
            cells = [line]
        for j, cell in enumerate(cells, 1):
            ws.cell(row=i, column=j, value=cell)

    wb.save(str(out_path))


def save_upload(file_content: bytes, original_filename: str) -> str:
    """Save uploaded file and return the path."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    ext = Path(original_filename).suffix
    safe_name = "".join(c for c in Path(original_filename).stem if c.isalnum() or c in "._- ")[:30]
    filename = f"{safe_name}_{timestamp}{ext}"
    path = UPLOAD_DIR / filename
    path.write_bytes(file_content)
    return str(path)
